# -*- coding: utf-8 -*-
# Tác giả: Nguyễn Vũ Hùng — Cục Quản lý Dược — 2026
"""
Tự động soạn + lưu NHÁP (hoặc TRÌNH thẳng) phiếu trình trên emoh.moh.gov.vn
- KHÔNG đăng nhập: dùng lại cookie phiên từ trình duyệt (bạn dán vào).
- Khung Xem trước có 2 nút: "Lưu dự thảo" (sign=0, mặc định, an toàn — chỉ lưu, không đi vào
  luồng ký duyệt) và "Trình văn bản" (sign=1 — trình thật, văn bản bắt đầu đi vào luồng ký
  duyệt ngay). Chọn "Lưu dự thảo" thì tự mở thùng nháp trên web, xem lại, tự bấm Trình sau.

Cần: Python 3.9+, thư viện requests  ->  pip install requests
Đặt 3 file cùng thư mục: trinh_van_ban.py, du_lieu.json, noi_nhan.json
Chạy:  python trinh_van_ban.py
"""
import atexit, base64, json, logging, os, re, shutil, subprocess, sys, tempfile, textwrap
import threading, time, traceback, unicodedata
from html import unescape as html_unescape   # tên "html" đã dùng làm biến cục bộ khắp file (nội dung
                                              # trang) — import tách riêng để khỏi đụng nhau
from contextlib import contextmanager
from functools import lru_cache
from datetime import datetime, timedelta
from logging.handlers import TimedRotatingFileHandler
from urllib.parse import quote, unquote, urljoin
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter import font as tkfont

try:
    import requests
except ImportError:
    raise SystemExit("Chưa có thư viện 'requests'. Chạy: pip install requests")

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None   # vẫn chạy được, chỉ tắt tính năng tự đọc PDF

try:
    import fitz   # PyMuPDF — dùng để tìm toạ độ chữ ký + ghi chú thích số thứ tự
except ImportError:
    fitz = None

try:
    import docx2pdf   # chuyển .docx -> .pdf (điều khiển Word cài sẵn trên máy) — không hỗ trợ .doc
except ImportError:
    docx2pdf = None

try:
    import docx   # python-docx — đọc thẳng .docx (không cần Word) để tự điền NGAY trong lúc
                   # đang chờ Word chuyển .docx -> PDF chạy song song ở luồng khác (xem
                   # convert_office_doc_to_pdf) — trước đây phải đợi PDF xong mới tách được chữ
except ImportError:
    docx = None

AUTHOR_MARK = "Nguyễn Vũ Hùng  ·  Cục Quản lý Dược  ·  2026"

BASE = "https://emoh.moh.gov.vn"
CAS = "https://emoh.moh.gov.vn:8443/passportv3"
LOGIN_URL = CAS + "/login"
CAPTCHA_URL = CAS + "/jcaptcha.jpg"
SERVICE = "https://emoh.moh.gov.vn"
# Đóng gói bằng PyInstaller (onefile) thì __file__ trỏ vào thư mục giải nén tạm (_MEIPASS),
# không phải nơi đặt file .exe — phải dùng sys.executable để tìm đúng thư mục chứa du_lieu.json,
# noi_nhan.json... (và để settings.json/id_cache.json/nguoi_dung.json ghi đúng chỗ, không mất
# khi thư mục tạm bị dọn sau khi tắt chương trình).
if getattr(sys, "frozen", False):
    HERE = os.path.dirname(os.path.abspath(sys.executable))
    # Trên macOS, sys.executable nằm trong TênApp.app/Contents/MacOS/ — nếu để nguyên thì các
    # file .json phải nhét vào tận trong ruột app (phải bấm "Show Package Contents" mới thấy),
    # bất tiện khi copy nguyên thư mục sang máy khác. Đi lên khỏi .app để dữ liệu nằm CẠNH
    # file .app trong Finder, giống hệt cách .exe + .json nằm cạnh nhau trên Windows.
    if sys.platform == "darwin" and "/Contents/MacOS" in HERE:
        HERE = os.path.dirname(os.path.dirname(os.path.dirname(HERE)))
else:
    HERE = os.path.dirname(os.path.abspath(__file__))

# ---------- Ghi log ra file (phục vụ tra lỗi khi dùng lâu dài, không ai ngồi cạnh theo dõi) ----
# CHỈ ghi CỤC BỘ trên máy — không tự gửi đi đâu cả (app xử lý nội dung phiếu trình có thể nhạy
# cảm, không phù hợp kiểu "tự động gửi báo lỗi về server hãng" như app thương mại hay làm). Khi
# có sự cố, tự tìm đúng file app_log.txt (cùng thư mục settings.json/nguoi_dung.json) gửi lại để
# tra — xem HUONG_DAN.md. Giữ tối đa 14 ngày gần nhất (TimedRotatingFileHandler tự xoay/xoá file
# cũ), không phình vô hạn theo thời gian dùng.
LOG_FILE = os.path.join(HERE, "app_log.txt")
APP_VERSION = "2026-08-21"

def _setup_file_logger():
    logger = logging.getLogger("trinh_van_ban")
    logger.setLevel(logging.DEBUG)
    handler = TimedRotatingFileHandler(LOG_FILE, when="midnight", backupCount=14,
                                        encoding="utf-8", delay=True)
    handler.setFormatter(logging.Formatter("%(asctime)s  %(levelname)-7s  %(message)s"))
    logger.addHandler(handler)
    return logger

_file_logger = _setup_file_logger()
_file_logger.info(f"=== Khởi động chương trình — bản {APP_VERSION} ===")

def _log_uncaught_exception(exc_type, exc_value, exc_tb):
    """Lưới hứng cuối cho lỗi KHÔNG lường trước ở luồng chính (ngoài mọi try/except đã có sẵn
    trong code) — ghi đầy đủ traceback vào app_log.txt trước khi hành vi mặc định (in ra
    stderr, thường không ai thấy vì app đóng gói không có cửa sổ console) tiếp tục chạy."""
    _file_logger.error("LỖI KHÔNG BẮT ĐƯỢC (luồng chính):\n"
                        + "".join(traceback.format_exception(exc_type, exc_value, exc_tb)))
    sys.__excepthook__(exc_type, exc_value, exc_tb)
sys.excepthook = _log_uncaught_exception

def _log_thread_exception(args):
    """Tương tự _log_uncaught_exception nhưng cho lỗi lọt ra ngoài 1 threading.Thread nền (mỗi
    worker trong code đã tự try/except riêng rồi, đây chỉ là lưới hứng cuối phòng khi sót)."""
    _file_logger.error(f"LỖI KHÔNG BẮT ĐƯỢC (luồng nền '{args.thread.name}'):\n"
                        + "".join(traceback.format_exception(args.exc_type, args.exc_value,
                                                               args.exc_traceback)))
threading.excepthook = _log_thread_exception

# Header kiểu trình duyệt (điều hướng, KHÔNG phải XHR) — dùng cho login & upload
BROWSER_HEADERS = {
    "X-Requested-With": None,
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "same-origin",
}

# ---------- Nạp dữ liệu ----------
def load_json(name):
    with open(os.path.join(HERE, name), encoding="utf-8") as f:
        return json.load(f)

DATA = load_json("du_lieu.json")
BOOK = load_json("noi_nhan.json")
try:
    CAY = load_json("cay_don_vi.json")     # cây đầy đủ: internal + lien_thong, path neo gốc
except Exception:
    CAY = {"internal": {"root_id": "52", "nodes": []}, "lien_thong": {"nodes": []}}
ENUMS = DATA["enums"]
TPL = DATA["templates"]
# Tên loại văn bản viết HOA (khớp dòng tiêu đề in trên trang 1, VD "QUYẾT ĐỊNH") -> tên gốc
# trong ENUMS["documentType"] (VD "Quyết định"), dùng để tự nhận loại VB từ file PDF.
DOC_TYPE_TITLES = {name.upper(): name for name in ENUMS["documentType"]}

FLOW_STORE_FILE = os.path.join(HERE, "luong_trinh.json")

def _migrate_old_flow_list(old):
    """Định dạng cũ: {"flows": [{"name": "Luồng Cục"/"Luồng Phòng"/"Luồng Bộ", "flowId": ...}]}
    — 3 luồng CỨNG dùng chung cho MỌI người chạy chương trình này, dù đó là 3 luồng CÁ NHÂN chỉ
    của riêng người tạo file. Chuyển sang ghim + quy tắc cục bộ (giữ nguyên hành vi hiện tại cho
    máy này, nhưng máy người khác sẽ có sổ ghim RIÊNG, rỗng lúc mới cài — xem thảo luận)."""
    by_name = {fl.get("name"): fl.get("flowId") for fl in old.get("flows", [])}
    bo, cuc, phong = by_name.get("Luồng Bộ"), by_name.get("Luồng Cục"), by_name.get("Luồng Phòng")
    rules = []
    if bo:   rules.append({"keyword": "byt", "flowId": bo})
    if cuc:  rules.append({"keyword": "qld", "flowId": cuc})
    if phong: rules.append({"keyword": "cl", "flowId": phong})
    return {
        "pinned": [x for x in (bo, cuc, phong) if x],
        "freq": {},
        "rules": rules,
        "doc_type_rules": ({"Giấy chứng nhận": bo} if bo else {}),
        "default_flow_id": phong,
    }

def load_flow_store():
    """Sổ ghim/tần suất luồng trình + quy tắc tự nhận luồng theo ký hiệu văn bản — CỤC BỘ theo
    máy đang chạy (y hệt cơ chế ghim/tần suất "Nơi nhận" ở nguoi_dung.json), KHÔNG PHẢI danh
    sách luồng cứng dùng chung cho mọi người. Danh sách luồng thật luôn lấy động từ web
    (fetch_prepare_insert_data) — file này chỉ nhớ: luồng nào máy này hay dùng/đã ghim, và quy
    tắc "thấy chữ X trong ký hiệu thì chọn luồng nào" của riêng máy này."""
    try:
        s = load_json("luong_trinh.json")
    except Exception:
        s = {}
    migrated = "flows" in s and "pinned" not in s
    if migrated:
        s = _migrate_old_flow_list(s)
    s.setdefault("pinned", [])
    s.setdefault("freq", {})
    s.setdefault("rules", [])
    s.setdefault("doc_type_rules", {})
    s.setdefault("default_flow_id", None)
    if migrated:
        save_flow_store(s)   # ghi lại ngay định dạng mới — lần chạy sau khỏi phải migrate lại
    return s

def save_flow_store(store, path=None):
    """`path`: chỉ để test/gọi tay trỏ sang file khác — mặc định (None) luôn là sổ thật của máy
    đang chạy. Trước đây hàm này ghi cứng vào FLOW_STORE_FILE bất kể `store` truyền vào là gì,
    nên 1 dict tạm/giả lập (vd trong test) vẫn âm thầm ghi đè sổ thật — nay phải tự truyền
    path khác thì mới ghi chỗ khác được."""
    try:
        atomic_write_json(path or FLOW_STORE_FILE, store, ensure_ascii=False, indent=1)
    except Exception:
        pass

def bump_flow_freq(store, flow_id, path=None):
    store["freq"][flow_id] = store["freq"].get(flow_id, 0) + 1
    save_flow_store(store, path)

def flow_keyword_from_code(code):
    """Lấy phần chữ có Ý NGHĨA từ ký hiệu văn bản để HỌC quy tắc mới — bỏ số thứ tự (đổi theo
    từng văn bản, không lặp lại), chỉ giữ phần sau dấu '/' hoặc '-' CUỐI CÙNG (vd '936/CL' ->
    'CL', '2205/QĐ-BYT' -> 'BYT', '/GM-QLD' -> 'QLD'). Trả None nếu phần đó rỗng hoặc toàn số —
    không có gì lặp lại được để học thành quy tắc."""
    if not code:
        return None
    part = re.split(r"[/\-]", code)[-1].strip()
    if not part or part.isdigit():
        return None
    return part

def signer_pref_key(flow_id, node_id):
    return f"{flow_id}:{node_id}"

def remember_signer_pick(store, flow_id, node_id, user_id, path=None):
    """Nhớ người vừa được chọn cho 1 bước (flowId, nodeId) — dùng chung cho MỌI quy tắc/từ khoá
    dẫn tới cùng luồng+bước đó (không tách riêng theo từ khoá, tránh trùng lặp dữ liệu)."""
    if user_id is None:
        return
    key = signer_pref_key(flow_id, node_id)
    pref = store.setdefault("signer_pref", {}).setdefault(key, {"last": None, "freq": {}})
    pref["last"] = user_id
    uid_s = str(user_id)
    pref["freq"][uid_s] = pref["freq"].get(uid_s, 0) + 1
    save_flow_store(store, path)

def preferred_signer(store, flow_id, node_id, candidates):
    """Người nên chọn mặc định trong `candidates` (list dict có 'userId') theo lịch sử đã chọn
    cho đúng (flow_id, node_id) này — ưu tiên LẦN GẦN NHẤT, rồi tới tần suất cao nhất; None nếu
    chưa có lịch sử hoặc người cũ không còn trong danh sách ứng viên hiện tại (vd đã đổi vai trò)."""
    pref = store.get("signer_pref", {}).get(signer_pref_key(flow_id, node_id))
    if not pref:
        return None
    cand_ids = {c.get("userId") for c in candidates}
    if pref.get("last") in cand_ids:
        return pref["last"]
    for uid_s, _cnt in sorted(pref.get("freq", {}).items(), key=lambda kv: -kv[1]):
        try:
            uid = int(uid_s)
        except ValueError:
            continue
        if uid in cand_ids:
            return uid
    return None

# ---------- Tiện ích ----------
def now_ms():
    return str(int(time.time() * 1000))

def atomic_write_json(path, data, **kw):
    """Ghi JSON AN TOÀN: ghi ra file tạm CÙNG THƯ MỤC với `path` rồi mới os.replace() đè lên
    file đích — os.replace là 1 thao tác nguyên tử của hệ điều hành, nên nếu chương trình bị
    tắt/lỗi/mất điện GIỮA CHỪNG thì file đích vẫn còn nguyên bản CŨ (chưa từng bị đụng tới),
    không rơi vào tình trạng dở dang. Khác với việc ghi thẳng bằng open(path, "w") — cách đó
    XOÁ TRẮNG file đích ngay khi mở, nên nếu json.dump lỗi sau đó (đĩa đầy, encoding lỗi...)
    thì mất luôn dữ liệu cũ, dù lỗi bị 'except Exception: pass' nuốt lặng lẽ ở nơi gọi."""
    folder = os.path.dirname(path) or "."
    fd, tmp_path = tempfile.mkstemp(prefix=".tmp_", suffix=".json", dir=folder)
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            json.dump(data, f, **kw)
        os.replace(tmp_path, path)
    except Exception:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
        raise

def decode_params(pairs):
    """Giải mã tên & giá trị (khuôn lưu ở dạng URL-encoded)."""
    return [[unquote(n), unquote(v)] for n, v in pairs]

def set_param(params, name, value):
    for p in params:
        if p[0] == name:
            p[1] = value
            return
    params.append([name, value])

def set_suffix(params, suffix, value):
    """Ghi đè MỌI trường có tên kết thúc bằng suffix (kể cả có/không tiền tố)."""
    hit = False
    for p in params:
        if p[0] == suffix or p[0].endswith("." + suffix):
            p[1] = value
            hit = True
    return hit

class PipelineError(Exception):
    pass

@contextmanager
def step(log, n, total, label):
    """Bọc mỗi bước: báo bắt đầu / xong / mắc-ở-bước-nào."""
    log(f"\n▶ BƯỚC {n}/{total}: {label} …")
    t0 = time.time()
    try:
        yield
    except PipelineError as e:
        log(f"✖ MẮC Ở BƯỚC {n}/{total} — {label}\n   Lý do: {e}")
        raise
    except Exception as e:
        log(f"✖ MẮC Ở BƯỚC {n}/{total} — {label}\n   Lỗi: {e!r}")
        raise
    log(f"✓ Bước {n}/{total} xong ({time.time()-t0:.1f}s)")

def http_log(log, r):
    log(f"   ↳ HTTP {r.status_code} · {len(r.text)} ký tự · {r.url.split('?')[0].split('/')[-1]}")

# ---------- Phiên & Đăng nhập ----------
CHROME_UA = ("Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 "
             "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36")

def make_session():
    s = requests.Session()
    # Mặc định là header cho các lời gọi .do (XHR) SAU khi đã vào hệ thống.
    s.headers.update({
        "User-Agent": CHROME_UA,
        "X-Requested-With": "XMLHttpRequest",
        "Origin": BASE,
        "Referer": BASE + "/Index.do",
    })
    return s

# Header cho bước ĐĂNG NHẬP: giống hệt gõ URL trên trình duyệt (KHÔNG origin/referer/xhr)
def _nav_headers():
    return {
        "User-Agent": CHROME_UA,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "vi,en-US;q=0.9,en;q=0.8",
        "Upgrade-Insecure-Requests": "1",
        # gỡ các header của app khỏi request điều hướng CAS
        "X-Requested-With": None, "Origin": None, "Referer": None,
    }

def _scrape_hidden(html):
    """Lấy các trường ẩn của form login (đặc biệt là 'lt'), chấp nhận mọi thứ tự thuộc tính."""
    fields = {}
    for m in re.finditer(r'<input[^>]+>', html):
        tag = m.group(0)
        if 'hidden' not in tag and 'name="lt"' not in tag:
            continue
        n = re.search(r'name="([^"]+)"', tag)
        v = re.search(r'value="([^"]*)"', tag)
        if n:
            fields[n.group(1)] = v.group(1) if v else ""
    return fields

def _is_login_page(html):
    return ('id="fm1"' in html) or ('name="lt"' in html)

def _is_interstitial(html):
    t = html.lower()
    return ("chưa thể đăng nhập" in t) or ("về trang đăng nhập" in t)

def _find_login_link(html):
    """Tìm link 'Về trang đăng nhập' trong trang đệm."""
    m = re.search(r'href="([^"]*passportv3/login[^"]*)"', html)
    if m:
        return m.group(1).replace("&amp;", "&")
    return LOGIN_URL + "?appCode=VOFFICE&service=" + SERVICE

def _scrape_form_action(html):
    m = re.search(r'<form[^>]*id="fm1"[^>]*action="([^"]*)"', html) \
        or re.search(r'<form[^>]*action="([^"]*passportv3/login[^"]*)"', html)
    return m.group(1).replace("&amp;", "&") if m else None

def _pass_waf(s, r, headers, log, depth=0):
    """Vượt lớp chống bot: trang trả về JS đặt cookie rồi reload.
    Tự bóc cookie, gắn vào phiên, tải lại URL. Trả về response mới."""
    txt = r.text
    if depth > 3 or 'document.cookie' not in txt:
        return r
    got = False
    for m in re.finditer(r'document\.cookie\s*=\s*"([^"=]+)=([^"+;]+)', txt):
        name, val = m.group(1), m.group(2)
        s.cookies.set(name, val, domain="emoh.moh.gov.vn")
        log(f"• Vượt lớp bảo vệ: gắn cookie {name}…")
        got = True
    if not got:
        return r
    r2 = s.get(r.url, headers=headers, timeout=30, allow_redirects=True)
    http_log(log, r2)
    return _pass_waf(s, r2, headers, log, depth + 1)   # lặp nếu còn tầng nữa

def _get(s, url, log, **kw):
    """GET có tự vượt lớp chống bot."""
    headers = kw.pop("headers", _nav_headers())
    r = s.get(url, headers=headers, timeout=30, allow_redirects=True, **kw)
    http_log(log, r)
    return _pass_waf(s, r, headers, log)

def _get_login_page(s, log):
    """Đi từ trang chủ để CAS tự dẫn tới form login. Trả về (html, url_form)."""
    log("• Vào trang chủ để hệ thống dẫn tới đăng nhập…")
    r = _get(s, BASE + "/", log)
    html, url = r.text, r.url

    if not _is_login_page(html) and _is_interstitial(html):
        link = _find_login_link(html)
        if link.startswith("/"):
            link = BASE + link
        log("• Gặp trang đệm → theo link 'Về trang đăng nhập'…")
        r = _get(s, link, log)
        html, url = r.text, r.url

    if not _is_login_page(html):
        log("• Chưa thấy form → thử mở thẳng trang đăng nhập CAS…")
        r = _get(s, LOGIN_URL + "?appCode=VOFFICE&service=" + SERVICE, log)
        html, url = r.text, r.url

    if not _is_login_page(html):
        log("   — Nội dung nhận được (để soi):")
        log("   " + " ".join(html.split())[:300])
        raise PipelineError("Không tới được form đăng nhập (xem nội dung ở trên).")
    return html, url

def _do_post(s, html, url_form, username, password, captcha, log):
    action = _scrape_form_action(html) or url_form
    action = urljoin(url_form, action)          # ghép tương đối vào ĐÚNG URL trang login (giữ cổng 8443)
    hidden = _scrape_hidden(html)
    if "lt" not in hidden:
        log("   ⚠ Không thấy vé 'lt' trong form.")
    data = dict(hidden)
    data.update({"username": username, "password": password, "captcha": captcha or "",
                 "_eventId": "submit", "submit": "ĐĂNG NHẬP"})
    data.setdefault("loginCount", "0")
    post_headers = _nav_headers()
    post_headers["Referer"] = url_form
    post_headers["Origin"] = "https://emoh.moh.gov.vn:8443"
    log(f"• Gửi thông tin đăng nhập → {action.split('?')[0]}")
    r = s.post(action, data=data, headers=post_headers, timeout=30, allow_redirects=True)
    http_log(log, r)
    r = _pass_waf(s, r, _nav_headers(), log)     # phòng khi phản hồi cũng bị WAF chặn
    return r

def cas_login(s, username, password, log, ask_captcha):
    """Đăng nhập CAS, nạp cookie vào session s. ask_captcha(path)->str hoặc None."""
    html, url_form = _get_login_page(s, log)

    r = _do_post(s, html, url_form, username, password, "", log)
    if _is_login_page(r.text):
        # Vẫn ở trang login → sai mật khẩu, hoặc cần captcha (loginCount>=3)
        if 'id="capchaRow"' in r.text and "display:none" not in r.text.split('id="capchaRow"')[1][:40]:
            need_captcha = True
        else:
            need_captcha = ("captcha" in r.text.lower() and "display:none" not in r.text)
        if not need_captcha:
            log("   — Phản hồi (để soi):")
            log("   " + " ".join(r.text.split())[:300])
            raise PipelineError("Đăng nhập thất bại: nhiều khả năng sai tài khoản/mật khẩu.")
        # Cần captcha
        log("• Cần captcha. Đang tải ảnh…")
        ci = s.get(CAPTCHA_URL, headers=_nav_headers(), timeout=30)
        path = os.path.join(HERE, "captcha.jpg")
        with open(path, "wb") as f:
            f.write(ci.content)
        code = ask_captcha(path)
        if not code:
            raise PipelineError("Chưa nhập captcha nên dừng.")
        # lấy lại form (lt mới) rồi post kèm captcha
        html, url_form = _get_login_page(s, log)
        r = _do_post(s, html, url_form, username, password, code, log)

    if not _is_login_page(r.text) and "passportv3/login" not in r.url:
        log("   → Đăng nhập thành công.")
        return True
    log("   — Phản hồi cuối (để soi):")
    log("   " + " ".join(r.text.split())[:300])
    raise PipelineError("Đăng nhập thất bại (sai mật khẩu hoặc captcha).")

# ---------- Các bước ----------
def open_forms(s, log, report_id=None):
    """Mở phiếu trình mới + form văn bản; trả về (html_form_van_ban, html_form_phieu_trinh_sua).
    `report_id`: có giá trị khi SỬA 1 phiếu trình đã có — mở đúng form Sửa (`prepareUpdate.do`)
    thay vì Tạo mới (`prepareInsert.do`) để lấy URL upload RIÊNG cho file phiếu trình lần sửa
    này — xác nhận qua HAR thật ('thay file.har'): khi sửa phiếu và đổi file phiếu trình, trình
    duyệt gọi `voReport!prepareUpdate.do?reportId=<id>` (KHÔNG gọi `prepareInsert.do`) ngay
    trước lần upload file phiếu trình mới; phần tử thứ 2 trả về (None nếu tạo mới) — nơi gọi tự
    tách URL upload riêng cho phiếu trình từ đó thay vì từ form tạo mới."""
    report_edit_html = None
    if report_id:
        log(f"• Mở phiếu trình #{report_id} để sửa (prepareUpdate)…")
        r_edit = s.post(BASE + "/voReport!prepareUpdate.do", params={"reportId": report_id},
                         data={"dojo.preventCache": now_ms()}, timeout=30)
        report_edit_html = r_edit.text
    else:
        log("• Mở phiếu trình mới (prepareInsert)…")
        s.post(BASE + "/voReport!prepareInsert.do", data={"dojo.preventCache": now_ms()}, timeout=30)
    log("• Mở form soạn văn bản (prepareCreateDraft)…")
    r = s.post(BASE + "/voPublishDocument!prepareCreateDraft.do",
               data={"dojo.preventCache": now_ms()}, timeout=30)
    http_log(log, r)
    html = r.text
    if "login" in r.url.lower() or "passport" in r.url.lower() or len(html) < 2000:
        raise PipelineError("Phiên có vẻ đã hết hạn (bị đẩy về đăng nhập). "
                            "Đăng nhập lại trên trình duyệt, lấy cookie mới rồi thử lại.")
    return html, report_edit_html

def extract_upload_urls(html):
    """Lấy URL upload (kèm id mã hóa) cho từng ô kẹp file."""
    urls = {}
    # Ghép mỗi action với ô uploader gần nhất phía trước
    for m in re.finditer(r"getElementById\('browse(upload\w+)'\)[\s\S]{0,400}?action:\s*'([^']+)'", html):
        urls[m.group(1)] = m.group(2).replace("&amp;", "&")
    # Dự phòng: bắt trực tiếp các action upload theo thứ tự
    if not urls:
        found = [u.replace("&amp;", "&") for u in
                 re.findall(r"action:\s*'([^']*vou/file/upload[^']*)'", html)]
        if len(found) >= 1: urls["uploadReportFile"] = found[0]
        if len(found) >= 2: urls["uploadDraftFile"] = found[1]
    return urls

def upload(s, url, filepath, log):
    """Upload 1 file. Dựng multipart GIỐNG HỆT trình duyệt: MAX_FILE_SIZE rồi Filedata. Gửi
    NGUYÊN TÊN file gốc (kể cả dấu tiếng Việt/khoảng trắng) — xác nhận qua HAR thật: server
    lưu/hiển thị đúng tên có dấu bình thường, không cần đổi sang ASCII."""
    name = os.path.basename(filepath)
    log(f"• Upload '{name}'…")
    up_headers = {
        "X-Requested-With": None,
        "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
                   "image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7"),
        "Referer": BASE + "/Index.do?request_locale=en_US&mainMenu=3&trId=2.0",
        "Upgrade-Insecure-Requests": "1",
        "Sec-Fetch-Dest": "iframe",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "same-origin",
    }
    with open(filepath, "rb") as fh:
        files = [
            ("MAX_FILE_SIZE", (None, "5")),                     # trường bắt buộc, đứng TRƯỚC (như trình duyệt)
            ("Filedata", (name, fh, "application/pdf")),
        ]
        r = s.post(url, files=files, headers=up_headers, timeout=120)
    http_log(log, r)
    txt = r.text
    # Response chứa postMessage([tên, "ID", "uploadXxx", token])
    m = re.search(r'\[\s*"[^"]*"\s*,\s*"(\d{4,})"', txt)
    if not m:
        m = re.search(r'"(\d{6,})"', txt) or re.search(r'(\d{6,})', txt)
    if not m:
        raise PipelineError("Không đọc được ID file từ kết quả upload.\n--- Phản hồi ---\n" + txt[:900])
    fid = m.group(1)
    log(f"   → ID file = {fid}")
    return fid

def upload_many(s, url, paths, log):
    """Upload nhiều file vào cùng ô. Trả về (chuỗi_id nối ';', id_file_chính)."""
    ids = []
    for fp in paths:
        ids.append(upload(s, url, fp, log))
    attach = ";".join(ids) + (";" if ids else "")   # server dùng dạng 'a;b;'
    main = ids[0] if ids else ""
    return attach, main

def reload_token(s, log):
    log("• Xin token mới (reloadToken)…")
    r = s.get(BASE + "/token!reloadToken.do", params={"dojo.preventCache": now_ms()}, timeout=30)
    http_log(log, r)
    m = re.search(r"[A-Z0-9]{24,}", r.text)
    if not m:
        raise PipelineError("Không tách được token từ reloadToken.\n--- Phản hồi ---\n" + repr(r.text[:300]))
    return m.group(0)

# ---------- Luồng trình (chuỗi người ký/duyệt) ----------
def _b64(s):
    return base64.b64encode((s or "").encode("utf-8")).decode("ascii")

def fetch_flow_nodes(s, flow_id, doc_abstract, log):
    """Lấy danh sách người ký/duyệt của 1 luồng trình (searchNodeInFlow)."""
    log(f"• Lấy danh sách người trong luồng trình (flowId={flow_id})…")
    url = (f"{BASE}/voReport!searchNodeInFlow.do?flowId={flow_id}"
           f"&token.getTokenParamString()&docAbs={quote(doc_abstract or '')}")
    r = s.post(url, data={"q": "*", "start": 0, "count": 200, "startval": 0}, timeout=30)
    http_log(log, r)
    try:
        data = r.json()
    except Exception:
        raise PipelineError("Không đọc được JSON luồng trình.\n--- Phản hồi ---\n" + r.text[:400])
    items = data.get("items") or []
    if not items:
        raise PipelineError(f"Luồng trình flowId={flow_id} không trả về ai cả (kiểm tra lại flowId).")
    log(f"   → {len(items)} người trong luồng.")
    return items

def build_flow_json(items):
    """Dựng lại chuỗi myJsonString (đúng định dạng server cần: tên/vai trò/đơn vị mã hoá base64)."""
    nodes = [{
        "deptId": n.get("deptId"), "userId": n.get("userId"), "roleId": n.get("roleId"),
        "actionType": n.get("actionType"), "nodeId": n.get("nodeId"), "order": n.get("order"),
        "roleName": _b64(n.get("roleName")), "fullName": _b64(n.get("fullName")),
        "deptName": _b64(n.get("deptName")),
    } for n in items]
    return json.dumps(nodes, ensure_ascii=False, separators=(",", ":"))

# ---------- Luồng trình KHÁC (lấy động từ web, ngoài 3 luồng quen) ----------
def _parse_select_options(html_text, select_name_substr):
    """Rút các <option value=ID>Tên</option> trong 1 <select> có tên chứa `select_name_substr`
    (dùng chung cho cả 2 select tĩnh mà prepareInsert.do nhúng sẵn trong HTML: danh sách luồng
    trình VÀ danh sách hồ sơ công việc — cả 2 đều theo đúng tài khoản đang đăng nhập)."""
    m = re.search(rf'<select[^>]*{re.escape(select_name_substr)}[^>]*>(.*?)</select>', html_text, re.S)
    if not m:
        return None
    out = []
    for opt_m in re.finditer(r'<option\b[^>]*\bvalue="([^"]*)"[^>]*>(.*?)</option>', m.group(1), re.S):
        val, name = opt_m.group(1).strip(), html_unescape(opt_m.group(2)).strip()
        if not val or val == "-1":
            continue   # dòng "---Chọn---"
        out.append({"id": val, "name": name})
    return out

def fetch_prepare_insert_data(s, log):
    """Gọi voReport!prepareInsert.do 1 lần, lấy 2 danh sách nhúng sẵn trong HTML theo ĐÚNG tài
    khoản đang đăng nhập — không cần sửa code/HAR mỗi khi tài khoản/luồng/hồ sơ khác nhau:
    - Luồng trình khác 3 luồng quen (<select name="reportForm.flowAsignId">).
    - Hồ sơ công việc (<select name="reportForm.profileFlowAsignId">) — BẮT BUỘC phải đúng của
      tài khoản đang trình, vì đây chính là chỗ trước đây bị hardcode theo hồ sơ của 1 người
      (xem save_report_draft/TPL) khiến phiếu trình tạo ra bị gắn nhầm hồ sơ khi đổi tài khoản
      khác — server vẫn báo "thành công" nhưng phiếu trình lạc mất, không thấy trong thùng nháp
      của tài khoản đang dùng."""
    log("• Lấy danh sách luồng trình + hồ sơ công việc từ web (prepareInsert)…")
    r = s.post(BASE + "/voReport!prepareInsert.do", data={"dojo.preventCache": now_ms()}, timeout=30)
    flow_opts = _parse_select_options(r.text, "flowAsignId") or []
    profile_opts = _parse_select_options(r.text, "profileFlowAsignId") or []
    flows = [{"flowId": o["id"], "name": o["name"]} for o in flow_opts]
    profiles = [{"fileId": o["id"], "name": o["name"]} for o in profile_opts]
    log(f"   → {len(flows)} luồng khác + {len(profiles)} hồ sơ công việc từ web.")
    return flows, profiles

def fetch_current_user_identity(s, log):
    """Lấy (userId, fullname) của TÀI KHOẢN ĐANG ĐĂNG NHẬP qua notice!getTop.do — xác nhận qua
    2 file HAR độc lập: field "userId"/"fullname" của message đầu tiên luôn khớp đúng người
    đang xem thông báo (không phải người được nhắc tới trong nội dung thông báo, đó là
    "userActionId"). Dùng để điền reportForm.creator/creatorId đúng theo tài khoản đang chạy,
    thay vì cứng theo 1 người (xem TPL["saveReport"]) — 2 trường này khác nhau giữa các tài
    khoản (đã kiểm chứng qua HAR) và có thể là lý do phiếu trình "lưu thành công" nhưng không
    hiện trong thùng nháp của tài khoản đang dùng (voReport!onSearchMyReport.do có vẻ lọc theo
    đúng danh tính phiên đăng nhập).
    Trả về None nếu không lấy được (vd tài khoản chưa có thông báo nào) — lúc đó gọi nơi dùng
    hàm này tự rơi về giá trị tĩnh cũ, không được tự bịa danh tính."""
    r = s.post(BASE + "/notice!getTop.do", data={"dojo.preventCache": now_ms()}, timeout=30)
    try:
        data = r.json()
    except Exception:
        log("   — Không đọc được danh tính tài khoản (notice!getTop.do lỗi JSON).")
        return None
    msg = data.get("message") or {}
    user_id, full_name = msg.get("userId"), msg.get("fullname")
    if not user_id or not full_name:
        log("   — Không lấy được danh tính tài khoản từ notice!getTop.do (có thể chưa có thông báo nào).")
        return None
    return user_id, full_name

def view_node_roles(s, flow_id, node_id):
    """Giải mã 1 bước của luồng (node) chưa có sẵn người: trả về TOÀN BỘ biến thể chức danh
    khả dĩ ở bước đó — list các (roleId, roleName, deptId, deptName). deptId đã được server tự
    quy đổi từ placeholder (-1=phòng người trình, -2=đơn vị người trình) sang phòng/đơn vị thật
    của người đang đăng nhập.
    LƯU Ý: dù nhiều biến thể chung 1 roleId (vd 4052 lặp lại cho cả "Cục Trưởng"/"Ký Văn
    Bản"/"Phó Cục Trưởng"), thử nghiệm thực tế cho thấy roleName CÓ ảnh hưởng tới danh sách
    người trả về ở viewPersonAsign.do (khác quan sát ban đầu từ dữ liệu mẫu) — nên phải tra
    riêng từng biến thể, không được chỉ lấy biến thể đầu tiên."""
    r = s.post(f"{BASE}/voReport!viewNode.do?flowId={flow_id}&nodeId={node_id}",
               data={"dojo.preventCache": now_ms()}, timeout=30)
    try:
        data = r.json()
    except Exception:
        raise PipelineError(f"Không đọc được vai trò/đơn vị cho bước (nodeId={node_id}).")
    items = data.get("items") or []
    if not items:
        raise PipelineError(f"Bước (nodeId={node_id}) không xác định được vai trò/đơn vị.")
    it = items[0]
    role_ids, role_names = it.get("roleId") or [], it.get("roleName") or []
    dept_ids, dept_names = it.get("deptId") or [], it.get("deptName") or []
    if not role_ids or not dept_ids:
        raise PipelineError(f"Bước (nodeId={node_id}) thiếu vai trò hoặc đơn vị.")
    dept_id = dept_ids[0]
    dept_name = dept_names[0] if dept_names else ""
    variants, seen = [], set()
    for i, rid in enumerate(role_ids):
        rname = role_names[i] if i < len(role_names) else ""
        key = (rid, rname)
        if key in seen:
            continue
        seen.add(key)
        variants.append({"roleId": rid, "roleName": rname, "deptId": dept_id, "deptName": dept_name})
    return variants

def view_person_candidates(s, role_id, dept_id, role_name):
    """Danh sách người đủ điều kiện giữ 1 vai trò tại 1 đơn vị (có thể 1 người, có thể nhiều)."""
    url = (f"{BASE}/voReport!viewPersonAsign.do?roleId={role_id}&deptId={dept_id}"
           f"&deptName=&roleName={quote(_b64(role_name))}")
    r = s.post(url, data={"dojo.preventCache": now_ms()}, timeout=30)
    try:
        data = r.json()
    except Exception:
        raise PipelineError(f"Không đọc được danh sách người cho vai trò roleId={role_id}.")
    return data.get("items") or []

def resolve_flow_signers(s, flow_id, doc_abstract, log):
    """Với 1 luồng CHƯA có sẵn người (luồng ngoài 3 luồng quen): lấy từng bước, bước nào thiếu
    người thì tự tra (viewNode + viewPersonAsign) — tra RIÊNG từng biến thể chức danh (vd "Cục
    Trưởng" và "Phó Cục Trưởng" ở cùng 1 bước "Lãnh đạo Cục ký"), không gộp/bỏ bớt biến thể nào,
    vì mỗi biến thể có thể ra danh sách người khác nhau (đã xác nhận qua thực tế dùng: "Cục
    Trưởng" chỉ 1 người trong khi "Phó Cục Trưởng" có 3 người — bỏ sót sẽ mất lựa chọn của
    người dùng). Trả về list bước, mỗi bước thêm khoá 'candidates' (list người, mỗi người kèm
    'roleId'/'roleName' của biến thể mà họ giữ) để GUI tự điền (tổng cộng đúng 1 người) hoặc
    hiện ô chọn (từ 2 người trở lên, có ghi kèm chức danh cho dễ phân biệt)."""
    items = fetch_flow_nodes(s, flow_id, doc_abstract, log)
    out = []
    for n in items:
        node = dict(n)
        if node.get("userId") is not None:
            node["candidates"] = []
            out.append(node)
            continue
        log(f"   • Tra người cho bước: {node.get('name')!r} …")
        variants = view_node_roles(s, flow_id, node["nodeId"])
        candidates = []
        for v in variants:
            for p in view_person_candidates(s, v["roleId"], v["deptId"], v["roleName"]):
                p = dict(p)
                p["roleId"], p["roleName"] = v["roleId"], v["roleName"]
                candidates.append(p)
        if variants:
            node["deptId"], node["deptName"] = variants[0]["deptId"], variants[0]["deptName"]
        node["candidates"] = candidates
        if len(candidates) == 1:
            node["userId"] = candidates[0].get("userId")
            node["fullName"] = candidates[0].get("fullName")
            node["roleId"] = candidates[0].get("roleId")
            node["roleName"] = candidates[0].get("roleName")
        out.append(node)
    return out

ID_CACHE_FILE = os.path.join(HERE, "id_cache.json")
def load_id_cache():
    try:
        with open(ID_CACHE_FILE, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}
def save_id_cache(c):
    try:
        atomic_write_json(ID_CACHE_FILE, c, ensure_ascii=False)
    except Exception:
        pass
ID_CACHE = load_id_cache()

def _match_name(target, pairs):
    """pairs = [(id, name)]. Khớp theo tên chuẩn hóa (bỏ dấu + gọn khoảng trắng)."""
    t = " ".join(_norm(target).split())
    for nid, nm in pairs:
        if " ".join(_norm(nm).split()) == t:
            return nid
    return None

def _get_children(s, parent_id, log):
    r = s.post(BASE + "/departmentAction!getChildrenNode.do",
               data={"parentItemId": parent_id, "checboxStatus": "false",
                     "dojo.preventCache": now_ms()}, timeout=30)
    try:
        arr = r.json()
    except Exception:
        arr = json.loads(r.text)
    if isinstance(arr, dict):
        arr = arr.get("items", [])
    return [(str(n.get("id")), n.get("name", "")) for n in arr if n.get("id") is not None]

def _lienthong_roots(s, log):
    r = s.post(BASE + "/departmentAction!getTreeLinkDocument.do",
               data={"dojo.preventCache": now_ms()}, timeout=30)
    try:
        d = r.json()
    except Exception:
        d = json.loads(r.text)
    items = d.get("items", []) if isinstance(d, dict) else d
    return [(str(n.get("id")), n.get("name", "")) for n in items if n.get("id") is not None]

def resolve_node(s, node, tree, log):
    """node={'name','path'}; tree='internal'|'lien_thong'. Trả về ID (đi cây, có cache)."""
    path = node["path"]
    key = tree + "|" + " > ".join(path)
    if key in ID_CACHE:
        return ID_CACHE[key]
    if tree == "internal":
        cur = str(CAY["internal"]["root_id"])     # 52 = Bộ Y Tế
        rest = path[1:]                            # bỏ 'Bộ Y Tế'
    else:
        roots = _lienthong_roots(s, log)
        cur = _match_name(path[0], roots)
        if not cur:
            raise PipelineError(f"Không thấy gốc liên thông '{path[0]}' trên hệ thống.")
        rest = path[1:]
    for seg in rest:
        kids = _get_children(s, cur, log)
        nid = _match_name(seg, kids)
        if not nid:
            raise PipelineError(f"Kẹt ở cấp '{seg.strip()}' (đơn vị: {node['name']}). "
                                f"Tên có thể lệch so với cây thật — gửi log này để chỉnh.")
        cur = nid
    ID_CACHE[key] = cur
    save_id_cache(ID_CACHE)
    return cur

def resolve_nodes(s, nodes, tree, log):
    """Danh sách node -> (chuỗi_tên, chuỗi_id)."""
    names, ids = [], []
    for nd in nodes:
        log(f"   • Giải ID: {nd['name']}…")
        i = resolve_node(s, nd, tree, log)
        names.append(nd["name"].strip()); ids.append(i)
    return ";".join(names), ";".join(ids)

def save_document(s, cfg, doc, draft_attach, draft_sign, log, existing_pid=None):
    """Lưu 1 văn bản dự thảo (onInsertDraft). `doc` mang loại VB/số ký hiệu/trích yếu RIÊNG
    của văn bản này (mỗi văn bản trong 1 phiếu trình có thể khác loại/khác số/khác trích yếu —
    chỉ nơi nhận/độ khẩn/độ mật/luồng trình là dùng chung, lấy từ cfg). Trả về publishDocumentId.
    `existing_pid`: có giá trị khi SỬA 1 văn bản đã có (thay vì tạo mới) — server hiểu là cập
    nhật đúng văn bản đó (xác nhận qua HAR: request giống hệt tạo mới, chỉ khác đúng field
    publishDocumentId mang ID cũ thay vì để trống)."""
    params = decode_params([list(p) for p in TPL["insertDraft"]["params"]])
    P = "publishDocumentCreateForm."
    dt = ENUMS["documentType"][doc["doc_type"]]
    set_param(params, P + "documentTypeId", dt)
    set_param(params, P + "documentType", doc["doc_type"])
    set_param(params, P + "code", doc["code"])
    set_param(params, P + "documentAbstract", doc["abstract"])
    if cfg.get("priority"):
        set_param(params, P + "priorityId", ENUMS["priority"][cfg["priority"]])
        set_param(params, P + "priority", cfg["priority"])
    if cfg.get("security"):
        set_param(params, P + "securityTypeId", ENUMS["security"][cfg["security"]])
        set_param(params, P + "securityType", cfg["security"])
    set_param(params, P + "attachDraftId", draft_attach)   # dự thảo + tài liệu gửi kèm (nối ';')
    set_param(params, P + "signRequere", draft_sign)       # chỉ dự thảo cần ký
    set_param(params, P + "publishDocumentId", existing_pid or "")
    set_param(params, P + "status", "")
    set_param(params, P + "createDatePublish", datetime.now().strftime("%Y-%m-%d"))
    # Nơi nhận: xóa hết rồi đặt lại theo lựa chọn
    for suf in ["receiveInside", "receiveInsideId", "receiveEdoc", "receiveEdocId",
                "receiveToKnow", "receiveToKnowId", "receiveReport", "receiveReportId",
                "receiveOutside", "receiveOutsideId",
                "receiveSaveDepartment", "receiveSaveDepartmentId"]:
        set_suffix(params, suf, "")
    for cat, (nm, idsuf, tree) in {
        "inside":  ("receiveInside",  "receiveInsideId",  "internal"),
        "report":  ("receiveReport",  "receiveReportId",  "internal"),
        "edoc":    ("receiveEdoc",    "receiveEdocId",    "lien_thong"),
        "save":    ("receiveSaveDepartment", "receiveSaveDepartmentId", "internal"),
        "know":    ("receiveToKnow",  "receiveToKnowId",  "internal"),
    }.items():
        nodes = cfg.get("recv_" + cat) or []
        if nodes:
            log(f"• Giải ID nơi nhận ({nm})…")
            n, i = resolve_nodes(s, nodes, tree, log)
            set_suffix(params, nm, n); set_suffix(params, idsuf, i)
    set_param(params, "dojo.preventCache", now_ms())

    tok = reload_token(s, log)
    log("• Lưu văn bản dự thảo (onInsertDraft)…")
    r = s.post(BASE + "/voPublishDocument!onInsertDraft.do",
               params={"struts.token.name": "token", "token": tok},
               data=dict(params), timeout=60)
    http_log(log, r)
    m = re.search(r'"?publishDocumentId"?\s*[:=]\s*"?(\d{4,})', r.text)
    if not m:
        m = re.search(r'(\d{9,})', r.text)
    if not m:
        raise PipelineError("Không lấy được publishDocumentId sau khi lưu.\n--- Phản hồi ---\n" + r.text[:400])
    pid = m.group(1)
    log(f"   → publishDocumentId = {pid}")
    return pid

def flow_items_for_cfg(s, cfg, log):
    """Node list của luồng đang chọn (order/actionType/roleName/userId...), dùng để đánh số chữ
    ký VÀ để dựng myJsonString khi lưu. Ưu tiên `flow_nodes_override` — danh sách đã resolve sẵn
    ở khung "Chọn người ký" (GUI luôn resolve luồng đang chọn, kể cả luồng đã có sẵn người, xem
    FlowSignerPanel) — chỉ tự fetch lại nếu cfg được dựng tay, không qua khung đó."""
    items = cfg.get("flow_nodes_override")
    if items is not None:
        return items
    flow_id = cfg.get("flow_id")
    if not flow_id:
        return []
    return fetch_flow_nodes(s, flow_id, cfg.get("report_content", ""), log)

def save_report_draft(s, cfg, report_attach, report_sign, documents, log, sign="0", report_id=None):
    """Lưu phiếu trình (onUpdate). `documents`: list các văn bản đã lưu (mỗi dict có sẵn '_pid'
    = publishDocumentId từ save_document()) — ghi thành nhiều dòng draftDocumentGridForm[0],
    [1], ... (giống hệt cách trang web tự thêm dòng khi bạn bấm "Thêm văn bản" nhiều lần trong
    1 phiếu trình — xác nhận qua file HAR).
    `sign`: "0" = LƯU NHÁP (mặc định, an toàn — không đi vào luồng ký duyệt), "1" = TRÌNH THẬT
    (văn bản bắt đầu đi vào luồng ký duyệt) — xác nhận qua HAR: request giống hệt lưu nháp,
    chỉ khác đúng tham số này, không có trường nào khác cần thêm.
    `report_id`: có giá trị khi SỬA 1 phiếu trình đã có (thay vì tạo mới) — server hiểu là cập
    nhật đúng phiếu đó (xác nhận qua HAR: cùng field reportId, chỉ khác để trống hay không)."""
    q = dict(TPL["saveReport"]["query"])
    q = {unquote(k): unquote(v) for k, v in q.items()}

    flow_id = cfg.get("flow_id") or q.get("flowId")
    if flow_id:
        items = flow_items_for_cfg(s, cfg, log)
        q["flowId"] = flow_id
        q["myJsonString"] = build_flow_json(items)

    # Hồ sơ công việc (fileId/profileFlowAsignId) — KHÔNG được để nguyên giá trị tĩnh trong mẫu
    # (đó là hồ sơ của riêng 1 tài khoản cụ thể). Đổi tài khoản mà vẫn dùng hồ sơ cũ, server vẫn
    # báo "thành công" nhưng phiếu trình bị gắn nhầm hồ sơ của người khác — tìm không thấy trong
    # thùng nháp của tài khoản đang dùng. Nếu cfg có chọn hồ sơ (từ danh sách lấy động theo đúng
    # tài khoản), dùng giá trị đó; nếu không mới rơi về giá trị tĩnh trong mẫu (tài khoản gốc).
    work_profile_id = cfg.get("work_profile_id")
    work_profile_name = cfg.get("work_profile_name")
    if work_profile_id:
        q["fileId"] = work_profile_id

    q["sign"] = sign
    q["token"] = reload_token(s, log)

    params = decode_params([list(p) for p in TPL["saveReport"]["params"]])
    if work_profile_id:
        set_param(params, "reportForm.profileFlowAsignId", work_profile_id)
        set_param(params, "reportForm.profileFlowAsign", work_profile_name or "")

    # creator/creatorId — KHÔNG được để nguyên giá trị tĩnh trong mẫu (đó là danh tính riêng 1
    # tài khoản). Server có vẻ lọc "thùng nháp phiếu trình" theo đúng creatorId khớp phiên đăng
    # nhập (xem onSearchMyReport.do) — nếu để nguyên, phiếu trình vẫn "lưu thành công" nhưng bị
    # gắn nhầm danh tính người khác nên không hiện ra. Lấy động theo tài khoản đang chạy; nếu
    # không lấy được thì rơi về giá trị tĩnh trong mẫu (an toàn cho tài khoản gốc).
    # (officeId/officeName TẠM giữ nguyên = Cục Quản lý Dược — hiện chỉ dùng nội bộ Cục này;
    # đơn vị khác dùng thì tự đổi cứng 2 giá trị đó trong mẫu, chưa cần tự động ở bước này.)
    identity = fetch_current_user_identity(s, log)
    if identity:
        user_id, full_name = identity
        set_param(params, "reportForm.creatorId", user_id)
        set_param(params, "reportForm.creator", full_name)
        log(f"   → Người trình (creator): {full_name} (userId={user_id})")

    set_param(params, "reportForm.content", cfg["report_content"])
    set_param(params, "reportForm.attachId", report_attach)   # phiếu trình + tài liệu không gửi (nối ';')
    set_param(params, "reportForm.signRequere", report_sign)  # chỉ phiếu trình cần ký
    set_param(params, "reportForm.reportId", report_id or "")
    set_param(params, "reportForm.editorDate", datetime.now().strftime("%Y-%m-%d"))
    now_iso = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    for i, doc in enumerate(documents):
        pfx = f"draftDocumentGridForm[{i}]."
        set_param(params, pfx + "publishDocumentId", doc["_pid"])
        set_param(params, pfx + "code", doc["code"])
        set_param(params, pfx + "documentType", doc["doc_type"])
        set_param(params, pfx + "documentAbstract", doc["abstract"])
        set_param(params, pfx + "createDate", now_iso)
    set_param(params, "dojo.preventCache", now_ms())

    action_label = "TRÌNH" if sign == "1" else "lưu NHÁP"
    log(f"• {action_label.capitalize()} phiếu trình ({len(documents)} văn bản, onUpdate, sign={sign})…")
    r = s.post(BASE + "/voReport!onUpdate.do", params=q, data=dict(params), timeout=60)
    http_log(log, r)
    # QUAN TRỌNG: server này có thể trả HTTP 200 ngay cả khi từ chối yêu cầu (báo lỗi trong nội
    # dung, không đổi mã HTTP) — status_code=200 KHÔNG chắc chắn là đã lưu được. Log nguyên văn
    # phản hồi để biết chắc, chứ không suy đoán.
    body = r.text.strip()
    log(f"   ↳ Nội dung phản hồi ({len(body)} ký tự): {body[:500]!r}")
    if r.status_code != 200:
        raise PipelineError(f"onUpdate trả mã {r.status_code}.\n--- Phản hồi ---\n{body[:400]}")
    log(f"   → Đã gửi yêu cầu {action_label} (xem dòng 'Nội dung phản hồi' ở trên — "
        "nếu trống hoàn toàn thì có thể là bình thường, nếu có chữ thì đọc kỹ xem có báo lỗi không).")

# ---------- Xác minh SAU KHI lưu (đọc lại từ server, không suy đoán từ phản hồi onUpdate) ----------
# Đã xác nhận qua 2 file HAR thật (trình lại 1 phiếu cũ, và tạo phiếu hoàn toàn mới): phản hồi
# của chính onUpdate.do KHÔNG mang lại reportId hay bất kỳ thông tin đọc được nào (cả 2 lần đều
# rỗng dù server báo HTTP 200) — nên không thể suy ra "đã lưu" chỉ từ phản hồi đó. Thay vào đó,
# gọi lại 1-2 API TRA CỨU (đọc, không ghi gì) để tìm đúng phiếu trình vừa lưu bằng cách khớp
# content + creatorId + thời điểm tạo — nếu tìm thấy, đó là bằng chứng THẬT (server đã ghi nhận),
# không phải suy đoán từ mã HTTP.
def _search_my_report(s, grid=None, date_from=None, date_to=None, count=50):
    """Gọi onSearchMyReport.do — `grid="prepareProcessDocument"` = đúng hộp "đang trình/đang xử
    lý" (đã xác nhận qua HAR: trả về gọn, vài dòng); `grid=None` = danh sách chung KHÔNG lọc
    trạng thái (bao gồm cả nháp lẫn đã trình — đây là hộp dùng để tìm thấy cả văn bản còn ở
    thùng nháp). Tham số postData khác nhau giữa 2 biến thể là ĐÚNG theo HAR thật, không phải
    thiếu sót — mỗi biến thể có bộ tham số riêng của đúng màn hình tương ứng trên web.
    `date_from`/`date_to` (chuỗi "YYYY-MM-DD"): mặc định None → giữ đúng hành vi cũ (đầu tháng
    hiện tại đến hôm nay) — dùng cho verify_report_saved (chỉ cần tìm phiếu vừa lưu trong tháng
    này); tab "Quản lý Phiếu trình" truyền khoảng ngày người dùng tự chọn."""
    now = datetime.now()
    data = {
        "searchForm.content": "",
        "reportSearchForm.createDateFrom": date_from or now.replace(day=1).strftime("%Y-%m-%d"),
        "reportSearchForm.createDateTo": date_to or now.strftime("%Y-%m-%d"), "reportSearchForm.content": "",
        "q": "*", "start": 0, "count": count, "startval": 0,
    }
    params = None
    if grid:
        params = {"grid": grid}
        data["reportSearchForm.reportOfficeNumber"] = ""
    else:
        data["reportSearchForm.reportType"] = "-1"
        data["reportSearchForm.stateId"] = "-1"
        data["reportSearchForm.status"] = "-1"
    r = s.post(BASE + "/voReport!onSearchMyReport.do", params=params, data=data, timeout=30)
    return r.json().get("items") or []

def _search_processed_report(s, date_from=None, date_to=None, count=50):
    """Danh sách "Hoàn thành" — gọi ENDPOINT KHÁC (`voReport!onSearchReport.do?grid=processed`,
    không phải onSearchMyReport.do như 2 hộp Đang xử lý/Nháp) — xác nhận qua HAR riêng, đúng bộ
    tham số postData của màn hình "Hoàn thành" trên web (có thêm reportSearchForm.creator/
    reportNumber/finishDateFrom/To/officeId mà 2 hộp kia không có). Mỗi item trả về có thêm
    `finishDate` (ngày hoàn thành) và `status == 3`."""
    now = datetime.now()
    data = {
        "searchForm.content": "", "reportSearchForm.creator": "", "reportSearchForm.reportNumber": "",
        "reportSearchForm.createDateFrom": date_from or now.replace(day=1).strftime("%Y-%m-%d"),
        "reportSearchForm.createDateTo": date_to or now.strftime("%Y-%m-%d"),
        "reportSearchForm.finishDateFrom": "", "reportSearchForm.finishDateTo": "",
        "reportSearchForm.reportOfficeNumber": "", "reportSearchForm.officeId": "",
        "reportSearchForm.content": "", "q": "*", "start": 0, "count": count, "startval": 0,
    }
    r = s.post(BASE + "/voReport!onSearchReport.do", params={"grid": "processed"}, data=data, timeout=30)
    return r.json().get("items") or []

def _match_report(items, content, creator_id, since_dt, expect_report_id=None):
    """Khớp đúng phiếu trình VỪA lưu trong `items` (kết quả _search_my_report).
    `expect_report_id`: có giá trị khi đang SỬA 1 phiếu trình đã có (biết chắc ID trước khi lưu)
    — khớp thẳng theo `reportId`, KHÔNG dùng bộ lọc content/thời gian bên dưới, vì `createdDate`
    server trả về là ngày TẠO GỐC (không đổi khi sửa) nên lúc nào cũng sớm hơn `since_dt` (thời
    điểm sửa) — bộ lọc thời gian sẽ luôn loại bỏ nhầm chính phiếu vừa sửa nếu áp dụng ở đây.
    Không có `expect_report_id` (đang TẠO MỚI, chưa có ID nào để dựa vào) mới rơi về cách cũ:
    content nguyên văn + tạo sau `since_dt` (trừ hao 10 giây cho lệch giờ máy/server) — đã kiểm
    chứng bằng HAR thật (createdDate khớp chính xác tới từng giây với lúc gọi onUpdate). Chỉ so
    thêm creatorId khi có (`creator_id` có thể None nếu không lấy được danh tính tài khoản — vẫn
    cứ khớp theo content + thời gian, không bỏ cuộc hoàn toàn chỉ vì thiếu 1 lớp so khớp phụ)."""
    if expect_report_id is not None:
        return next((it for it in items if str(it.get("reportId")) == str(expect_report_id)), None)
    target = (content or "").strip()
    matches = []
    for it in items:
        if (it.get("content") or "").strip() != target:
            continue
        if creator_id is not None and str(it.get("creatorId")) != str(creator_id):
            continue
        try:
            created = datetime.strptime(it["createdDate"], "%Y-%m-%dT%H:%M:%S")
        except Exception:
            continue
        if created < since_dt - timedelta(seconds=10):
            continue   # phiếu trình cũ trùng nội dung, tạo từ trước lúc mình lưu — bỏ qua
        matches.append((created, it))
    if not matches:
        return None
    matches.sort(key=lambda x: x[0], reverse=True)
    return matches[0][1]

def fetch_report_process(s, report_id, log=lambda *a: None, count=200):
    """Toàn bộ các bước (đã qua/đang chờ) trong luồng ký duyệt của 1 phiếu trình — mỗi bước có
    receiveUser/receiveUserId/receiveRoleId/receiveGroupId/displayPositionName/actionType/
    processOrder/status (1=đang chờ, 0=đã qua/chưa tới). Trả [] nếu không tra được (không
    raise — chỉ là chưa hiện được tiến trình, không phải lỗi nghiêm trọng)."""
    try:
        r = s.post(BASE + "/voReport!onSearchReportProcess.do", params={"reportId": report_id},
                   data={"q": "*", "start": 0, "count": count, "startval": 0}, timeout=30)
        return r.json().get("items") or []
    except Exception as e:
        log(f"   • Không tra được tiến trình ký của phiếu trình: {e!r}")
        return []

def find_current_pending_step(s, report_id, log=lambda *a: None):
    """Bước đang chờ xử lý (status=1) — trả về dict có receiveUser/displayPositionName, hoặc
    None nếu không tra được/không có bước nào đang chờ."""
    items = fetch_report_process(s, report_id, log, count=12)
    return next((it for it in items if it.get("status") == 1), None)

def fetch_report_history(s, report_id, log=lambda *a: None, count=50):
    """Toàn bộ nhật ký hành động của 1 phiếu trình (Trình ký/Hủy trình ký/...), mới nhất trước —
    mỗi dòng có note/createAt/fullname/effectType. Trả [] nếu không tra được."""
    try:
        r = s.post(BASE + "/voReport!getReportHistory.do", params={"objectId": report_id},
                   data={"q": "*", "start": 0, "count": count, "startval": 0}, timeout=30)
        return r.json().get("items") or []
    except Exception as e:
        log(f"   • Không tra được nhật ký phiếu trình: {e!r}")
        return []

def find_latest_history_note(s, report_id, log=lambda *a: None):
    """Dòng mới nhất trong nhật ký hành động của phiếu trình — trả về dict có
    note/createAt/fullname, hoặc None nếu không tra được."""
    items = fetch_report_history(s, report_id, log, count=20)
    return items[0] if items else None

def fetch_report_attachs(s, report_id, log=lambda *a: None):
    """Danh sách file đính kèm (văn bản dự thảo) của 1 phiếu trình — mỗi item có
    draftDocumentName/draftDocumentPath/documentAbstract/documentId. Trả [] nếu không tra được."""
    try:
        r = s.post(BASE + "/voReport!getAttachs.do",
                   params={"reportId": report_id, "attachType": "draftSubmission"},
                   data={"q": "*", "start": 0, "count": 20, "startval": 0}, timeout=30)
        return r.json().get("items") or []
    except Exception as e:
        log(f"   • Không tra được file đính kèm của phiếu trình: {e!r}")
        return []

def fetch_document_of_report(s, report_id, log=lambda *a: None):
    """Chi tiết từng văn bản (Loại VB/Số ký hiệu/Trích yếu/Nơi nhận/Độ khẩn-mật/Người ký) của 1
    phiếu trình đã có — nguồn duy nhất cho các field này khi sửa nháp (không đọc được từ file
    PDF). Trả [] nếu không tra được."""
    try:
        r = s.post(BASE + "/voPublishDocument!onSearchDocumentOfReport.do",
                   params={"reportId": report_id},
                   data={"q": "*", "start": 0, "count": 200, "startval": 0}, timeout=30)
        return r.json().get("items") or []
    except Exception as e:
        log(f"   • Không tra được chi tiết văn bản của phiếu trình: {e!r}")
        return []

def fetch_edit_draft_upload_url(s, publish_document_id, log=lambda *a: None):
    """Mở form Sửa 1 văn bản đã có (onEditDraft) để lấy URL upload riêng cho lần sửa này — tái
    dùng đúng extract_upload_urls() đã có (dùng cho form tạo mới), chỉ khác nguồn HTML."""
    r = s.post(BASE + "/voPublishDocument!onEditDraft.do",
               params={"publishDocumentId": publish_document_id, "moduleCall": "reportForm"},
               data={"dojo.preventCache": now_ms()}, timeout=30)
    urls = extract_upload_urls(r.text)
    if not urls:
        raise PipelineError(f"Không thấy URL upload trong form Sửa (publishDocumentId={publish_document_id}).")
    return urls.get("uploadDraftFile") or list(urls.values())[-1]

def fetch_draft_attach_tokens(s, publish_document_id, log=lambda *a: None):
    """Token tải file cho từng file đính kèm của 1 văn bản đã có — xác nhận qua HAR thật ('har
    đính kèm để báo cáo.har'): gọi thẳng draftDocumentPath (đường dẫn thô trả về từ
    fetch_report_attachs/getAttachs.do) để tải bị chặn 403 — server chỉ cho tải qua đúng cặp
    attachId + token lấy từ endpoint này (`uploadiframe!getAttachFile.do`, objectType=2 =
    'uploadDraftFile'), rồi ghép vào `uploadiframe!openFile.do?token=...&attachId=...` (cùng
    cơ chế đã dùng để tải file phiếu trình qua attachPathIcons — CÓ hoạt động, đã tự xác nhận).
    Trả về dict {attachId (khớp draftDocumentId của getAttachs.do): {"name", "token"}}."""
    try:
        r = s.post(BASE + "/uploadiframe!getAttachFile.do",
                   params={"objectId": publish_document_id, "objectType": 2, "id": "uploadDraftFile"},
                   data={}, timeout=30)
        items = r.json().get("items") or []
        if len(items) < 3:
            return {}
        ids, names, tokens = items[0], items[1], items[2]
        return {aid: {"name": name, "token": token} for aid, name, token in zip(ids, names, tokens)}
    except Exception as e:
        log(f"   • Không tra được token tải file văn bản (publishDocumentId={publish_document_id}): {e!r}")
        return {}

def download_attach(s, url_or_path, dest_path, log=lambda *a: None):
    """Tải 1 file đính kèm đã có trên hệ thống về máy (dùng session đã đăng nhập) — cho tính
    năng Sửa nháp tự điền lại file cũ. `url_or_path` là URL đầy đủ (từ attachPathIcons hoặc từ
    fetch_draft_attach_tokens() + uploadiframe!openFile.do — xem 2 nơi gọi)."""
    url = url_or_path if url_or_path.startswith("http") else BASE + url_or_path
    r = s.get(url, timeout=60)
    if r.status_code != 200 or not r.content:
        raise PipelineError(f"Tải file thất bại (status={r.status_code}): {url}")
    with open(dest_path, "wb") as f:
        f.write(r.content)
    log(f"   → Đã tải file cũ về: {dest_path}")
    return dest_path

def open_file_with_default_app(path):
    """Mở 1 file bằng ứng dụng mặc định của hệ điều hành (PDF reader, Word...) — dùng cho nút
    "Mở" ở khung Chi tiết phiếu trình (xem ReportDetailWindow._open_attach). Ném lỗi ra ngoài
    nếu không mở được, để nơi gọi tự quyết định báo cho người dùng ra sao."""
    if sys.platform == "win32":
        os.startfile(path)   # type: ignore[attr-defined]
    elif sys.platform == "darwin":
        subprocess.Popen(["open", path])
    else:
        subprocess.Popen(["xdg-open", path])

def remove_attach_file(s, attach_id, log=lambda *a: None):
    """Xoá 1 file đính kèm CŨ khỏi hệ thống (`uploadiframe!removeFile.do`) — BẮT BUỘC khi Sửa 1
    văn bản/phiếu trình đã có file và upload lại: xác nhận qua 2 HAR thật ('thay file.har',
    'bấm trình đi.har') — trình duyệt luôn gọi removeFile cho ID cũ ngay trước khi gắn file mới.
    Nếu bỏ qua bước này: `attachDraftId` là danh sách nối ';' (xem save_document) — file mới
    upload lại chỉ CỘNG THÊM vào chứ không thay thế, càng Sửa nhiều lần văn bản càng tích file
    trùng lặp. CHỈ gọi ở nhánh ghi thật (sau khi check_only đã return) — đây là hành động XOÁ
    THẬT trên hệ thống, không phải bước đọc. Best-effort: lỗi thì log rõ để tự kiểm tra lại trên
    web, không raise (1 file xoá lỗi không nên chặn toàn bộ phiếu)."""
    try:
        s.post(BASE + "/uploadiframe!removeFile.do", params={"attachId": attach_id},
               data={"dojo.preventCache": now_ms()}, timeout=30)
        log(f"   → Đã xoá file cũ (attachId={attach_id}) trước khi thay bằng file mới.")
    except Exception as e:
        log(f"   • Không xoá được file cũ (attachId={attach_id}): {e!r} — có thể còn sót file "
            "trùng lặp, tự kiểm tra lại trên web.")

def cancel_report(s, report_id, log=lambda *a: None):
    """Thu hồi (hủy trình ký) 1 phiếu trình đang xử lý. Không có mẫu HAR nào đọc được response
    của chính onCancelReport.do (Chrome không giữ cache kịp lúc chụp) — áp dụng đúng triết lý
    'không tin response, verify lại bằng truy vấn đọc' đã dùng cho việc lưu/trình (xem
    verify_report_saved): sau khi gọi, tra lại getReportHistory và tìm dòng mới nhất có
    effectType==7 ("Hủy trình ký phiếu trình" — đã xác nhận đúng chữ này qua HAR thật).
    Trả (ok: bool, note: dict|None) — ok=False không có nghĩa chắc chắn thất bại, chỉ là chưa
    xác minh được (giống verify_report_saved), KHÔNG raise trừ khi bản thân request lỗi mạng."""
    tok = reload_token(s, log)
    log(f"• Thu hồi phiếu trình (reportId={report_id})…")
    r = s.post(BASE + "/voReport!onCancelReport.do",
               params={"reportId": report_id, "struts.token.name": "token", "token": tok},
               data={"dojo.preventCache": now_ms()}, timeout=30)
    http_log(log, r)
    if r.status_code != 200:
        raise PipelineError(f"onCancelReport trả mã {r.status_code}.\n--- Phản hồi ---\n{r.text[:400]}")
    for delay in VERIFY_RETRY_DELAYS:
        log(f"   • Chờ {delay}s rồi tra lại lịch sử để xác minh…")
        time.sleep(delay)
        items = fetch_report_history(s, report_id, log, count=5)
        if items and items[0].get("effectType") == 7:
            log("   ✓ Lịch sử đã có dòng 'Hủy trình ký phiếu trình' — thu hồi thành công.")
            return True, items[0]
    log("   ⚠ Đã thử đủ số lần vẫn chưa thấy dòng xác nhận trong lịch sử — không coi là lỗi, "
        "chỉ là chưa xác minh được (tự kiểm tra lại trên web nếu cần chắc chắn).")
    return False, None

VERIFY_RETRY_DELAYS = [3, 4, 5, 6]   # giây nghỉ TRƯỚC mỗi lần thử — hệ thống này khá chậm, nên
                                     # thử vài lần cách quãng thay vì tin ngay 1 lần đầu; DỪNG
                                     # NGAY khi có kết quả, không chờ hết các lần còn lại (tổng
                                     # cộng tối đa ~18 giây nếu chưa lần nào thấy).

def verify_report_saved(s, cfg, sign, since_dt, log=lambda *a: None):
    """Xác minh (đọc lại từ server) rằng phiếu trình vừa lưu/trình thật sự tồn tại — KHÔNG bao
    giờ raise (best-effort): thất bại ở đây không có nghĩa là việc lưu thất bại, chỉ là chưa
    xác minh được, để nơi gọi hiển thị đúng trạng thái 'chưa chắc' thay vì báo lỗi giả.
    Mỗi lần thử tra ĐỒNG THỜI cả 2 nơi — hộp "đang trình" (grid=prepareProcessDocument) VÀ danh
    sách chung (bao gồm cả thùng nháp) — thấy ở đâu trước thì biết trạng thái luôn: thấy ở hộp
    "đang trình" là bằng chứng mạnh nhất (đã thật sự vào luồng ký duyệt); chỉ thấy ở danh sách
    chung (không thấy ở "đang trình") nghĩa là đã lưu nhưng có thể CHƯA thực sự vào luồng — đáng
    chú ý riêng nếu đang Trình (sign=1), vì lẽ ra phải thấy ở cả 2 nơi.
    Trả về dict: {"verified", "report_id", "in_process", "pending", "history_note"}."""
    result = {"verified": False, "report_id": None, "in_process": False,
              "pending": None, "history_note": None}
    identity = None
    try:
        identity = fetch_current_user_identity(s, log)
    except Exception as e:
        log(f"   • Không lấy được danh tính tài khoản: {e!r} — vẫn thử xác minh, chỉ bớt 1 lớp so khớp.")
    if not identity:
        log("   • Không xác định được tài khoản đang đăng nhập — vẫn thử khớp theo nội dung + thời gian.")
    creator_id = identity[0] if identity else None
    content = cfg.get("report_content", "")
    # Đang SỬA phiếu cũ (cfg["report_id"] đã biết trước khi lưu) — khớp thẳng theo reportId thay
    # vì content+since_dt (xem _match_report): createdDate không đổi khi sửa nên luôn "cũ hơn"
    # since_dt, bộ lọc thời gian sẽ loại bỏ nhầm chính phiếu vừa sửa nếu vẫn dùng cách cũ.
    expect_report_id = cfg.get("report_id")

    item, in_process = None, False
    for attempt, delay in enumerate(VERIFY_RETRY_DELAYS, start=1):
        log(f"   • Chờ {delay}s rồi tra lại lần {attempt}/{len(VERIFY_RETRY_DELAYS)}…")
        time.sleep(delay)
        try:
            item_process = _match_report(_search_my_report(s, grid="prepareProcessDocument"),
                                          content, creator_id, since_dt, expect_report_id)
        except Exception as e:
            log(f"   • Tra hộp 'đang trình' lỗi: {e!r}")
            item_process = None
        try:
            item_all = _match_report(_search_my_report(s, grid=None), content, creator_id, since_dt,
                                      expect_report_id)
        except Exception as e:
            log(f"   • Tra danh sách phiếu trình (kể cả nháp) lỗi: {e!r}")
            item_all = None
        if item_process or item_all:
            item, in_process = (item_process or item_all), bool(item_process)
            where = "hộp 'đang trình'" if in_process else "danh sách phiếu trình (chưa thấy ở 'đang trình')"
            log(f"   ✓ Tìm thấy ở {where} — dừng thử lại.")
            break
        log("   • Chưa thấy ở đâu cả, thử lại.")
    else:
        log("   ⚠ Đã thử đủ số lần vẫn chưa thấy — không coi là lỗi lưu, chỉ là chưa xác minh được.")

    if not item:
        return result
    result["verified"] = True
    result["report_id"] = item.get("reportId")
    result["in_process"] = in_process
    if sign == "1":
        try:
            result["pending"] = find_current_pending_step(s, result["report_id"], log)
        except Exception as e:
            log(f"   • Không tra được người đang giữ: {e!r}")
        try:
            result["history_note"] = find_latest_history_note(s, result["report_id"], log)
        except Exception as e:
            log(f"   • Không tra được nhật ký: {e!r}")
    return result

# ---------- Luồng chính ----------
# Các "phase" (mốc lớn) để GUI hiện checklist tiến trình thân thiện, gộp nhiều bước kỹ thuật
# nhỏ (xem step() ở dưới) lại thành các mốc dễ hiểu cho người dùng phổ thông.
PIPELINE_PHASES = [
    ("prepare", "Chuẩn bị file"),
    ("upload", "Tải văn bản lên hệ thống"),
    ("save_docs", "Lưu văn bản"),
    ("save_report", "Lưu phiếu trình"),
    ("verify", "Xác minh với hệ thống"),
]

def run_pipeline(s, cfg, log, check_only=False, phase_cb=lambda key: None):
    # Mỗi văn bản dự thảo trong cfg["documents"] có file + loại VB/số/trích yếu RIÊNG (xác nhận
    # qua HAR "luồng trình 2 văn bản một lúc"): mỗi văn bản upload riêng, onInsertDraft riêng
    # (ra publishDocumentId riêng), rồi phiếu trình liệt kê tất cả qua draftDocumentGridForm[i].
    all_documents = [dict(d) for d in (cfg.get("documents") or [])]
    documents = [d for d in all_documents if d.get("file_draft_main")]
    report_main = cfg.get("file_report_main") or None
    report_files = ([report_main] if report_main else []) + list(cfg.get("files_report_extra") or [])
    borrowed_report_as_draft = False
    if not documents:            # thử nghiệm: chưa chọn văn bản nào thì dùng tạm file phiếu trình
        documents = [{"doc_type": "", "code": "", "abstract": "",
                      "file_draft_main": None, "files_draft_extra": []}]
        borrowed_report_as_draft = True

    n_docs = len(documents)
    total = (5 + n_docs) if check_only else (6 + 2 * n_docs)   # +1 so với trước = bước "Xác minh"
    log(f"Chế độ: {'CHỈ KIỂM TRA (không ghi)' if check_only else 'LƯU NHÁP thật'} · {total} bước"
        + (f" · {n_docs} văn bản" if not borrowed_report_as_draft else ""))
    n = [0]
    def nn():
        n[0] += 1
        return n[0]

    phase_cb("prepare")
    report_id = cfg.get("report_id")
    with step(log, nn(), total, "Mở 2 form (phiếu trình + văn bản)"):
        html, report_edit_html = open_forms(s, log, report_id=report_id)

    with step(log, nn(), total, "Tìm ô kẹp file trong form"):
        urls = extract_upload_urls(html)
        log(f"   Tìm thấy {len(urls)} ô: {', '.join(urls) or '(KHÔNG có!)'}")
        if not urls:
            raise PipelineError("Không thấy URL upload trong form.")
        url_report = urls.get("uploadReportFile") or list(urls.values())[0]
        url_draft  = urls.get("uploadDraftFile")  or list(urls.values())[-1]
        if report_id:
            # Sửa phiếu trình đã có — URL upload của file phiếu trình phải lấy từ form Sửa
            # (report_edit_html), không phải từ form tạo mới ở trên (xem open_forms()).
            edit_urls = extract_upload_urls(report_edit_html or "")
            if not edit_urls:
                raise PipelineError(
                    f"Không thấy URL upload trong form Sửa phiếu trình (reportId={report_id}).")
            url_report = edit_urls.get("uploadReportFile") or list(edit_urls.values())[0]

    for doc in documents:
        main = doc.get("file_draft_main")
        doc["_files"] = (report_files[:1] if borrowed_report_as_draft else
                          ([main] if main else []) + list(doc.get("files_draft_extra") or []))

    with step(log, nn(), total, "Đánh số chữ ký lên Phiếu trình + các Văn bản (theo luồng trình)"):
        if not cfg.get("auto_stamp", True):
            log("   (Bỏ qua — đã tắt tự đánh số.)")
        else:
            flow_items = flow_items_for_cfg(s, cfg, log)
            if borrowed_report_as_draft:
                if report_main and report_main.lower().endswith(".pdf"):
                    log("  › File phiếu trình:")
                    stamped = stamp_signature_numbers(report_main, flow_items, log,
                                                       stamps=cfg.get("stamps_report_override"))
                    report_files = [stamped] + list(cfg.get("files_report_extra") or [])
                log("   (Nhóm VĂN BẢN đang dùng tạm file phiếu trình để thử — không đánh số riêng.)")
            else:
                if report_main and report_main.lower().endswith(".pdf"):
                    log("  › File phiếu trình:")
                    stamped = stamp_signature_numbers(report_main, flow_items, log,
                                                       stamps=cfg.get("stamps_report_override"))
                    report_files = [stamped] + list(cfg.get("files_report_extra") or [])
                for i, doc in enumerate(documents):
                    main = doc.get("file_draft_main")
                    if main and main.lower().endswith(".pdf"):
                        log(f"  › Văn bản {i+1}/{n_docs}: {os.path.basename(main)}")
                        stamped = stamp_signature_numbers(main, flow_items, log,
                                                           stamps=doc.get("stamps_override"))
                        doc["_files"] = [stamped] + list(doc.get("files_draft_extra") or [])

    phase_cb("upload")
    with step(log, nn(), total, f"Upload nhóm PHIẾU TRÌNH ({len(report_files)} file) → lấy ID"):
        report_attach, report_sign = upload_many(s, url_report, report_files, log)
        log(f"   attachId={report_attach}  (ký: {report_sign})")

    for i, doc in enumerate(documents):
        label = f"Upload Văn bản {i+1}/{n_docs} ({len(doc['_files'])} file) → lấy ID"
        with step(log, nn(), total, label):
            # Sửa 1 văn bản đã có (existing_pid) — dùng URL upload RIÊNG của form Sửa
            # (onEditDraft), không phải URL upload của form tạo mới (url_draft) — xác nhận
            # qua HAR: web mở lại đúng dialog Sửa để lấy URL này trước khi upload file mới.
            existing_pid = doc.get("_existing_pid")
            doc_upload_url = fetch_edit_draft_upload_url(s, existing_pid, log) if existing_pid else url_draft
            attach, sign = upload_many(s, doc_upload_url, doc["_files"], log)
            doc["_attach"], doc["_sign"] = attach, sign
            log(f"   attachDraftId={attach}  (ký: {sign})")

    if check_only:
        with step(log, nn(), total, "Xin token (kiểm tra) — KHÔNG ghi gì"):
            reload_token(s, log)
        log("\n✔ KIỂM TRA OK: phiên, upload, token đều hoạt động. Chưa tạo gì cả.")
        return

    phase_cb("save_docs")
    for i, doc in enumerate(documents):
        with step(log, nn(), total, f"Xin token + LƯU văn bản {i+1}/{n_docs} (onInsertDraft)"):
            # Sửa văn bản đã có (existing_pid) — xoá file CŨ trước khi lưu, nếu không file mới
            # (vừa upload lại ở bước trên) sẽ CỘNG THÊM vào attachDraftId chứ không thay thế,
            # càng Sửa/Trình lại nhiều lần văn bản càng tích file trùng lặp (xác nhận qua HAR
            # thật — xem remove_attach_file). CHỈ chạy ở đây (sau khi check_only đã return phía
            # trên) — đây là bước ghi/xoá thật.
            for old_id in doc.get("_existing_attach_ids") or []:
                remove_attach_file(s, old_id, log)
            doc["_pid"] = save_document(s, cfg, doc, doc["_attach"], doc["_sign"], log,
                                         existing_pid=doc.get("_existing_pid"))

    submit_sign = cfg.get("submit_sign", "0")
    step_label = "Xin token + TRÌNH phiếu trình (onUpdate sign=1)" if submit_sign == "1" \
        else "Xin token + LƯU NHÁP phiếu trình (onUpdate sign=0)"
    phase_cb("save_report")
    since_dt = datetime.now()   # mốc thời gian TRƯỚC lúc gọi onUpdate — dùng để lọc khi dò lại
                                # reportId ở bước xác minh (loại bỏ phiếu trình cũ trùng nội dung)
    with step(log, nn(), total, step_label):
        # Sửa phiếu trình đã có (report_id) — xoá file CŨ của chính phiếu trình trước khi lưu,
        # cùng lý do với file văn bản ở trên (reportForm.attachId cũng là danh sách nối ';').
        if cfg.get("report_id"):
            for old_id in cfg.get("report_existing_attach_ids") or []:
                remove_attach_file(s, old_id, log)
        save_report_draft(s, cfg, report_attach, report_sign, documents, log, sign=submit_sign,
                           report_id=cfg.get("report_id"))

    phase_cb("verify")
    with step(log, nn(), total, "Xác minh lại với hệ thống (đọc lại, không ghi gì)"):
        verify_result = verify_report_saved(s, cfg, submit_sign, since_dt, log)
        if verify_result["verified"]:
            log(f"   ✓ Đã xác minh: reportId={verify_result['report_id']} có tồn tại trên hệ thống.")
        else:
            log("   ⚠ Chưa xác minh được — không có nghĩa là chưa lưu, chỉ là chưa đọc lại "
                "được ngay, tự kiểm tra thêm trên web nếu cần.")

    if submit_sign == "1":
        log(f"\n✔ XONG TẤT CẢ {total} BƯỚC. Phiếu trình đã được TRÌNH.")
    else:
        log(f"\n✔ XONG TẤT CẢ {total} BƯỚC. Phiếu trình đã lưu NHÁP.")
        log("→ Mở web → thùng nháp phiếu trình → kiểm tra → tự bấm TRÌNH.")
    return {"submit_sign": submit_sign, **verify_result}

# ---------- Giao diện ----------
# ================= LƯU CÀI ĐẶT & DỮ LIỆU THÔNG MINH =================
SETTINGS_FILE = os.path.join(HERE, "settings.json")
STORE_FILE = os.path.join(HERE, "nguoi_dung.json")
KR_SERVICE = "emoh_trinh_vb"
try:
    import keyring
except Exception:
    keyring = None

def load_settings():
    try:
        with open(SETTINGS_FILE, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def save_settings(d):
    try:
        atomic_write_json(SETTINGS_FILE, d, ensure_ascii=False, indent=1)
    except Exception:
        pass

def save_password(username, pw):
    if keyring:
        try:
            keyring.set_password(KR_SERVICE, username, pw); return True
        except Exception:
            return False
    return False

def load_password(username):
    if keyring and username:
        try:
            return keyring.get_password(KR_SERVICE, username)
        except Exception:
            return None
    return None

def load_store():
    try:
        with open(STORE_FILE, encoding="utf-8") as f:
            s = json.load(f)
    except Exception:
        s = {}
    s.setdefault("freq", {}); s.setdefault("pinned", []); s.setdefault("templates", {})
    return s

def save_store(store, path=None):
    """`path`: chỉ để test/gọi tay trỏ sang file khác — mặc định (None) luôn là sổ thật."""
    try:
        atomic_write_json(path or STORE_FILE, store, ensure_ascii=False, indent=1)
    except Exception:
        pass

def bump_freq(store, name, path=None):
    store["freq"][name] = store["freq"].get(name, 0) + 1
    save_store(store, path)

def toggle_pin(store, name, path=None):
    p = store["pinned"]
    p.remove(name) if name in p else p.append(name)
    save_store(store, path)

# ---------- Tìm kiếm không dấu + viết tắt + mờ ----------
# @lru_cache: tên đơn vị/vai trò lặp lại hàng nghìn lần trong cây (mỗi phím gõ quét lại
# toàn bộ pool) nhưng bản thân chuỗi tên không đổi giữa các lần gõ -> nhớ kết quả để
# khỏi normalize/regex lại từ đầu mỗi phím (đỡ giật khi gõ tiếng Việt qua bộ gõ Telex,
# vì 1 ký tự có dấu thường phát sinh vài sự kiện KeyRelease liên tiếp).
@lru_cache(maxsize=None)
def _norm(s):
    s = unicodedata.normalize("NFD", s.lower())
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return s.replace("đ", "d")

@lru_cache(maxsize=None)
def _initials(name):
    return "".join(w[0] for w in re.split(r"[\s\-_/.]+", _norm(name)) if w)

def _is_subseq(q, t):
    it = iter(t)
    return all(c in it for c in q)

def _score(q, name):
    if not q:
        return 0
    n = _norm(name); ini = _initials(name)
    if n.startswith(q): return 100
    if ini.startswith(q): return 95
    if q in n: return 80
    if _is_subseq(q, ini): return 70
    if _is_subseq(q, n): return 40
    return -1

def _node_key(node):
    return " > ".join(node["path"])

def search_nodes(query, nodes, store, k=15):
    """Tìm trên danh sách node {name,parent,path}. Trả về list node."""
    q = _norm(query.strip())
    freq = store.get("freq", {}); pinned = set(store.get("pinned", []))
    if not q:
        cand = [nd for nd in nodes if _node_key(nd) in pinned or freq.get(_node_key(nd), 0) > 0]
        cand.sort(key=lambda nd: (-(1000 if _node_key(nd) in pinned else 0) - freq.get(_node_key(nd), 0), nd["name"]))
        return cand[:k]
    scored = []
    for nd in nodes:
        sc = _score(q, nd["name"])
        if sc < 0:
            continue
        kkey = _node_key(nd)
        boost = (1000 if kkey in pinned else 0) + min(freq.get(kkey, 0), 100)
        scored.append((sc, boost, nd))
    scored.sort(key=lambda x: (-x[0], -x[1], len(x[2]["name"])))
    return [nd for _, _, nd in scored[:k]]

# ---------- Đọc PDF Dự thảo văn bản: tự điền Loại VB + Số/ký hiệu + Trích yếu ----------
# Mẫu Công văn: "Số: <số>/<ký hiệu>" ở đầu văn bản; <số> có thể trống (chưa cấp số).
# Một số mẫu có tên loại (Giấy chứng nhận...) lại ghi "Số hiệu:" thay vì "Số:".
CODE_RE = re.compile(r'S[ốôo]\s*(?:hi[eệ]u)?\s*:?\s*(\d*)\s*/\s*([A-Za-zÀ-ỹĐđ0-9\-]+)')
# Trích yếu (mẫu Công văn): đoạn ngay sau "V/v" (hoặc "Về việc") tới khi gặp dòng trống
# hoặc khối tiêu đề bên cột phải (Kính gửi / Cộng hòa.../ Độc lập.../ Hà Nội, ngày...).
ABSTRACT_RE = re.compile(r'(?:V\s*/\s*v|Về\s+việc)\s*:?\s*(.*)', re.IGNORECASE)
ABSTRACT_STOP = ("kinh gui", "cong hoa", "doc lap", "ha noi, ngay", "ha noi,ngay")

# Trích yếu (mẫu có tên loại — Quyết định/Kế hoạch/Giấy mời...): không có "V/v", trích yếu
# nằm ngay dưới dòng tên loại (VD "QUYẾT ĐỊNH") tới khi gặp dòng mở đầu phần nội dung/căn cứ.
TITLE_ABSTRACT_STOP = ("can cu", "thuc hien", "nham", "xet de nghi", "theo de nghi", "xet ")

def extract_code_from_text(text):
    m = CODE_RE.search(text)
    if not m:
        return None
    num, code = m.group(1), m.group(2)
    return (num + "/" + code) if num else ("/" + code)

def extract_abstract_from_text(text):
    m = ABSTRACT_RE.search(text)
    if not m:
        return None
    tail = text[m.start(1):]
    out = []
    for ln in tail.splitlines()[:6]:      # tối đa 6 dòng cho an toàn
        s = ln.strip()
        if not s or any(_norm(s).startswith(mk) for mk in ABSTRACT_STOP):
            break
        out.append(s)
    joined = " ".join(" ".join(out).split())
    return joined or None

def _sentence_case(s):
    """VIẾT HOA TOÀN BỘ -> Chỉ hoa chữ đầu (dùng cho khối tiêu đề 2 dòng kiểu Giấy chứng nhận)."""
    s = " ".join(s.split())
    return s[:1].upper() + s[1:].lower() if s else s

_DOC_TYPE_NAMES_BY_LEN = sorted(DOC_TYPE_TITLES.keys(), key=len, reverse=True)

def extract_doc_type_title(lines):
    """Dò 15 dòng đầu trang 1 tìm dòng viết HOA khớp 1 tên loại trong ENUMS[documentType]
    (VD "QUYẾT ĐỊNH" -> "Quyết định"). Trả về (tên loại, chỉ số dòng) hoặc (None, None) nếu
    không thấy — mẫu Công văn vốn không có dòng tên loại nên sẽ luôn rơi vào trường hợp này.

    2 bước, ưu tiên độ chắc chắn: (1) quét hết tìm dòng khớp TUYỆT ĐỐI cả dòng trước (đáng tin
    nhất) — có thì dùng luôn; (2) chỉ khi không dòng nào khớp tuyệt đối mới quét lại tìm dòng
    BẮT ĐẦU BẰNG 1 tên loại rồi còn chữ khác phía sau (VD "GIẤY CHỨNG NHẬN THỰC HÀNH TỐT" —
    tên loại "Giấy chứng nhận" + phụ đề dính cùng dòng). Nếu 1 dòng bắt đầu bằng nhiều tên loại
    (VD tên loại này là tiền tố của tên loại khác), ưu tiên tên DÀI hơn (cụ thể hơn)."""
    upper_lines = []
    for i, raw in enumerate(lines[:15]):
        s = " ".join(raw.split())
        if s and _is_upper_title(s):
            upper_lines.append((i, s))
    for i, s in upper_lines:
        dt = DOC_TYPE_TITLES.get(s)
        if dt:
            return dt, i
    for i, s in upper_lines:
        for name in _DOC_TYPE_NAMES_BY_LEN:
            if s.startswith(name + " "):
                return DOC_TYPE_TITLES[name], i
    return None, None

def extract_abstract_after_title(lines, title_idx):
    """Trích yếu cho văn bản có tên loại. Nếu dòng ngay sau tên loại cũng viết HOA (khối
    tiêu đề 2 dòng, VD "GIẤY CHỨNG NHẬN" + "ĐỦ ĐIỀU KIỆN KINH DOANH DƯỢC") thì ghép 2 dòng
    và chỉ hoa chữ đầu. Ngược lại lấy các dòng thường ngay dưới, dừng khi gặp dòng mở đầu
    phần nội dung (Căn cứ/Thực hiện/Nhằm...) hoặc 1 dòng viết HOA khác (dòng chức danh, VD
    "BỘ TRƯỞNG BỘ Y TẾ")."""
    nxt = lines[title_idx + 1].strip() if title_idx + 1 < len(lines) else ""
    if nxt and _is_upper_title(nxt):
        return _sentence_case(lines[title_idx] + " " + nxt)
    out = []
    for ln in lines[title_idx + 1: title_idx + 1 + 6]:
        s = ln.strip()
        if not s or _is_upper_title(s):
            break
        if any(_norm(s).startswith(mk) for mk in TITLE_ABSTRACT_STOP):
            break
        out.append(s)
    joined = " ".join(" ".join(out).split())
    return joined or None

def extract_draft_fields(path):
    """Đọc trang 1 file PDF dự thảo. Trả về (loại VB, code, abstract) — phần không suy ra
    được = None. loại VB chỉ có giá trị khi nhận ra được dòng tên loại (Quyết định, Kế
    hoạch...); mẫu Công văn không có dòng này nên loại VB luôn là None (giữ nguyên combobox)."""
    if PdfReader is None:
        raise RuntimeError("Chưa cài thư viện 'pypdf'. Chạy: pip install pypdf")
    reader = PdfReader(path)
    if not reader.pages:
        return None, None, None
    text = reader.pages[0].extract_text() or ""
    text = unicodedata.normalize("NFC", text)
    code = extract_code_from_text(text)

    doc_type, title_idx = extract_doc_type_title(text.splitlines())
    if doc_type is not None:
        abstract = extract_abstract_after_title(text.splitlines(), title_idx)
        return doc_type, code, abstract
    return None, code, extract_abstract_from_text(text)

# ---------- Thư mục tạm cho file PHÁI SINH (chuyển .docx, đánh số chữ ký) ----------
# Không bao giờ ghi đè lên đúng đường dẫn file gốc (dễ lỗi nếu file gốc đang mở ở chương trình
# khác) và không đổi tên (thêm ngày/hậu tố) để tránh rác/lẫn lộn trong thư mục của người dùng —
# thay vào đó luôn tạo 1 thư mục tạm RIÊNG (do chương trình toàn quyền sở hữu, không ai khác
# đụng tới) cho mỗi file phái sinh, giữ NGUYÊN TÊN gốc bên trong đó. Dọn lúc thoát chương trình
# (atexit) — các file này chỉ dùng 1 lần cho đúng lượt upload hiện tại, không cần giữ lại.
_GEN_TMPDIRS = []

def _cleanup_gen_tmpdirs():
    for d in _GEN_TMPDIRS:
        shutil.rmtree(d, ignore_errors=True)
    _GEN_TMPDIRS.clear()

atexit.register(_cleanup_gen_tmpdirs)

def _new_gen_tmpdir(prefix):
    d = tempfile.mkdtemp(prefix=prefix)
    _GEN_TMPDIRS.append(d)
    return d

def convert_office_doc_to_pdf(path):
    """Chuyển .docx sang .pdf bằng gói docx2pdf (điều khiển Word cài sẵn trên máy). KHÔNG hỗ
    trợ .doc (định dạng cũ) — gói docx2pdf tự chặn cứng chỉ nhận .docx. Lưu PDF vào 1 thư mục
    tạm riêng (xem _new_gen_tmpdir ở trên) — KHÔNG lưu cạnh file gốc nữa (tránh ghi đè nếu 1
    bản PDF cùng tên đã tồn tại và đang mở ở chương trình khác). Tên file giữ nguyên, chỉ đổi
    đuôi .docx -> .pdf. Trả về đường dẫn PDF vừa tạo."""
    if docx2pdf is None:
        raise RuntimeError("Chưa cài thư viện 'docx2pdf'. Chạy: pip install docx2pdf")
    if not path.lower().endswith(".docx"):
        raise RuntimeError("Chỉ hỗ trợ .docx — file .doc (định dạng cũ) hãy tự mở bằng Word, "
                            "'Save As' sang .docx hoặc .pdf rồi chọn lại.")
    src = os.path.abspath(path)
    out_dir = _new_gen_tmpdir("voffice_conv_")
    pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
    docx2pdf.convert(src, pdf_path)
    if not os.path.exists(pdf_path):
        raise RuntimeError("Word không tạo ra file PDF (không rõ lý do) — tự xuất PDF tay rồi chọn lại.")
    return pdf_path

def extract_phieu_trinh_content(path):
    """Đọc trang 1 file PHIẾU TRÌNH (không phải văn bản dự thảo): tìm dòng "PHẦN I..." làm
    mốc, lấy nội dung ngay sau CỤM IN ĐẬM ĐẦU TIÊN xuất hiện sau mốc đó (nhãn trường — không
    cần biết trước tên nhãn cụ thể như "Tên văn bản trình:"/"Nội dung xin ý kiến:"/..., tự
    nhận theo định dạng đậm/thường chung của mẫu phiếu trình) cho tới cụm in đậm tiếp theo
    (nhãn trường kế tiếp), gộp tất cả thành 1 dòng. Trả None nếu không tìm được mốc/nhãn (mẫu
    phiếu trình khác, hoặc PDF quét ảnh không có lớp chữ) — không suy đoán, để điền tay.
    Cần pymupdf (fitz) vì cần biết chữ nào IN ĐẬM — pypdf (dùng cho extract_draft_fields) chỉ
    đọc được chữ thô, không giữ định dạng."""
    if fitz is None:
        raise RuntimeError("Chưa cài thư viện 'pymupdf'. Chạy: pip install pymupdf")
    doc = fitz.open(path)
    try:
        if len(doc) == 0:
            return None
        page = doc[0]
        spans = []
        for block in page.get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                for sp in line.get("spans", []):
                    text = unicodedata.normalize("NFC", sp.get("text", "")).strip()
                    if not text:
                        continue
                    bold = bool(sp.get("flags", 0) & 16) or "bold" in (sp.get("font") or "").lower()
                    bbox = sp["bbox"]
                    spans.append({"y0": bbox[1], "x0": bbox[0], "text": text, "bold": bold})
        spans.sort(key=lambda s: (round(s["y0"]), s["x0"]))

        anchor_idx = next((i for i, s in enumerate(spans)
                            if re.match(r"ph[ầâ]n\s*i\b", s["text"], re.IGNORECASE)), None)
        if anchor_idx is None:
            return None
        label_idx = next((i for i in range(anchor_idx + 1, len(spans)) if spans[i]["bold"]), None)
        if label_idx is None:
            return None
        i = label_idx
        while i < len(spans) and spans[i]["bold"]:   # nhãn có thể tách nhiều mảnh in đậm liền nhau
            i += 1
        parts = []
        while i < len(spans) and not spans[i]["bold"]:
            parts.append(spans[i]["text"])
            i += 1
        content = " ".join(parts).strip()
        return content or None
    finally:
        doc.close()

# ---------- Đọc thẳng .docx (không cần đợi Word chuyển sang PDF) ----------
# Mục tiêu: tự điền NGAY khi vừa chọn file .docx, thay vì phải đợi Word convert xong (có thể
# mất vài giây tới vài chục giây). Việc chuyển sang PDF (convert_office_doc_to_pdf) vẫn chạy
# như cũ — vẫn cần PDF thật để xem trước/đóng dấu/upload — chỉ là bước ĐỌC ĐỂ TỰ ĐIỀN không
# còn phải chờ nó nữa. Dùng lại đúng các hàm tách chữ (extract_code_from_text/
# extract_abstract_from_text/extract_doc_type_title/extract_abstract_after_title) như bản PDF,
# chỉ khác nguồn lấy "lines".

DOCX_PAGE1_MAX_LINES = 40   # .docx không có ranh giới trang thật (khác PDF) — lấy tạm N đoạn
                            # văn đầu tiên làm "trang 1", đủ rộng cho quốc hiệu/số ký hiệu/tên
                            # loại/trích yếu của mọi mẫu văn bản hành chính đã gặp

def _iter_docx_paragraphs(document):
    """Duyệt TOÀN BỘ đoạn văn theo đúng thứ tự xuất hiện trong file, kể cả đoạn nằm trong Ô
    BẢNG — mẫu Công văn chuẩn hành chính thường đặt khối "Số:.../V/v..." trong 1 bảng 2 cột
    (quốc hiệu bên phải, tên cơ quan bên trái), python-docx mặc định tách riêng
    document.paragraphs/document.tables và BỎ QUA chữ trong bảng, nên phải tự duyệt cây XML
    để không bỏ sót đúng đoạn quan trọng nhất."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    def _walk(parent_elm, parent):
        for child in parent_elm:
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                table = Table(child, parent)
                for row in table.rows:
                    for cell in row.cells:
                        yield from _walk(cell._tc, cell)

    yield from _walk(document.element.body, document)

def _run_is_bold(run, paragraph):
    """run.bold trả None nếu đậm/không đậm là do KẾ THỪA từ style (không đặt riêng ở run) —
    phải lần lên style cha của đoạn văn mới biết chắc, nếu không sẽ hiểu nhầm là "không đậm"."""
    if run.bold is not None:
        return run.bold
    style = paragraph.style
    while style is not None:
        b = style.font.bold
        if b is not None:
            return b
        style = style.base_style
    return False

def extract_draft_fields_docx(path):
    """Như extract_draft_fields() nhưng đọc thẳng .docx bằng python-docx — không cần Word."""
    if docx is None:
        raise RuntimeError("Chưa cài thư viện 'python-docx'. Chạy: pip install python-docx")
    d = docx.Document(path)
    lines = [unicodedata.normalize("NFC", p.text) for p in _iter_docx_paragraphs(d) if p.text.strip()]
    lines = lines[:DOCX_PAGE1_MAX_LINES]
    if not lines:
        return None, None, None
    text = "\n".join(lines)
    code = extract_code_from_text(text)
    doc_type, title_idx = extract_doc_type_title(lines)
    if doc_type is not None:
        abstract = extract_abstract_after_title(lines, title_idx)
        return doc_type, code, abstract
    return None, code, extract_abstract_from_text(text)

def extract_phieu_trinh_content_docx(path):
    """Như extract_phieu_trinh_content() nhưng đọc thẳng .docx bằng python-docx — không cần
    Word. In đậm đọc từ run.bold (kèm suy ra từ style cha qua _run_is_bold), đáng tin hơn cách
    đoán qua tên font trong bản PDF. Các run liền kề CÙNG trạng thái đậm/thường trong 1 đoạn
    văn được gộp thành 1 "khối" — Word hay tách run vụn vặt (gõ sửa, gạch chân chính tả...)
    không liên quan gì tới định dạng đậm/thường thật, gộp lại để khỏi vỡ thuật toán dò nhãn."""
    if docx is None:
        raise RuntimeError("Chưa cài thư viện 'python-docx'. Chạy: pip install python-docx")
    d = docx.Document(path)
    spans = []
    for p in _iter_docx_paragraphs(d):
        cur_text, cur_bold = "", None
        for r in p.runs:
            t = unicodedata.normalize("NFC", r.text)
            if not t:
                continue
            b = _run_is_bold(r, p)
            if b == cur_bold:
                cur_text += t
            else:
                if cur_text.strip():
                    spans.append({"text": cur_text.strip(), "bold": cur_bold})
                cur_text, cur_bold = t, b
        if cur_text.strip():
            spans.append({"text": cur_text.strip(), "bold": cur_bold})

    anchor_idx = next((i for i, s in enumerate(spans)
                        if re.match(r"ph[ầâ]n\s*i\b", s["text"], re.IGNORECASE)), None)
    if anchor_idx is None:
        return None
    label_idx = next((i for i in range(anchor_idx + 1, len(spans)) if spans[i]["bold"]), None)
    if label_idx is None:
        return None
    i = label_idx
    while i < len(spans) and spans[i]["bold"]:
        i += 1
    parts = []
    while i < len(spans) and not spans[i]["bold"]:
        parts.append(spans[i]["text"])
        i += 1
    content = " ".join(parts).strip()
    return content or None

def flow_id_for_code(code, flow_store):
    """Suy ra luồng trình theo ký hiệu văn bản, dùng quy tắc "từ khoá -> luồng" của CHÍNH máy
    này (flow_store["rules"], thứ tự ưu tiên = thứ tự trong list) — không còn hardcode flowId
    trong code, vì mỗi máy/mỗi tài khoản có thể có luồng khác nhau. Không khớp quy tắc nào thì
    rơi về flow_store["default_flow_id"] (nếu máy này đã đặt)."""
    n = _norm(code or "")
    for rule in flow_store.get("rules", []):
        kw = rule.get("keyword")
        if kw and _norm(kw) in n:
            return rule.get("flowId")
    return flow_store.get("default_flow_id")

def flow_id_for_doc(doc_type, code, flow_store):
    """Như flow_id_for_code(), nhưng loại VB có luồng cố định (flow_store["doc_type_rules"])
    thì ưu tiên luồng đó, không cần xét ký hiệu."""
    forced = flow_store.get("doc_type_rules", {})
    if doc_type in forced:
        return forced[doc_type]
    return flow_id_for_code(code, flow_store)

# ---------- Đánh số chữ ký lên PDF dự thảo (theo luồng trình) ----------
# Vị trí ký = chú thích PDF (Text annot, icon "Comment", nội dung = số) — không phải chữ in
# lên trang. Số đóng dấu = đúng "order" của bước trong luồng đang chọn (actionType∈{1,4,5} =
# bước "ký" thật, khác bước phê duyệt/cấp số/ban hành) — xác nhận khớp 100% qua HAR trên cả 3
# luồng cá nhân cũ (order 1/2-3/5 khớp đúng số đã hardcode trước đây). Chức danh trên PDF được
# so khớp với `roleName` thật của luồng đang chọn (không còn giới hạn 4 tên cứng) — nên áp dụng
# được cho BẤT KỲ luồng nào, không riêng 3 luồng cũ.
SIG_CHUYEN_VIEN_ROLE = "CHUYÊN VIÊN"   # quy ước cố định, luôn = 0 — KHÔNG phải 1 bước thật của
                                        # luồng nào (không xuất hiện trong searchNodeInFlow.do ở
                                        # cả 3 luồng cá nhân đã kiểm chứng), là người soạn thảo,
                                        # đứng trước cả bước ký đầu tiên của luồng.

def _norm_role_text(s):
    return unicodedata.normalize("NFC", (s or "")).strip().upper()

def build_role_number_map(items):
    """Dựng {chức danh viết HOA: số đóng dấu} từ node list của luồng đang chọn (kết quả
    searchNodeInFlow.do / resolve_flow_signers, đã có order/actionType/roleName) — thay cho
    bảng cứng trước đây. Chỉ lấy bước "ký" thật (actionType 1/4/5); actionType 0 (phê duyệt),
    2/3 (văn thư cấp số/ban hành) không đóng dấu số, không đưa vào bảng."""
    m = {SIG_CHUYEN_VIEN_ROLE: 0}
    for n in items or []:
        if n.get("actionType") in (1, 4, 5):
            rn = n.get("roleName")
            if rn:
                m[_norm_role_text(rn)] = n.get("order")
    return m

def _is_upper_title(text):
    """Dòng chức danh thật luôn viết HOA hết (phân biệt với câu văn thường có nhắc tới
    tên vai trò, ví dụ 'Kính gửi: ..., Phó Cục trưởng ...' — câu này không viết hoa hết)."""
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c == c.upper() for c in letters)

def _sig_page_lines(page):
    """Toạ độ từng dòng chữ trên trang, sắp theo thứ tự trên->dưới."""
    lines = []
    d = page.get_text("dict")
    for block in d.get("blocks", []):
        for ln in block.get("lines", []):
            text = "".join(sp.get("text", "") for sp in ln.get("spans", [])).strip()
            if not text:
                continue
            # Font nhúng trong PDF đôi khi ánh xạ chữ có dấu tiếng Việt ra Unicode không
            # chuẩn NFC (khác dạng chữ trong role_map) — chuẩn hoá về NFC để so khớp chuỗi
            # ("CỤC TRƯỞNG"...) không bị trật, dù nhìn bằng mắt vẫn giống hệt nhau.
            text = unicodedata.normalize("NFC", text)
            x0, y0, x1, y1 = ln["bbox"]
            lines.append({"text": text, "x0": x0, "y0": y0, "x1": x1, "y1": y1})
    lines.sort(key=lambda l: (round(l["y0"]), l["x0"]))
    return lines

# Mục lớn đánh số La Mã đứng riêng 1 dòng (vd "I.", "II.", "V.") — CHỮ HOA hết nên qua được
# _is_upper_title(), nhưng KHÔNG PHẢI chức danh. Nếu không loại, dòng 1 chữ cái như "V" khớp
# nhầm qua vế "up in role" vì "V" tình cờ là 1 ký tự con nằm trong "CHUYÊN VIÊN" — gây phát
# hiện "TRÙNG số" giả (2 mục La Mã khác nhau cùng bị gán số 0), chặn cả file dù không có gì sai
# thật (xác nhận qua ca thực tế: mục "I." và "V." trên cùng 1 trang, cả 2 báo trùng số 0).
_ROMAN_NUMERAL_RE = re.compile(r'^[IVXLCDM]+$')
SIG_MIN_TITLE_LEN = 4   # chức danh thật luôn là cụm nhiều chữ (vd "CỤC TRƯỞNG") — dòng ngắn hơn
                        # ngưỡng này không đủ để so khớp kiểu "nằm trong nhau" (tier 1/2), chỉ
                        # còn so khớp NGUYÊN VĂN (tier 3) là an toàn dù ngắn.

def _sig_find_hits(lines, role_map):
    """Các dòng CHỮ HOA khớp 1 chức danh trong role_map (của luồng đang chọn). Trả về
    [(chỉ số dòng, tên chức danh khớp được)]. So khớp CẢ HAI CHIỀU — vì roleName lấy từ server
    (vd "Phó Cục trưởng") có thể dài/ngắn hơn dòng in thật trên PDF (vd chỉ in "CỤC TRƯỞNG",
    không có "PHÓ"): ưu tiên khớp đúng nguyên văn, rồi tới chức danh NẰM TRONG dòng in (dòng in
    dài hơn, có thể có tiền tố "KT."/"PHÓ" phía trước), cuối cùng mới tới dòng in NẰM TRONG chức
    danh (dòng in ngắn/chung chung hơn chức danh thật của luồng). Loại trước các dòng quá ngắn/
    số La Mã đứng riêng — không thể là chức danh thật, chỉ gây khớp nhầm ở vế so khớp lỏng."""
    hits = []
    for i, ln in enumerate(lines):
        if not _is_upper_title(ln["text"]):
            continue
        up = ln["text"].upper()
        core = up.rstrip(".)").strip()
        if _ROMAN_NUMERAL_RE.match(core):
            continue
        best = None   # (mức ưu tiên, độ dài chức danh, tên chức danh)
        for role in role_map:
            if role == up:
                tier = 3
            elif len(up) >= SIG_MIN_TITLE_LEN and role in up:
                tier = 2
            elif len(up) >= SIG_MIN_TITLE_LEN and up in role:
                tier = 1
            else:
                continue
            cand = (tier, len(role), role)
            if best is None or cand > best:
                best = cand
        if best:
            hits.append((i, best[2]))
    return hits

def _sig_cluster_hits(lines, hits):
    """Gộp các dòng liền kề, cùng vai trò, không có khoảng trống (VD 'KT. CỤC TRƯỞNG' +
    'PHÓ CỤC TRƯỞNG') thành 1 vị trí ký — lấy dòng CUỐI cụm (chức danh thật) làm mốc.
    So sánh trực tiếp 2 hit liên tiếp (không đòi hỏi liền chỉ số trong 'lines'), vì cùng
    hàng ngang có thể xen dòng chấm chấm/nội dung cột khác không liên quan (VD khung
    'PHẦN III' bên trái chạy song song với khối chữ ký bên phải)."""
    clusters = []
    i = 0
    while i < len(hits):
        j = i
        while j + 1 < len(hits) and hits[j + 1][1] == hits[j][1]:
            cur, nxt = lines[hits[j][0]], lines[hits[j + 1][0]]
            gap = nxt["y0"] - cur["y1"]
            overlap = min(nxt["x1"], cur["x1"]) - max(nxt["x0"], cur["x0"])
            if gap <= 6 and overlap > 0:   # cùng cột (chồng lấn x) và không có khoảng trống
                j += 1
            else:
                break
        clusters.append((hits[j][0], hits[j][1]))   # (chỉ số dòng mốc, vai trò)
        i = j + 1
    return clusters

def _sig_find_name_line(lines, anchor_idx, max_gap=220):
    """Dòng tên người ký gần nhất phía dưới mốc, cùng cột (chồng lấn theo trục x)."""
    anchor = lines[anchor_idx]
    width = anchor["x1"] - anchor["x0"]
    for k in range(anchor_idx + 1, len(lines)):
        cand = lines[k]
        if cand["y0"] <= anchor["y1"]:
            continue
        if cand["y0"] - anchor["y1"] > max_gap:
            break
        overlap = min(cand["x1"], anchor["x1"]) - max(cand["x0"], anchor["x0"])
        if overlap > 0.3 * width:
            return cand
    return None

def find_signature_stamps(doc, flow_items, log):
    """Quét TOÀN BỘ file (mọi trang, kể cả phụ lục), trả về các vị trí cần đánh số.
    `flow_items`: node list của luồng đang chọn (searchNodeInFlow.do / resolve_flow_signers)."""
    role_map = build_role_number_map(flow_items)
    stamps = []
    for pno in range(len(doc)):
        lines = _sig_page_lines(doc[pno])
        hits = _sig_find_hits(lines, role_map)
        if not hits:
            continue
        for anchor_idx, role in _sig_cluster_hits(lines, hits):
            number = role_map.get(role)
            if number is None:
                log(f"   • Trang {pno+1}: thấy '{role}' nhưng luồng hiện tại không có số "
                    f"cho vai trò này — bỏ qua.")
                continue
            anchor = lines[anchor_idx]
            name_line = _sig_find_name_line(lines, anchor_idx)
            y_mid = ((anchor["y1"] + name_line["y0"]) / 2) if name_line else (anchor["y1"] + 40)
            x_mid = (anchor["x0"] + anchor["x1"]) / 2
            stamps.append({"page": pno, "x": x_mid, "y": y_mid, "number": number, "role": role,
                            "title": anchor["text"]})
            log(f"   • Trang {pno+1}: '{anchor['text']}' → số {number} (x={x_mid:.0f}, y={y_mid:.0f})")
    return stamps

# Watermark "đã xem/tải" hệ thống VOffice tự chèn MỖI LẦN 1 file được xem/tải qua web (kể cả
# lúc chương trình này tự tải file cũ về lúc "Sửa") — dạng chữ chéo "<tên đăng nhập>_<Họ
# tên>_<ngày/tháng/năm giờ:phút:giây>" in ở mọi trang. Không xoá thì file tải về lúc Sửa đã dính
# sẵn watermark, Trình lại xong xem/tải lần nữa sẽ bị chồng thêm 1 lớp mới nữa. Đã tự mở 1 file
# PDF thật (dry-run) để xác nhận: watermark là 1 khối lệnh vẽ chữ TÁCH BIỆT khỏi nội dung thật
# trong content stream của trang (không phải annotation, không phải lớp ẩn/OCG) — nên chỉ cần
# làm rỗng ĐÚNG câu lệnh vẽ ra đúng chuỗi chữ đó (khớp theo khuôn, neo theo tên đăng nhập đang
# dùng — không hardcode tên/tài khoản cụ thể nào, nên áp dụng được cho mọi tài khoản), không đụng
# tới bất kỳ lệnh vẽ nào khác — an toàn hơn nhiều so với xoá theo vùng hình chữ nhật (dễ ăn lẹm
# chữ thật nằm gần watermark do watermark bị xoay chéo).
_WATERMARK_TJ_RE = re.compile(rb'\(([^()\\]*)\)Tj')

def _watermark_pattern(username):
    return re.compile(re.escape(username).encode() + rb'_.+_\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2}:\d{2}')

def strip_view_watermark(path, username, log=lambda *a: None):
    """Xoá watermark "đã xem/tải" (nếu có) khỏi file PDF `path`, GHI ĐÈ ngay tại chỗ — chỉ nên
    gọi trên bản tải tạm về máy (vd lúc "Sửa"), không gọi trên file người dùng tự chọn. Trả về
    số watermark đã xoá; 0 nếu không tìm thấy (file có thể chưa dính watermark, hoặc không phải
    PDF, hoặc khuôn watermark khác đi — im lặng bỏ qua, không coi là lỗi)."""
    if fitz is None or not username or not path.lower().endswith(".pdf"):
        return 0
    pat = _watermark_pattern(username)
    try:
        doc = fitz.open(path)
    except Exception:
        return 0
    try:
        removed = 0
        for page in doc:
            for xref in page.get_contents():
                raw = doc.xref_stream(xref)
                def repl(m):
                    nonlocal removed
                    if pat.match(m.group(1)):
                        removed += 1
                        return b"()Tj"
                    return m.group(0)
                new_raw = _WATERMARK_TJ_RE.sub(repl, raw)
                if new_raw != raw:
                    doc.update_stream(xref, new_raw)
        if not removed:
            return 0
        out_path = path + ".nowm.tmp"
        doc.save(out_path, garbage=0, deflate=False)
    finally:
        doc.close()
    os.replace(out_path, path)
    log(f"   • Đã xoá {removed} watermark 'đã xem/tải' của hệ thống trên file vừa tải về.")
    return removed

def _sig_tagged_path(path):
    """Đường dẫn cho bản PDF đã đánh số — 1 thư mục tạm riêng (xem _new_gen_tmpdir), giữ
    NGUYÊN TÊN gốc bên trong (không ghi đè file gốc, không thêm ngày/hậu tố)."""
    out_dir = _new_gen_tmpdir("voffice_stamp_")
    return os.path.join(out_dir, os.path.basename(path))

def _clear_old_signature_stamps(doc, log):
    """Xoá các annotation số chữ ký do CHÍNH hàm này từng ghi ở lần đánh số trước (nếu file này
    đã đánh số rồi, VD Thu hồi → sửa → Trình lại đúng file cũ) — không xoá thì số mới bị chồng
    lên số cũ, nhìn như trùng/lộn xộn. Nhận diện đúng annotation của mình (không đụng ghi chú
    khác lỡ có sẵn trong file): dạng Text/icon "Comment" VÀ nội dung CHỈ là 1 số thuần — xem
    stamp_signature_numbers() (page.add_text_annot(point, str(số), icon="Comment"))."""
    removed = 0
    for page in doc:
        for annot in list(page.annots() or []):
            info = annot.info or {}
            if (annot.type[1] == "Text" and info.get("name") == "Comment"
                    and (info.get("content") or "").strip().isdigit()):
                page.delete_annot(annot)
                removed += 1
    if removed:
        log(f"   • Đã xoá {removed} số chữ ký cũ (đánh số lại từ đầu theo luồng hiện tại).")

def stamp_signature_numbers(path, flow_items, log, stamps=None):
    """Đọc 1 file PDF (phiếu trình hoặc dự thảo văn bản), tìm vị trí ký theo luồng, ghi
    chú thích (Text annot) số thứ tự. Lưu bản đã đánh số vào 1 thư mục tạm riêng, GIỮ NGUYÊN
    TÊN gốc (file gốc giữ nguyên, không sửa/không ghi đè). Trả về đường dẫn để upload —
    là bản đã đánh số nếu có đánh được gì, hoặc chính path gốc nếu không tìm thấy vị trí ký.
    Nếu `stamps` được truyền vào (VD: đã được người dùng chỉnh trong khung xem trước),
    dùng nguyên danh sách đó thay vì tự quét lại file. `flow_items`: node list của luồng
    đang chọn (xem flow_items_for_cfg)."""
    if fitz is None:
        raise PipelineError("Chưa cài thư viện 'pymupdf'. Chạy: pip install pymupdf")
    if not flow_items:
        raise PipelineError("Chưa chọn Luồng trình — không biết đánh số theo luồng nào.")

    doc = fitz.open(path)
    try:
        _clear_old_signature_stamps(doc, log)
        if stamps is None:
            stamps = find_signature_stamps(doc, flow_items, log)
        if not stamps:
            log("   ⚠ Không tìm thấy vị trí ký nào trong file — giữ nguyên file gốc, không đánh số.")
            return path

        by_number = {}
        for st in stamps:
            by_number.setdefault(st["number"], []).append(st)
        dups = {n: v for n, v in by_number.items() if len(v) > 1}
        if dups:
            detail = "; ".join(
                f"số {n} lặp {len(v)} lần (trang {', '.join(str(x['page']+1) for x in v)}, "
                f"chức danh: {', '.join(x['title'] for x in v)})"
                for n, v in dups.items())
            raise PipelineError(
                f"Phát hiện TRÙNG số chữ ký trong file (1 file không được trùng số) — "
                f"dừng lại để kiểm tra tay: {detail}")

        expected_max = max((n.get("order") for n in flow_items if n.get("actionType") in (1, 4, 5)),
                            default=None)
        actual_max = max(st["number"] for st in stamps)
        if expected_max is not None and actual_max != expected_max:
            log(f"   ⚠ Cảnh báo: luồng đang chọn thường có số cao nhất = {expected_max}, "
                f"nhưng file này chỉ thấy tới số {actual_max}. Kiểm tra lại file/luồng trước khi trình.")

        for st in stamps:
            page = doc[st["page"]]
            point = fitz.Point(st["x"] - 8, st["y"] - 8)
            annot = page.add_text_annot(point, str(st["number"]), icon="Comment")
            annot.update()

        out_path = _sig_tagged_path(path)
        try:
            doc.save(out_path)
        except Exception as e:
            raise PipelineError(f"Không lưu được file đã đánh số vào '{out_path}': {e}")
        log(f"   → Đã đánh {len(stamps)} số chữ ký, lưu tại: {out_path}")
        return out_path
    finally:
        doc.close()

# ---------- Ô chọn nơi nhận (tìm kiếm + chip + bộ mẫu) ----------
class RecipientBox(ttk.LabelFrame):
    CATS = [("inside", "Nhận nội bộ"), ("report", "Báo cáo"), ("edoc", "Liên thông"),
            ("save", "Nơi lưu"), ("know", "Để biết")]

    def __init__(self, parent, cay, store):
        super().__init__(parent, text="Nơi nhận  (gõ để tìm — không dấu/viết tắt đều được)", padding=8)
        self.store = store
        self.pools = {
            "inside": cay["internal"]["nodes"], "save": cay["internal"]["nodes"],
            "know": cay["internal"]["nodes"], "report": cay["internal"]["nodes"],
            "edoc": cay["lien_thong"]["nodes"],
        }
        self.buckets = {c: [] for c, _ in self.CATS}   # list các node
        self.cat = tk.StringVar(value="inside")

        rowc = ttk.Frame(self); rowc.pack(fill="x")
        ttk.Label(rowc, text="Thêm vào:").pack(side="left")
        for c, label in self.CATS:
            ttk.Radiobutton(rowc, text=label, value=c, variable=self.cat,
                            command=self._refilter).pack(side="left", padx=2)

        self.entry = ttk.Entry(self); self.entry.pack(fill="x", pady=(4, 0))
        self.entry.bind("<KeyRelease>", self._on_key)
        self.entry.bind("<Down>", lambda e: self._focus_list())
        self.entry.bind("<Return>", self._add_top)

        self.lb = tk.Listbox(self, height=8, activestyle="dotbox")
        self.lb.pack(fill="x")
        self.lb.bind("<Return>", self._add_selected)
        self.lb.bind("<Double-Button-1>", self._add_selected)
        self.lb.bind("<Button-3>", self._context)
        self.lb.bind("<Button-2>", self._context)

        self.chips = ttk.Frame(self); self.chips.pack(fill="x", pady=(6, 0))
        rowt = ttk.Frame(self); rowt.pack(fill="x", pady=(6, 0))
        ttk.Button(rowt, text="💾 Lưu bộ mẫu", command=self._save_template).pack(side="left")
        ttk.Button(rowt, text="📂 Nạp bộ mẫu", command=self._load_template).pack(side="left", padx=4)
        ttk.Label(rowt, text="  (chuột phải mục để Ghim ★)", foreground="gray").pack(side="left")

        self._matches = []
        self._debounce_id = None
        self._refilter(); self._render_chips()

    def _pool(self):
        return self.pools[self.cat.get()]

    def _label(self, nd):
        par = nd.get("parent")
        return nd["name"].strip() + (f"   ‹{par.strip()}›" if par else "")

    def _on_key(self, e):
        if e.keysym in ("Up", "Down", "Return"):
            return
        # Gộp các sự kiện KeyRelease liên tiếp (bộ gõ tiếng Việt hay phát sinh nhiều
        # sự kiện cho 1 ký tự có dấu) thành 1 lần lọc duy nhất sau khi ngừng gõ.
        if self._debounce_id is not None:
            self.after_cancel(self._debounce_id)
        self._debounce_id = self.after(120, self._debounced_refilter)

    def _debounced_refilter(self):
        self._debounce_id = None
        try:
            self._refilter()
        except tk.TclError:
            pass

    def _refilter(self, *_):
        self._matches = search_nodes(self.entry.get(), self._pool(), self.store)
        pinned = set(self.store.get("pinned", []))
        self.lb.delete(0, "end")
        for nd in self._matches:
            mark = "★ " if _node_key(nd) in pinned else "    "
            self.lb.insert("end", mark + self._label(nd))

    def _focus_list(self):
        if self._matches:
            self.lb.focus_set(); self.lb.selection_clear(0, "end")
            self.lb.selection_set(0); self.lb.activate(0)

    def _add_top(self, *_):
        if self._matches:
            self._add(self._matches[0])

    def _add_selected(self, *_):
        sel = self.lb.curselection()
        if sel:
            self._add(self._matches[sel[0]])

    def _add(self, node):
        c = self.cat.get()
        if not any(_node_key(x) == _node_key(node) for x in self.buckets[c]):
            self.buckets[c].append(node)
            bump_freq(self.store, _node_key(node))
        self.entry.delete(0, "end")
        self._refilter(); self._render_chips()
        self.entry.focus_set()

    def _remove(self, c, key):
        self.buckets[c] = [x for x in self.buckets[c] if _node_key(x) != key]
        self._render_chips()

    def _render_chips(self):
        for w in self.chips.winfo_children():
            w.destroy()
        any_sel = False
        for c, label in self.CATS:
            if not self.buckets[c]:
                continue
            any_sel = True
            row = ttk.Frame(self.chips); row.pack(fill="x", anchor="w", pady=1)
            ttk.Label(row, text=label + ":", width=12).pack(side="left", anchor="n")
            wrap = ttk.Frame(row); wrap.pack(side="left", fill="x")
            for nd in self.buckets[c]:
                chip = tk.Frame(wrap, bg="#e3f2fd", bd=1, relief="solid")
                chip.pack(side="left", padx=2, pady=1)
                tk.Label(chip, text=nd["name"].strip(), bg="#e3f2fd").pack(side="left", padx=(4, 0))
                tk.Button(chip, text="✕", bd=0, bg="#e3f2fd", padx=2,
                          command=lambda cc=c, k=_node_key(nd): self._remove(cc, k)).pack(side="left")
        if not any_sel:
            ttk.Label(self.chips, text="(chưa chọn nơi nhận nào)", foreground="gray").pack(anchor="w")

    def _context(self, e):
        idx = self.lb.nearest(e.y)
        if idx < 0 or idx >= len(self._matches):
            return
        key = _node_key(self._matches[idx])
        pinned = set(self.store.get("pinned", []))
        m = tk.Menu(self, tearoff=0)
        m.add_command(label=("Bỏ ghim" if key in pinned else "Ghim ★"),
                      command=lambda: (toggle_pin(self.store, key), self._refilter()))
        m.tk_popup(e.x_root, e.y_root)

    def _save_template(self):
        from tkinter import simpledialog
        if not any(self.buckets.values()):
            messagebox.showinfo("Bộ mẫu", "Chưa chọn nơi nhận nào để lưu."); return
        name = simpledialog.askstring("Lưu bộ mẫu", "Đặt tên bộ mẫu:")
        if name:
            self.store["templates"][name] = {c: list(v) for c, v in self.buckets.items()}
            save_store(self.store)
            messagebox.showinfo("Bộ mẫu", f"Đã lưu bộ mẫu '{name}'.")

    def _load_template(self):
        tpls = self.store.get("templates", {})
        if not tpls:
            messagebox.showinfo("Bộ mẫu", "Chưa có bộ mẫu nào."); return
        m = tk.Menu(self, tearoff=0)
        for name in tpls:
            m.add_command(label=name, command=lambda n=name: self._apply_template(n))
        m.tk_popup(self.winfo_pointerx(), self.winfo_pointery())

    def _apply_template(self, name):
        tpl = self.store["templates"].get(name, {})
        for c, _ in self.CATS:
            existing = {_node_key(x) for x in self.buckets[c]}
            for nd in tpl.get(c, []):
                if isinstance(nd, dict) and _node_key(nd) not in existing:
                    self.buckets[c].append(nd)
        self._render_chips()

    def get(self, cat):
        return list(self.buckets[cat])   # danh sách node

    def clear(self):
        for c, _ in self.CATS:
            self.buckets[c] = []
        self._render_chips()


class FileList(ttk.Frame):
    """Danh sách file: nút thêm + listbox + nút bỏ. get() -> [đường dẫn].
    Listbox tự giãn chiều cao theo đúng số file đang có (tối thiểu MIN_HEIGHT, tối đa
    MAX_HEIGHT dòng) — bộ trình có nhiều file thì thấy hết luôn, không bị cắt còn 2 dòng như
    trước; quá MAX_HEIGHT thì để phần cuộn của cả tab (_make_scrollable) lo tiếp."""
    MIN_HEIGHT = 2
    MAX_HEIGHT = 12

    def __init__(self, parent, label):
        super().__init__(parent)
        self.pack(fill="x", pady=2)
        self.paths = []
        top = ttk.Frame(self); top.pack(fill="x")
        ttk.Label(top, text=label, width=22).pack(side="left")
        ttk.Button(top, text="Thêm file…", command=self._add).pack(side="left")
        ttk.Button(top, text="Bỏ chọn", command=self._remove).pack(side="left", padx=4)
        self.lb = tk.Listbox(self, height=self.MIN_HEIGHT)
        self.lb.pack(fill="x", padx=(0, 0))

    def _resize(self):
        self.lb.config(height=max(self.MIN_HEIGHT, min(len(self.paths), self.MAX_HEIGHT)))

    def _add(self):
        ps = filedialog.askopenfilenames(filetypes=[("PDF", "*.pdf"), ("Tất cả", "*.*")])
        for p in ps:
            if p not in self.paths:
                self.paths.append(p); self.lb.insert("end", os.path.basename(p))
        self._resize()

    def _remove(self):
        sel = list(self.lb.curselection())
        for i in reversed(sel):
            self.lb.delete(i); del self.paths[i]
        self._resize()

    def add_path(self, p):
        """Thêm 1 đường dẫn từ NGOÀI (vd _apply_edit_data khi Sửa) — khác _add() ở chỗ không tự
        mở hộp thoại chọn file, chỉ nạp sẵn + tự giãn chiều cao như _add()."""
        if p not in self.paths:
            self.paths.append(p); self.lb.insert("end", os.path.basename(p))
            self._resize()

    def get(self):
        return list(self.paths)

    def clear(self):
        self.paths = []
        self.lb.delete(0, "end")
        self._resize()


class DocumentSection(ttk.LabelFrame):
    """1 văn bản trong nhóm VĂN BẢN — file (chính + tài liệu gửi kèm) + loại VB/số ký hiệu/
    trích yếu RIÊNG của văn bản này (1 phiếu trình có thể gửi nhiều văn bản khác loại/khác số).
    Chọn file chính sẽ tự điền loại VB/số/trích yếu (và tự chọn Luồng trình dùng chung) như
    màn hình cũ, chỉ khác là áp dụng riêng cho văn bản này."""

    def __init__(self, parent, app, on_remove):
        super().__init__(parent, text="Văn bản", padding=6)
        self.app = app
        self.on_remove = on_remove
        self._extract_seq = 0
        self.file_draft = tk.StringVar()

        # BƯỚC bạn cần làm — nổi bật, chữ đậm, luôn hiện.
        f = ttk.Frame(self); f.pack(fill="x", pady=2)
        ttk.Label(f, text="  Dự thảo văn bản:", width=18, font=("", 10, "bold")).pack(side="left")
        ttk.Entry(f, textvariable=self.file_draft).pack(side="left", fill="x", expand=True)
        ttk.Button(f, text="Chọn…", command=self._pick).pack(side="left")
        self.btn_remove = ttk.Button(f, text="✕", width=3, command=lambda: self.on_remove(self))
        self.btn_remove.pack(side="left", padx=(4, 0))

        self.extra = FileList(self, "  + Tài liệu gửi kèm:")

        f = ttk.Frame(self); f.pack(fill="x", pady=2)
        ttk.Label(f, text="  Loại văn bản:", width=18).pack(side="left")
        self.doc_type = ttk.Combobox(f, values=sorted(ENUMS["documentType"].keys()), state="readonly")
        self.doc_type.set("Công văn"); self.doc_type.pack(side="left", fill="x", expand=True)

        f = ttk.Frame(self); f.pack(fill="x", pady=2)
        ttk.Label(f, text="  Số/ký hiệu:", width=18).pack(side="left")
        self.code = ttk.Entry(f); self.code.pack(side="left", fill="x", expand=True)

        f = ttk.Frame(self); f.pack(fill="x", pady=2)
        ttk.Label(f, text="  Trích yếu:", width=18).pack(side="left")
        # tk.Text nhiều dòng thay vì ttk.Entry 1 dòng — trích yếu thường dài, Entry 1 dòng phải
        # cuộn ngang bằng chuột mới đọc hết được.
        self.abstract = tk.Text(f, height=2, wrap="word")
        self.abstract.pack(side="left", fill="x", expand=True)

        # Cảnh báo đọc file tự động thất bại — hiện ngay tại văn bản liên quan thay vì chỉ ghi
        # vào log kỹ thuật (đã ẩn khỏi giao diện chính, xem App.log).
        self.warn_label = ttk.Label(self, text="", foreground="#c62828", wraplength=420, justify="left")
        self.warn_label.pack(anchor="w", pady=(2, 0))

    def _set_warning(self, msg):
        self.warn_label.config(text="⚠ " + msg)

    def _clear_warning(self):
        self.warn_label.config(text="")

    def _pick(self):
        p = filedialog.askopenfilename(
            filetypes=[("PDF/Word", "*.pdf *.docx"), ("Tất cả", "*.*")])
        if not p:
            return
        self.file_draft.set(p)
        self._extract_seq += 1
        seq = self._extract_seq
        threading.Thread(target=self._prepare_worker, args=(p, seq), daemon=True).start()

    def _prepare_worker(self, path, seq):
        """Đọc nhanh để tự điền ngay khi vừa chọn file — .docx qua python-docx (không cần
        Word), .pdf qua pypdf như cũ. Việc chuyển .docx sang PDF THẬT không còn làm ở đây nữa —
        dời sang lúc bấm CHẠY (xem App._run/_run_after_conversion): gộp lại thành 1 lượt gọi
        Word DUY NHẤT, tuần tự cho mọi văn bản trong phiếu trình, tránh nhiều luồng nền cùng lúc
        tranh nhau 1 tiến trình Word (từng gây lỗi 'Word.Application.Quit'), đồng thời đảm bảo
        PDF luôn sẵn sàng trước khi khung Xem trước mở ra."""
        if path.lower().endswith(".docx"):
            self._extract_docx_worker(path, seq)
            return
        if not path.lower().endswith(".pdf"):
            return   # định dạng khác PDF/Word — không tự đọc được, để nguyên cho người dùng tự điền
        self._extract_worker(path, seq)

    def _extract_docx_worker(self, path, seq):
        """Đọc thẳng .docx bằng python-docx, không cần Word."""
        doc_type = code = abstract = None
        err = None
        try:
            doc_type, code, abstract = extract_draft_fields_docx(path)
        except Exception as e:
            err = str(e)
        if err:
            msg = (f"Không đọc nhanh được .docx để tự điền: {err} — vẫn có thể tự điền lại "
                   "sau khi bấm CHẠY (lúc đó có bản PDF).")
            self.app.after(0, lambda: self._set_warning(msg))
            return
        self.app.after(0, self._apply_extract, seq, doc_type, code, abstract, None)

    def _extract_worker(self, path, seq):
        doc_type = code = abstract = None
        err = None
        try:
            doc_type, code, abstract = extract_draft_fields(path)
        except Exception as e:
            err = str(e)
        self.app.after(0, self._apply_extract, seq, doc_type, code, abstract, err)

    def _apply_extract(self, seq, doc_type, code, abstract, err):
        if seq != self._extract_seq:
            return   # đã chọn file khác cho văn bản này trong lúc đọc — bỏ kết quả cũ
        if err:
            self._set_warning(f"Không đọc được PDF để tự điền: {err}")
            return
        self._clear_warning()
        if doc_type:
            self.doc_type.set(doc_type)
        if code:
            self.code.delete(0, "end"); self.code.insert(0, code)
            fid = flow_id_for_doc(doc_type, code, self.app.flow_store)
            fname = self.app._flow_name_for_id(fid) if fid else None
            if fname:
                self.app.flow.set(fname)
                self.app._on_flow_changed()
        if abstract:
            self.abstract.delete("1.0", "end"); self.abstract.insert("1.0", abstract)
            if not self.app.report_content.get("1.0", "end-1c").strip():   # không ghi đè nếu đã có/văn bản khác đã điền
                self.app.report_content.delete("1.0", "end"); self.app.report_content.insert("1.0", abstract)

    def get(self):
        return {
            "doc_type": self.doc_type.get(),
            "code": self.code.get(),
            "abstract": self.abstract.get("1.0", "end-1c"),
            "file_draft_main": self.file_draft.get(),
            "files_draft_extra": self.extra.get(),
            "_existing_pid": getattr(self, "_existing_pid", None),
            "_existing_attach_ids": getattr(self, "_existing_attach_ids", None),
        }


# ---------- Cuộn bằng con lăn chuột (dùng chung cho mọi Canvas cuộn được) ----------
# Không dùng cách bind_all/unbind_all trên "Enter"/"Leave" của chính Canvas: các widget con
# (Entry, Label, ảnh...) phủ gần kín Canvas nên "Leave" kích hoạt gần như ngay khi con trỏ
# chạm vào bất kỳ widget con nào — khiến cuộn chuột gần như không hoạt động. Thay vào đó,
# đăng ký MỖI Canvas cuộn được vào 1 sổ dùng chung, và bind_all 1 LẦN DUY NHẤT cho cả ứng
# dụng; khi có sự kiện, tra xem con trỏ đang ở trong Canvas nào (dò lên theo .master) rồi
# cuộn đúng Canvas đó. Nhiều cửa sổ cùng có Canvas cuộn được (VD màn chính + khung xem trước)
# vẫn hoạt động độc lập, không "unbind" mất phần cuộn của cửa sổ kia khi 1 cửa sổ đóng lại.
_SCROLLABLES = []   # [(canvas, cuộn_ngang: bool)]
_WHEEL_HANDLERS_BOUND = False

def _wheel_target(canvas, x_root, y_root):
    try:
        w = canvas.winfo_containing(x_root, y_root)
    except (tk.TclError, KeyError):
        return False
    while w is not None:
        if w is canvas:
            return True
        w = w.master
    return False

def _dispatch_wheel(e, horizontal):
    for canvas, allow_h in list(_SCROLLABLES):
        try:
            if not canvas.winfo_exists():
                _SCROLLABLES.remove((canvas, allow_h)); continue
        except tk.TclError:
            continue
        if horizontal and not allow_h:
            continue
        if _wheel_target(canvas, e.x_root, e.y_root):
            try:
                (canvas.xview_scroll if horizontal else canvas.yview_scroll)(
                    int(-e.delta / 120), "units")
            except tk.TclError:
                pass
            return

def bind_mousewheel_scroll(canvas, horizontal=False):
    """Cho phép cuộn `canvas` bằng con lăn chuột dù con trỏ đang ở trên widget con nào
    bên trong nó. `horizontal=True` để thêm Shift+cuộn = cuộn ngang."""
    global _WHEEL_HANDLERS_BOUND
    _SCROLLABLES.append((canvas, horizontal))
    canvas.bind("<Destroy>", lambda e: _SCROLLABLES.remove((canvas, horizontal))
                if (canvas, horizontal) in _SCROLLABLES else None)
    if not _WHEEL_HANDLERS_BOUND:
        canvas.bind_all("<MouseWheel>", lambda e: _dispatch_wheel(e, False))
        canvas.bind_all("<Shift-MouseWheel>", lambda e: _dispatch_wheel(e, True))
        _WHEEL_HANDLERS_BOUND = True

def make_scrollable_frame(parent):
    """Bọc nội dung trong Canvas + thanh cuộn dọc + cuộn chuột. Trả về frame bên trong để
    pack nội dung vào."""
    outer = ttk.Frame(parent); outer.pack(fill="both", expand=True)
    # tk.Canvas không tự đổi màu theo theme sáng/tối như các ô ttk khác (mặc định trắng) — dò
    # đúng màu nền ttk hiện tại để không nổi thành 1 mảng trắng lạc quẻ khi hệ thống đang Dark
    # Mode (đã xảy ra thật, xem phản hồi người dùng).
    bg = ttk.Style().lookup("TFrame", "background") or parent.winfo_toplevel().cget("bg")
    canvas = tk.Canvas(outer, highlightthickness=0, bg=bg)
    vsb = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=vsb.set)
    canvas.pack(side="left", fill="both", expand=True)
    vsb.pack(side="right", fill="y")

    inner = ttk.Frame(canvas)
    win = canvas.create_window((0, 0), window=inner, anchor="nw")

    def _on_inner_configure(_e):
        canvas.configure(scrollregion=canvas.bbox("all"))
    inner.bind("<Configure>", _on_inner_configure)

    def _on_canvas_configure(e):
        canvas.itemconfig(win, width=e.width)
    canvas.bind("<Configure>", _on_canvas_configure)

    bind_mousewheel_scroll(canvas)
    return inner


class FlowSignerPanel(ttk.LabelFrame):
    """Tải danh sách bước của LUỒNG ĐANG CHỌN (bất kỳ luồng nào, không riêng luồng nào) — bước
    nào đã có sẵn người (vd luồng cá nhân tự tạo, bake sẵn người ký) thì bỏ qua âm thầm; bước
    nào máy tự biết chắc người (đúng 1 ứng viên) thì tự điền; bước nào không chắc (≥2 ứng viên)
    thì hiện ô cho người dùng chọn. Tự ẩn nếu luồng đang chọn đã có sẵn đủ người (không cần hỏi
    gì thêm) — tự hiện nếu có ít nhất 1 bước cần tự điền/chọn, để minh bạch cho người dùng thấy."""
    def __init__(self, parent, session, log, flow_store):
        super().__init__(parent, text="Chọn người ký cho luồng vừa chọn", padding=8)
        self.session = session
        self.log = log
        self.flow_store = flow_store   # sổ cục bộ — nhớ người hay chọn cho từng (luồng, bước)
        self._nodes = []          # kết quả resolve_flow_signers() cho luồng đang hiện
        self._pick_vars = {}      # nodeId -> (tk.StringVar, [candidates])
        self._loading = False
        self._flow_id = None
        self.status = ttk.Label(self, text="", foreground="gray")
        self.status.pack(anchor="w")
        self.rows_frame = ttk.Frame(self)
        self.rows_frame.pack(fill="x", pady=(4, 0))

    def load(self, flow_id, doc_abstract):
        """Tải lại toàn bộ cho `flow_id` mới — huỷ kết quả cũ. Tự ẩn trong lúc tải (tránh hiện
        nhầm kết quả của luồng trước đó); CHẠY/GỬI bị chặn qua is_ready() cho tới khi xong."""
        self._flow_id = flow_id
        self._nodes = []
        self._pick_vars = {}
        self._loading = True
        self.pack_forget()
        for w in self.rows_frame.winfo_children():
            w.destroy()
        self.status.config(text="Đang tải người ký cho luồng này…", foreground="gray")

        def worker():
            try:
                nodes = resolve_flow_signers(self.session, flow_id, doc_abstract, self.log)
            except Exception as e:
                err_msg = str(e)   # tính ngay trong khối except — "e" bị Python xoá khi except kết thúc
                self.after(0, lambda: self._load_failed(flow_id, err_msg))
                return
            self.after(0, lambda: self._load_done(flow_id, nodes))
        threading.Thread(target=worker, daemon=True).start()

    def _load_failed(self, flow_id, e):
        if flow_id != self._flow_id:
            return   # đã đổi sang luồng khác trong lúc tải — bỏ kết quả cũ
        self._loading = False
        self.status.config(text=f"✖ Không tải được: {e}", foreground="#c62828")
        self.pack(fill="x", pady=(6, 0))

    def _load_done(self, flow_id, nodes):
        if flow_id != self._flow_id:
            return
        self._loading = False
        self._nodes = nodes
        need_pick = [n for n in nodes if len(n.get("candidates") or []) >= 2]
        auto = [n for n in nodes if n.get("userId") is not None and len(n.get("candidates") or []) < 2]
        if not need_pick and not any(n.get("candidates") for n in nodes):
            self.pack_forget()   # luồng đã có sẵn đủ người (vd luồng cá nhân) — không cần hỏi gì
            return
        self.pack(fill="x", pady=(6, 0))
        self.status.config(
            text=f"{len(nodes)} bước — {len(auto)} tự điền, {len(need_pick)} cần bạn chọn.",
            foreground="gray")
        for n in nodes:
            # Xếp DỌC (tên bước ở trên, combobox/tên người ở dưới) thay vì nằm ngang — nằm
            # ngang trước đây chia đôi bề rộng hàng cho nhãn + combobox, tên người dài dễ bị
            # cắt chữ, phải kéo rộng cả cửa sổ mới đọc hết. Xếp dọc thì combobox luôn được toàn
            # bộ bề rộng khung, không phụ thuộc cửa sổ to hay nhỏ.
            row = ttk.Frame(self.rows_frame); row.pack(fill="x", pady=(2, 4))
            ttk.Label(row, text=n.get("name") or "(không rõ tên bước)", wraplength=380,
                      justify="left", foreground="gray").pack(anchor="w", fill="x")
            cands = n.get("candidates") or []
            if len(cands) >= 2:
                # Ghi kèm chức danh vào nhãn — 1 bước có thể gồm nhiều biến thể chức danh khác
                # nhau (vd "Cục Trưởng" chỉ 1 người, "Phó Cục Trưởng" 3 người khác), phải phân
                # biệt rõ đang chọn ai VỚI TƯ CÁCH gì, không chỉ chọn tên.
                pref_id = preferred_signer(self.flow_store, flow_id, n["nodeId"], cands)
                default_idx = 0
                if pref_id is not None:
                    default_idx = next(i for i, c in enumerate(cands) if c.get("userId") == pref_id)
                names = [f"{c.get('fullName') or ''} — {c.get('roleName') or ''}"
                         + ("  (lần trước)" if i == default_idx and pref_id is not None else "")
                         for i, c in enumerate(cands)]
                var = tk.StringVar(value=names[default_idx])
                n["userId"] = cands[default_idx].get("userId")
                n["fullName"] = cands[default_idx].get("fullName")
                n["roleId"] = cands[default_idx].get("roleId")
                n["roleName"] = cands[default_idx].get("roleName")
                cb = ttk.Combobox(row, values=names, textvariable=var, state="readonly")
                cb.pack(fill="x")
                def on_pick(_e, n=n, var=var, cands=cands, names=names):
                    idx = names.index(var.get())
                    n["userId"] = cands[idx].get("userId")
                    n["fullName"] = cands[idx].get("fullName")
                    n["roleId"] = cands[idx].get("roleId")
                    n["roleName"] = cands[idx].get("roleName")
                cb.bind("<<ComboboxSelected>>", on_pick)
                self._pick_vars[n["nodeId"]] = (var, cands)
            else:
                name = n.get("fullName") or "(chưa xác định được người)"
                fg = "black" if n.get("userId") is not None else "#c62828"
                ttk.Label(row, text=name, foreground=fg).pack(anchor="w")

    def is_ready(self):
        """False nếu còn đang tải, hoặc còn bước nào chưa có người (kể cả chưa chọn xong)."""
        if self._loading or not self._nodes:
            return False
        return all(n.get("userId") is not None for n in self._nodes)

    def get_nodes(self):
        return self._nodes


class _ConvertingDialog(tk.Toplevel):
    """Cửa sổ chờ nhỏ, hiện trong lúc chuyển .docx -> PDF tuần tự ngay sau khi bấm CHẠY (trước
    khi mở khung Xem trước hoặc chạy Chỉ kiểm tra) — để người dùng biết đang có việc chạy ngầm
    (Word có thể mất vài giây tới vài chục giây), không phải chương trình bị treo."""

    # Kích thước CỐ ĐỊNH (không tự co theo độ dài chữ) — đã đo qua các dòng trạng thái thực tế
    # dùng cửa sổ này (tên file/số văn bản dài, đường dẫn...); chữ dài hơn tự xuống dòng
    # (wraplength) thay vì đẩy cửa sổ giãn ra hoặc bị cắt.
    WIDTH = 460
    HEIGHT = 120

    def __init__(self, master, text):
        super().__init__(master)
        self.title("Đang chuẩn bị…")
        self.resizable(False, False)
        self.transient(master)
        self.protocol("WM_DELETE_WINDOW", lambda: None)   # không cho tự đóng giữa chừng
        self.label = ttk.Label(self, text=text, padding=24, wraplength=self.WIDTH - 48,
                                justify="left", anchor="center")
        self.label.pack(fill="both", expand=True)
        self.update_idletasks()
        x = master.winfo_rootx() + master.winfo_width() // 2 - self.WIDTH // 2
        y = master.winfo_rooty() + master.winfo_height() // 2 - self.HEIGHT // 2
        self.geometry(f"{self.WIDTH}x{self.HEIGHT}+{max(x, 0)}+{max(y, 0)}")
        self.grab_set()

    def set_status(self, text):
        self.label.config(text=text)

    def close(self):
        try:
            self.grab_release()   # xem _close_dlg() ở _open_role_dialog — tránh kẹt input trên macOS
        except Exception:
            pass
        self.destroy()


class SendLogWindow(tk.Toplevel):
    """Cửa sổ trạng thái riêng cho Lưu dự thảo/Trình văn bản (xem PreviewWindow._send) — ĐỘC LẬP
    với PreviewWindow, để người dùng luôn thấy tiến trình từng bước ngay cả khi PreviewWindow bị
    che khuất. Cùng kiểu nhẹ nhàng (nền sáng, chỉ 1 khung chữ) như _ConvertingDialog ("Sửa phiếu
    trình", "Tải danh sách phiếu trình") — KHÔNG phải khung log kỹ thuật nền đen; nội dung chỉ
    gồm đúng các bước lớn (PIPELINE_PHASES) người dùng đã quen thấy ở checklist, không phải log
    HTTP chi tiết (xem PreviewWindow._set_phase — nơi duy nhất gọi append() cho cửa sổ này).
    Không có nút X (không tự đóng giữa chừng) — chỉ đóng được qua nút "Đóng" (tự đếm ngược khi đã
    có kết quả, xem start_close_countdown), và đóng cửa sổ này sẽ đóng LUÔN PreviewWindow đi kèm
    (xem `on_close` truyền vào lúc tạo — PreviewWindow._finish_post_send)."""

    WIDTH = 460
    HEIGHT = 260

    def __init__(self, master, title, on_close):
        super().__init__(master)
        # transient (không phải grab_set — cửa sổ này vẫn tương tác được bình thường, không
        # khoá gì) để Windows gom nó vào cùng nhóm với cửa sổ cha trên taskbar thay vì hiện
        # thành 1 nút riêng biệt (macOS Dock vốn đã gom sẵn nên không thấy khác biệt).
        self.transient(master)
        self.title(title)
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", lambda: None)
        self._on_close_cb = on_close
        self._countdown_id = None
        self._lines = []

        self.label = ttk.Label(self, text="", padding=20, justify="left", anchor="nw",
                                wraplength=self.WIDTH - 40)
        self.label.pack(fill="both", expand=True)

        bottom = ttk.Frame(self, padding=(12, 8))
        bottom.pack(fill="x")
        # Vô hiệu tới khi có kết quả (xem start_close_countdown) — trong lúc đang chạy chưa có
        # gì để "đóng xong xuôi", đóng giữa chừng dễ hiểu nhầm là huỷ được việc đang gửi.
        self.btn_close = ttk.Button(bottom, text="Đóng", command=self._close_now, state="disabled")
        self.btn_close.pack(side="right")

        self.update_idletasks()
        x = master.winfo_rootx() + master.winfo_width() // 2 - self.WIDTH // 2
        y = master.winfo_rooty() + master.winfo_height() // 2 - self.HEIGHT // 2
        self.geometry(f"{self.WIDTH}x{self.HEIGHT}+{max(x, 0)}+{max(y, 0)}")

    def append(self, msg):
        self._lines.append(msg)
        self.label.config(text="\n".join(self._lines))

    def start_close_countdown(self, seconds=5):
        self.btn_close.config(state="normal")
        self._tick(seconds)

    def _tick(self, n):
        if n <= 0:
            self._close_now()
            return
        self.btn_close.config(text=f"Đóng ({n})")
        self._countdown_id = self.after(1000, lambda: self._tick(n - 1))

    def _close_now(self):
        # Bấm tay lúc nào cũng đóng ngay (không "huỷ đếm ngược" — đây là nút Đóng thật, không
        # phải nút Huỷ) — chỉ huỷ đúng cái hẹn giờ còn treo để khỏi gọi đóng 2 lần.
        if self._countdown_id is not None:
            self.after_cancel(self._countdown_id)
            self._countdown_id = None
        self.destroy()
        self._on_close_cb()


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Trợ lý trình văn bản")
        self.session = None
        self.settings = load_settings()
        self.store = load_store()
        self.flow_store = load_flow_store()
        self.file_report = tk.StringVar()
        self.doc_sections = []
        self._edit_tmpdirs = []   # thư mục tạm chứa file tải về khi "Sửa" — dọn ở _reset_form()/thoát
        self._readiness_after_id = None   # id lịch self.after() của _refresh_readiness — xem _cancel_readiness_loop
        atexit.register(self._cleanup_edit_tmpdirs)
        self.container = ttk.Frame(self); self.container.pack(fill="both", expand=True)
        self._show_login()

    def report_callback_exception(self, exc_type, exc_value, exc_tb):
        """Tkinter tự gọi hàm này cho MỌI lỗi xảy ra trong callback GUI (bấm nút, gõ phím...) —
        mặc định chỉ in ra stderr rồi thôi (không ai thấy ở bản đóng gói không có cửa sổ
        console). Đây là nguồn lỗi không lường trước phổ biến nhất trong 1 app tkinter, nên ghi
        vào app_log.txt trước khi vẫn giữ hành vi mặc định (in stderr) để debug tại chỗ khi cần."""
        _file_logger.error("LỖI TRONG GIAO DIỆN (callback):\n"
                            + "".join(traceback.format_exception(exc_type, exc_value, exc_tb)))
        super().report_callback_exception(exc_type, exc_value, exc_tb)

    def _clear(self):
        for w in self.container.winfo_children():
            w.destroy()

    def _row(self, parent, label, bold=False):
        f = ttk.Frame(parent); f.pack(fill="x", padx=12, pady=3)
        font = ("", 10, "bold") if bold else ("", 10)
        ttk.Label(f, text=label, width=18, font=font).pack(side="left")
        return f

    # ---------- MÀN 1: ĐĂNG NHẬP ----------
    def _cancel_readiness_loop(self):
        """Huỷ vòng self.after(700, self._refresh_readiness) đang chạy (nếu có) — bắt buộc phải
        gọi trước khi rời màn hình chính (_show_login gọi ở đây), nếu không mỗi lần Đăng xuất
        rồi Đăng nhập lại trong CÙNG 1 lần mở app sẽ cộng dồn thêm 1 vòng lặp 700ms chạy song
        song vĩnh viễn (vòng cũ không tự dừng vì nó tự lịch lại chính nó trên `self`, mà `self`
        — App/cửa sổ gốc — không bị huỷ khi đăng xuất, chỉ các widget con bên trong bị xoá)."""
        if self._readiness_after_id is not None:
            try:
                self.after_cancel(self._readiness_after_id)
            except tk.TclError:
                pass
            self._readiness_after_id = None

    def _show_login(self, auto=True):
        """`auto`: True (mặc định — lúc mở chương trình) — nếu đã "Nhớ đăng nhập" từ trước thì tự
        bấm ĐĂNG NHẬP luôn, không cần người dùng làm gì. `_do_logout` gọi auto=False để KHÔNG tự
        nhảy lại vào đúng tài khoản vừa đăng xuất — bắt buộc phải tự bấm lại."""
        self._cancel_readiness_loop()
        self._clear()
        self.geometry("440x380")
        pad = ttk.Frame(self.container, padding=16); pad.pack(fill="both", expand=True)
        ttk.Label(pad, text="Đăng nhập hệ thống", font=("", 12, "bold")).pack(anchor="w", pady=(0, 8))
        f = self._row(pad, "Tên đăng nhập:")
        self.username = ttk.Entry(f); self.username.pack(side="left", fill="x", expand=True)
        f = self._row(pad, "Mật khẩu:")
        self.password = ttk.Entry(f, show="•"); self.password.pack(side="left", fill="x", expand=True)
        self.password.bind("<Return>", lambda e: self._do_login())

        self.remember = tk.BooleanVar(value=bool(self.settings.get("remember")))
        ttk.Checkbutton(pad, text="Nhớ đăng nhập (mật khẩu lưu an toàn trong Keychain)",
                        variable=self.remember).pack(anchor="w", pady=2)

        # Điền sẵn từ lần trước
        saved_user = self.settings.get("username", "")
        saved_pw = None
        if saved_user:
            self.username.insert(0, saved_user)
            if self.settings.get("remember"):
                saved_pw = load_password(saved_user)
                if saved_pw:
                    self.password.insert(0, saved_pw)
        self.password.focus() if saved_user else self.username.focus()

        self.btn_login = ttk.Button(pad, text="ĐĂNG NHẬP", command=self._do_login)
        self.btn_login.pack(pady=10)
        if not keyring:
            ttk.Label(pad, text="(Muốn nhớ mật khẩu: chạy  pip install keyring)",
                      foreground="gray").pack(anchor="w")
        self.login_log = tk.Text(pad, height=6, bg="#101418", fg="#d0d0d0")
        self.login_log.pack(fill="both", expand=True)

        ttk.Label(pad, text=AUTHOR_MARK, font=("", 8), foreground="#999999").pack(
            anchor="e", pady=(4, 0))

        # Tự đăng nhập nếu có sẵn mật khẩu đã lưu — chạy qua after() để màn hình kịp vẽ ra
        # trước (không "nhảy cóc" thẳng sang màn chính khiến người dùng không kịp hiểu chuyện gì
        # xảy ra), _do_login vẫn dùng chung logic/xử lý lỗi như bấm tay bình thường.
        if auto and saved_pw:
            self._llog("• Đã lưu mật khẩu — tự đăng nhập…")
            self.after(50, self._do_login)

    def _llog(self, msg):
        self.login_log.insert("end", msg + "\n"); self.login_log.see("end"); self.update_idletasks()

    def _do_login(self):
        u, p = self.username.get().strip(), self.password.get()
        if not u or not p:
            messagebox.showwarning("Thiếu", "Nhập tên đăng nhập và mật khẩu."); return
        self.btn_login.config(state="disabled")
        self.login_log.delete("1.0", "end")
        s = make_session()
        try:
            cas_login(s, u, p, self._llog, self._ask_captcha)   # chạy luồng chính (captcha an toàn)
        except PipelineError as e:
            self._llog("✖ " + str(e)); self.btn_login.config(state="normal"); return
        except Exception as e:
            self._llog("✖ Lỗi đăng nhập: " + repr(e)); self.btn_login.config(state="normal"); return
        self.session = s
        self._logged_user = u          # lưu lại trước khi xóa màn login
        # Nhớ đăng nhập
        if self.remember.get():
            self.settings["username"] = u
            self.settings["remember"] = True
            if save_password(u, p):
                pass
            else:
                self._llog("   (Không lưu được mật khẩu — thiếu 'keyring'. Chỉ nhớ tên đăng nhập.)")
        else:
            self.settings["username"] = u   # vẫn nhớ tên cho tiện
            self.settings["remember"] = False
            try:
                if keyring:
                    keyring.delete_password(KR_SERVICE, u)
            except Exception:
                pass
        save_settings(self.settings)
        self._show_main()   # thành công → mở giao diện chính

    def _do_logout(self):
        """Đăng xuất — chỉ bỏ session đang dùng trên máy (KHÔNG gọi gì lên server, chưa xác nhận
        được URL đăng xuất CAS thật qua HAR nên không đoán). Mật khẩu đã lưu (nếu có) vẫn giữ
        nguyên trong Keychain — tự bỏ "Nhớ đăng nhập" ở màn đăng nhập lần sau nếu muốn quên hẳn."""
        if not messagebox.askyesno("Đăng xuất", "Đăng xuất khỏi tài khoản hiện tại?"):
            return
        self.session = None
        self._logged_user = None
        self._show_login(auto=False)   # auto=False: không tự nhảy lại vào ngay tài khoản vừa thoát

    def _open_voffice_web(self):
        """Mở trình duyệt tới màn hình "Quản lý phiếu trình" trên chính hệ thống VOffice (không
        phải tab "Quản lý Phiếu trình" của chương trình này) — cùng URL đã dùng ở
        PreviewWindow._start_post_send_close sau khi Lưu/Trình xong."""
        import webbrowser
        url = BASE + "/Index.do?request_locale=en_US&mainMenu=3&trId=2.2"
        try:
            webbrowser.open(url)
        except Exception as e:
            messagebox.showerror("Không mở được trình duyệt", str(e))

    # ---------- Khung cuộn (dùng chung) ----------
    def _make_scrollable(self, parent):
        return make_scrollable_frame(parent)

    def _compute_nav_colors(self):
        """Tự dò xem giao diện hệ thống đang Sáng hay Tối (macOS Dark Mode…) để chọn màu sidebar
        phù hợp — hardcode màu sáng khiến sidebar thành 1 mảng trắng lạc quẻ giữa ứng dụng đang
        hiển thị tối (xem phản hồi người dùng kèm ảnh chụp màn hình)."""
        style = ttk.Style(self)
        base_bg = style.lookup("TFrame", "background") or self.cget("bg") or "#f0f0f0"
        try:
            r, g, b = self.winfo_rgb(base_bg)
            luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 65535
        except tk.TclError:
            luminance = 1.0
        if luminance < 0.5:   # nền tối
            self._NAV_BG = base_bg
            self._NAV_BG_ACTIVE = "#2d4a63"
            self._NAV_FG = style.lookup("TLabel", "foreground") or "#e0e0e0"
            self._NAV_FG_ACTIVE = "#7ec3ff"
            self._NAV_ACCENT_ACTIVE = "#4da3ff"
        else:                  # nền sáng
            self._NAV_BG = "#f5f5f5"
            self._NAV_BG_ACTIVE = "#e3f2fd"
            self._NAV_FG = "#333333"
            self._NAV_FG_ACTIVE = "#0d47a1"
            self._NAV_ACCENT_ACTIVE = "#1976d2"

        # Màu sọc xen kẽ cho bảng danh sách phiếu trình (xem _build_report_tree) — tính từ
        # đúng màu nền Treeview thật của theme, không hardcode, để không lệch tông sáng/tối.
        tree_bg = style.lookup("Treeview", "background") or base_bg
        try:
            tr, tg, tb = self.winfo_rgb(tree_bg)
            tree_luminance = (0.299 * tr + 0.587 * tg + 0.114 * tb) / 65535
        except tk.TclError:
            tree_luminance = luminance
        self._TREE_ZEBRA = self._shade_color(tree_bg, 0.08, lighten=(tree_luminance < 0.5))

    def _shade_color(self, color, amount, lighten):
        """Làm sáng (lighten=True) hoặc tối (False) 1 màu đi `amount` (0-1) — dùng để tự suy ra
        màu sọc xen kẽ/hover từ đúng màu nền hiện có, thay vì đoán 1 màu cố định dễ lệch tông."""
        r, g, b = (c / 257 for c in self.winfo_rgb(color))   # 0-65535 -> 0-255
        if lighten:
            r += (255 - r) * amount; g += (255 - g) * amount; b += (255 - b) * amount
        else:
            r *= (1 - amount); g *= (1 - amount); b *= (1 - amount)
        return f"#{int(r):02x}{int(g):02x}{int(b):02x}"

    # ---------- Sidebar 4 tab dọc của MÀN 2 (Phiếu trình / Văn bản / Nơi nhận / Luồng) ----------
    # Tách theo nhóm việc thay vì gộp hết field lên 1 màn hình — mỗi lúc chỉ thấy 1 nhóm, đỡ
    # "ngợp" cho người dùng phổ thông. Điều hướng tự do (không phải wizard bắt buộc theo thứ
    # tự), nên mỗi tab có dấu ⚠ riêng nếu còn thiếu field bắt buộc (xem _refresh_readiness).
    def _show_compose_tab(self, name):
        for n, page in self._compose_tab_pages.items():
            w = self._compose_tab_widgets[n]
            active = (n == name)
            bg = self._NAV_BG_ACTIVE if active else self._NAV_BG
            fg = self._NAV_FG_ACTIVE if active else self._NAV_FG
            accent_bg = self._NAV_ACCENT_ACTIVE if active else self._NAV_BG
            w["row"].config(bg=bg)
            w["accent"].config(bg=accent_bg)
            w["label"].config(bg=bg, fg=fg, font=("", 10, "bold" if active else "normal"))
            w["badge"].config(bg=bg)
            if active:
                page.pack(fill="both", expand=True)
            else:
                page.pack_forget()
        self._compose_active_tab = name

    def _set_compose_tab_badge(self, name, has_warning):
        self._compose_tab_widgets[name]["badge"].config(text="⚠" if has_warning else "")

    def _set_report_warning(self, msg):
        self.report_warning_label.config(text="⚠ " + msg)

    def _clear_report_warning(self):
        self.report_warning_label.config(text="")

    def _set_profile_warning(self, msg):
        self.profile_warning_label.config(text="⚠ " + msg)

    def _clear_profile_warning(self):
        self.profile_warning_label.config(text="")

    def _set_flow_warning(self, msg):
        self.flow_warning_label.config(text="⚠ " + msg)

    def _clear_flow_warning(self):
        self.flow_warning_label.config(text="")

    def _build_compose_tab_phieu_trinh(self, parent):
        body = self._make_scrollable(parent)
        g1 = ttk.LabelFrame(body, text="Phiếu trình (không gửi đi)", padding=6)
        g1.pack(fill="x", padx=12, pady=(12, 0))
        f = self._row(g1, "  File phiếu trình:", bold=True)
        ttk.Entry(f, textvariable=self.file_report).pack(side="left", fill="x", expand=True)
        ttk.Button(f, text="Chọn…", command=self._pick_file_report).pack(side="left")
        self.extra_report = FileList(g1, "  + Tài liệu thêm (không gửi):")

        f = self._row(body, "Hồ sơ công việc:")
        self.work_profile = ttk.Combobox(f, values=[], state="readonly")
        self.work_profile.pack(side="left", fill="x", expand=True)
        self.profile_warning_label = ttk.Label(body, text="", foreground="#c62828",
                                                wraplength=420, justify="left")
        self.profile_warning_label.pack(anchor="w", padx=12, pady=(0, 4))

        f = self._row(body, "Nội dung phiếu:")
        # tk.Text nhiều dòng thay vì ttk.Entry 1 dòng — cùng lý do với "Trích yếu" ở DocumentSection.
        self.report_content = tk.Text(f, height=3, wrap="word")
        self.report_content.pack(side="left", fill="x", expand=True)
        self.report_warning_label = ttk.Label(body, text="", foreground="#c62828",
                                               wraplength=420, justify="left")
        self.report_warning_label.pack(anchor="w", padx=12, pady=(0, 4))

    def _build_compose_tab_van_ban(self, parent):
        body = self._make_scrollable(parent)
        self.doc_sections_frame = ttk.Frame(body)
        self.doc_sections_frame.pack(fill="x", padx=12, pady=(12, 0))
        self.doc_sections = []
        self._add_document_section()
        ttk.Button(body, text="+ Thêm văn bản", command=self._add_document_section).pack(
            anchor="w", padx=12, pady=(4, 0))
        ttk.Label(body, text="(File chính của mỗi văn bản = cái cần ký. Bỏ trống hết để thử "
                              "nghiệm sẽ dùng tạm file phiếu trình.)",
                  foreground="gray", wraplength=420).pack(anchor="w", padx=12, pady=(4, 10))

        f = self._row(body, "Độ khẩn:")
        self.priority = ttk.Combobox(f, values=[""] + sorted(ENUMS["priority"].keys()), state="readonly")
        self.priority.set("Khẩn"); self.priority.pack(side="left", fill="x", expand=True)
        f = self._row(body, "Độ mật:")
        self.security = ttk.Combobox(f, values=[""] + sorted(ENUMS["security"].keys()), state="readonly")
        self.security.set("Bình thường"); self.security.pack(side="left", fill="x", expand=True)

    def _build_compose_tab_noi_nhan(self, parent):
        body = self._make_scrollable(parent)
        self.recip = RecipientBox(body, CAY, self.store)
        self.recip.pack(fill="x", padx=12, pady=(12, 0))
        self.auto_stamp_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(body, text="Tự đánh số chữ ký lên Phiếu trình + các Văn bản theo Luồng trình đã chọn",
                        variable=self.auto_stamp_var).pack(anchor="w", padx=12, pady=(10, 0))

    def _build_compose_tab_luong(self, parent):
        body = self._make_scrollable(parent)
        # Danh sách luồng LUÔN lấy động từ web (theo đúng tài khoản đang đăng nhập) — không còn
        # 3 luồng cứng trong code nữa. Trong lúc chờ tải, combobox tạm hiện 1 dòng placeholder.
        f = self._row(body, "Luồng trình:")
        self.flow = ttk.Combobox(f, values=[self._FLOW_LOADING], state="readonly")
        self.flow.set(self._FLOW_LOADING)
        self.flow.pack(side="left", fill="x", expand=True)
        self.flow.bind("<<ComboboxSelected>>", self._on_flow_changed)
        self.flow_warning_label = ttk.Label(body, text="", foreground="#c62828",
                                             wraplength=420, justify="left")
        self.flow_warning_label.pack(anchor="w", padx=12, pady=(0, 4))

        self.flow_panel = FlowSignerPanel(body, self.session, self.log, self.flow_store)
        # chưa pack() — panel tự hiện/ẩn tuỳ luồng đang chọn đã có sẵn đủ người hay chưa

    # ---------- MÀN 2: SOẠN & LƯU NHÁP ----------
    def _show_main(self):
        self._clear()
        # Tự tính theo ĐÚNG màn hình đang chạy — không đoán 1 số cố định cho mọi máy (dễ tràn
        # màn hình nhỏ hoặc quá bé so với màn hình lớn). Có trần hợp lý để không quá khổ trên
        # màn hình rất lớn.
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        w, h = min(int(sw * 0.85), 1400), min(int(sh * 0.85), 950)
        self.geometry(f"{w}x{h}")
        self.minsize(480, 360)

        # Thanh trên cùng, NGOÀI notebook — hiện xuyên suốt mọi tab (khác nhãn "Đã đăng nhập"
        # kiểu cũ, chỉ nằm trong tab "Soạn văn bản" nên khuất khi đang ở "Quản lý Phiếu trình").
        topbar = ttk.Frame(self.container, padding=(12, 6)); topbar.pack(fill="x")
        ttk.Label(topbar, text=f"Đã đăng nhập: {getattr(self, '_logged_user', '')}",
                  foreground="#2e7d32", font=("", 9, "bold")).pack(side="left")
        ttk.Button(topbar, text="Đăng xuất", command=self._do_logout).pack(side="right")
        ttk.Button(topbar, text="Mở VOffice", command=self._open_voffice_web).pack(side="right", padx=(0, 6))

        notebook = ttk.Notebook(self.container)
        notebook.pack(fill="both", expand=True)
        compose_tab = ttk.Frame(notebook)
        manage_tab = ttk.Frame(notebook)
        notebook.add(compose_tab, text="Soạn văn bản")
        notebook.add(manage_tab, text="Quản lý Phiếu trình")
        # Giữ lại để "Sửa" (xem _edit_in_compose) tự chuyển đúng sang tab này, và để Lưu/Trình
        # xong tự nhảy sang "Quản lý Phiếu trình" (xem _open_manage_reports_tab).
        self._notebook, self._compose_tab, self._manage_tab = notebook, compose_tab, manage_tab

        # ---- Sidebar 4 tab dọc (trái) + nội dung tab (phải); CHẠY/trạng thái cố định ở dưới
        # cùng, thấy được dù đang xem tab nào — xem _show_compose_tab()/_build_compose_tab_*().
        self._FLOW_LOADING = "(đang tải danh sách luồng…)"
        self._flow_by_name = {}
        self._profile_by_name = {}

        compose_root = ttk.Frame(compose_tab); compose_root.pack(fill="both", expand=True)
        body = ttk.Frame(compose_root); body.pack(fill="both", expand=True)

        self._compute_nav_colors()

        nav = tk.Frame(body, bg=self._NAV_BG, width=150); nav.pack(side="left", fill="y")
        nav.pack_propagate(False)
        ttk.Separator(body, orient="vertical").pack(side="left", fill="y")
        content_area = ttk.Frame(body); content_area.pack(side="left", fill="both", expand=True)

        COMPOSE_TABS = ("Phiếu trình", "Văn bản", "Nơi nhận", "Luồng")
        self._compose_tab_pages = {}
        self._compose_tab_widgets = {}
        for name in COMPOSE_TABS:
            page = ttk.Frame(content_area)
            self._compose_tab_pages[name] = page

            row = tk.Frame(nav, bg=self._NAV_BG); row.pack(fill="x")
            accent = tk.Frame(row, width=4, bg=self._NAV_BG); accent.pack(side="left", fill="y")
            label = tk.Label(row, text=name, bg=self._NAV_BG, fg=self._NAV_FG, anchor="w",
                              padx=12, pady=10, font=("", 10))
            label.pack(side="left", fill="both", expand=True)
            badge = tk.Label(row, text="", bg=self._NAV_BG, fg="#c62828", font=("", 10, "bold"))
            badge.pack(side="right", padx=(0, 10))
            for w in (row, accent, label, badge):
                w.bind("<Button-1>", lambda _e, n=name: self._show_compose_tab(n))
            self._compose_tab_widgets[name] = {"row": row, "accent": accent, "label": label, "badge": badge}

        self._build_compose_tab_phieu_trinh(self._compose_tab_pages["Phiếu trình"])
        self._build_compose_tab_van_ban(self._compose_tab_pages["Văn bản"])
        self._build_compose_tab_noi_nhan(self._compose_tab_pages["Nơi nhận"])
        self._build_compose_tab_luong(self._compose_tab_pages["Luồng"])
        self._show_compose_tab(COMPOSE_TABS[0])

        # ---- Cố định dưới cùng: nút CHẠY + trạng thái sẵn sàng ----
        footer = ttk.Frame(compose_root); footer.pack(fill="x", padx=12, pady=8)
        self.check_var = tk.BooleanVar(value=False)   # "Chỉ kiểm tra" — ẩn khỏi UI, không dùng nữa
        ttk.Button(footer, text="CHẠY", command=self._run).pack(side="left", pady=4)
        ttk.Button(footer, text="Làm mới form", command=self._confirm_reset_form).pack(
            side="left", padx=(8, 0), pady=4)
        self.readiness_label = ttk.Label(footer, text="", font=("", 9))
        self.readiness_label.pack(side="left", padx=(10, 0))

        # Log kỹ thuật — không còn hiện trong giao diện chính (quá phức tạp với người dùng phổ
        # thông); self.log() vẫn cần 1 nơi để ghi vào nên giữ 1 Text KHÔNG pack() ra màn hình.
        self.logbox = tk.Text(compose_root)

        ttk.Label(compose_root, text=AUTHOR_MARK, font=("", 8), foreground="#999999").pack(
            anchor="e", padx=12, pady=(0, 6))

        self._refresh_readiness()   # tự cập nhật định kỳ — xem _refresh_readiness()

        self._fetch_flow_data()   # sau cùng — logbox đã có sẵn để self.log() gọi từ luồng nền

        self._build_manage_reports_tab(manage_tab)

    def _pick(self, var):
        p = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf"), ("Tất cả", "*.*")])
        if p: var.set(p)

    def _pick_file_report(self):
        """Chọn File phiếu trình — khác _pick() thường ở chỗ: nếu là .docx, tự chuyển
        sang PDF trước (nền); rồi đọc "Nội dung phiếu" ngay từ chính phiếu trình (xem
        extract_phieu_trinh_content), GHI ĐÈ ô đó (nội dung phiếu trình luôn đáng tin hơn trích
        yếu văn bản — trích yếu chỉ điền khi ô còn trống, xem DocumentSection._apply_extract)."""
        p = filedialog.askopenfilename(
            filetypes=[("PDF/Word", "*.pdf *.docx"), ("Tất cả", "*.*")])
        if not p:
            return
        self.file_report.set(p)
        threading.Thread(target=self._prepare_report_worker, args=(p,), daemon=True).start()

    def _prepare_report_worker(self, path):
        """Đọc nhanh để tự điền "Nội dung phiếu" ngay khi vừa chọn file — .docx qua python-docx
        (không cần Word), .pdf qua pymupdf như cũ. Việc chuyển .docx sang PDF thật dời sang lúc
        bấm CHẠY (xem _run()/_run_after_conversion) — xem lý do ở DocumentSection._prepare_worker."""
        if path.lower().endswith(".docx"):
            self._extract_report_docx_worker(path)
            return
        if not path.lower().endswith(".pdf"):
            return   # định dạng khác PDF/Word — không tự đọc được, để nguyên cho người dùng tự điền
        self._extract_report_worker(path)

    def _extract_report_docx_worker(self, path):
        """Đọc thẳng .docx bằng python-docx, không cần Word."""
        content, err = None, None
        try:
            content = extract_phieu_trinh_content_docx(path)
        except Exception as e:
            err = str(e)
        if err:
            msg = (f"Không đọc nhanh được .docx để tự điền nội dung phiếu: {err} — vẫn có thể "
                   "tự điền lại sau khi bấm CHẠY (lúc đó có bản PDF).")
            self.after(0, lambda: self._set_report_warning(msg))
            return
        self.after(0, self._apply_report_extract, path, content, None)

    def _extract_report_worker(self, path):
        content, err = None, None
        try:
            content = extract_phieu_trinh_content(path)
        except Exception as e:
            err = str(e)
        self.after(0, self._apply_report_extract, path, content, err)

    def _apply_report_extract(self, path, content, err):
        if path != self.file_report.get():
            return   # đã chọn file phiếu trình khác trong lúc đọc — bỏ kết quả cũ
        if err:
            self._set_report_warning(f"Không đọc được nội dung phiếu trình để tự điền: {err}")
            return
        if content:
            self.report_content.delete("1.0", "end")
            self.report_content.insert("1.0", content)
            self._clear_report_warning()
        else:
            self._set_report_warning("Không nhận ra được nội dung trong file phiếu trình (mẫu "
                                      "khác/PDF quét ảnh không có lớp chữ) — điền tay nếu cần.")

    FLOW_SEPARATOR = "──────── Luồng khác (từ web) ────────"
    FLOW_QUEN_MAX = 5   # số luồng "quen" tối đa hiện ở đầu combobox (ghim + hay dùng nhất)

    def _fetch_flow_data(self):
        """Lấy (1 lần, cùng 1 request prepareInsert.do): toàn bộ luồng của tài khoản đang đăng
        nhập, VÀ toàn bộ hồ sơ công việc — cả 2 đều đúng theo tài khoản đang đăng nhập, không
        phải giá trị tĩnh ghi cứng trong luong_trinh.json/du_lieu.json."""
        def worker():
            try:
                flows, profiles = fetch_prepare_insert_data(self.session, self.log)
            except Exception as e:
                msg = f"Không lấy được danh sách luồng/hồ sơ công việc từ web: {e}"
                self.after(0, lambda: self._set_flow_warning(msg))
                self.after(0, lambda: self._set_profile_warning(msg))
                return
            self.after(0, lambda: self._apply_flow_list(flows))
            self.after(0, lambda: self._apply_work_profiles(profiles))
        threading.Thread(target=worker, daemon=True).start()

    def _apply_flow_list(self, live_flows):
        """Xếp lên đầu combobox tối đa FLOW_QUEN_MAX luồng: ưu tiên đã ghim, rồi theo tần suất
        dùng TRÊN MÁY NÀY (luong_trinh.json — sổ cục bộ, không dùng chung cho mọi người) — phần
        còn lại xếp dưới 1 dòng phân cách, giữ nguyên thứ tự web trả về."""
        by_id = {fl["flowId"]: fl["name"] for fl in live_flows}
        pinned = set(self.flow_store.get("pinned", []))
        freq = self.flow_store.get("freq", {})
        known_ids = [fid for fid in by_id if fid in pinned or freq.get(fid, 0) > 0]
        known_ids.sort(key=lambda fid: (0 if fid in pinned else 1, -freq.get(fid, 0)))
        quen_ids = known_ids[:self.FLOW_QUEN_MAX]

        self._flow_by_name = {by_id[fid]: fid for fid in quen_ids}
        rest = [fl for fl in live_flows if fl["flowId"] not in quen_ids]
        values = list(self._flow_by_name.keys())
        if rest:
            values.append(self.FLOW_SEPARATOR)
            for fl in rest:
                self._flow_by_name.setdefault(fl["name"], fl["flowId"])
            values += [fl["name"] for fl in rest]

        if not values:
            self._set_flow_warning("Không lấy được luồng nào từ web (giữ nguyên ô trống, tự kiểm tra lại mạng).")
            return
        self._clear_flow_warning()
        self.flow.config(values=values)
        self._select_default_flow()   # tải khung chọn người ngay cho lựa chọn mặc định

    def _select_default_flow(self):
        """Chọn lại luồng đầu danh sách hiện có trong combobox (không fetch lại mạng) — dùng
        khi tải xong lần đầu VÀ khi reset form về trạng thái ban đầu (xem _reset_form)."""
        values = list(self.flow.cget("values"))
        if values:
            self.flow.set(values[0])
        self._on_flow_changed()

    def _apply_work_profiles(self, profiles):
        if not profiles:
            self._set_profile_warning("Không tìm thấy hồ sơ công việc nào của tài khoản này trên web.")
            return
        self._clear_profile_warning()
        self._profile_by_name = {p["name"]: p["fileId"] for p in profiles}
        self.work_profile.config(values=list(self._profile_by_name.keys()))
        self._select_default_work_profile()

    def _select_default_work_profile(self):
        """Chọn lại hồ sơ công việc mặc định trong danh sách hiện có (không fetch lại mạng):
        ưu tiên hồ sơ có chữ "chung" (dùng chung, không gắn 1 vụ việc cụ thể); nếu không có thì
        lấy hồ sơ đầu tiên trong danh sách của tài khoản này."""
        if not self._profile_by_name:
            self.work_profile.set("")
            return
        default_name = next((n for n in self._profile_by_name if "chung" in n.lower()),
                             next(iter(self._profile_by_name)))
        self.work_profile.set(default_name)

    def _on_flow_changed(self, _event=None):
        name = self.flow.get()
        if name == self.FLOW_SEPARATOR:
            # dòng phân cách không phải luồng thật — quay lại luồng đầu danh sách
            self.flow.set(next(iter(self._flow_by_name), self._FLOW_LOADING))
            name = self.flow.get()
        flow_id = self._flow_by_name.get(name)
        if not flow_id:
            return   # danh sách luồng chưa tải xong (vẫn đang hiện placeholder) — chưa làm gì được
        self.flow_panel.load(flow_id, self.report_content.get("1.0", "end-1c"))

    def _flow_name_for_id(self, flow_id):
        for name, fid in self._flow_by_name.items():
            if fid == flow_id:
                return name
        return None

    def _add_document_section(self):
        ds = DocumentSection(self.doc_sections_frame, self, self._remove_document_section)
        ds.pack(fill="x", pady=(0, 6))
        self.doc_sections.append(ds)
        self._update_doc_section_remove_buttons()

    def _remove_document_section(self, ds):
        if len(self.doc_sections) <= 1:
            return   # luôn giữ ít nhất 1 văn bản
        self.doc_sections.remove(ds)
        ds.destroy()
        self._update_doc_section_remove_buttons()

    def _update_doc_section_remove_buttons(self):
        only_one = len(self.doc_sections) <= 1
        for ds in self.doc_sections:
            ds.btn_remove.config(state=("disabled" if only_one else "normal"))

    def _confirm_reset_form(self):
        if messagebox.askyesno(
                "Làm mới form",
                "Xóa hết dữ liệu đang nhập (file, văn bản, nơi nhận, nội dung phiếu...) "
                "để bắt đầu phiếu mới?"):
            self._reset_form()

    def _reset_form(self):
        """Đưa form chính về đúng trạng thái ban đầu như lúc _show_main() vừa chạy — gọi tự
        động sau khi 1 phiếu đã Lưu/Trình THÀNH CÔNG (xem PreviewWindow._on_close), hoặc do
        người dùng tự bấm "Làm mới form". Không đụng self.session/self.settings/self.store/
        self.flow_store (trạng thái đăng nhập/sổ dùng chung, không phải dữ liệu theo phiếu),
        và không fetch lại danh sách luồng/hồ sơ từ mạng — chỉ chọn lại giá trị mặc định trong
        danh sách đã có sẵn trong bộ nhớ."""
        self._editing_report_id = None   # xoá dấu "đang sửa phiếu X" (xem _edit_in_compose)
        self._editing_report_existing_attach_ids = None
        self._cleanup_edit_tmpdirs()
        self.file_report.set("")
        self.extra_report.clear()

        for ds in list(self.doc_sections):
            ds.destroy()
        self.doc_sections = []
        self._add_document_section()

        self.report_content.delete("1.0", "end")
        self.recip.clear()

        self.priority.set("Khẩn")
        self.security.set("Bình thường")
        self.auto_stamp_var.set(True)
        self.check_var.set(False)

        self._select_default_flow()
        self._select_default_work_profile()

        self._clear_report_warning()
        self._clear_profile_warning()
        self._clear_flow_warning()
        self.logbox.delete("1.0", "end")
        self.log("— Đã tự làm mới form, sẵn sàng cho phiếu trình mới —")
        self._refresh_readiness()

    def _cleanup_edit_tmpdirs(self):
        """Xoá các thư mục tạm chứa file tải về lúc "Sửa" (xem _edit_in_compose) — gọi khi làm
        mới form (dữ liệu cũ không còn cần) và khi thoát chương trình (atexit), tránh để lại rác
        chứa văn bản đã tải về vô thời hạn trong thư mục tạm hệ điều hành."""
        for d in self._edit_tmpdirs:
            shutil.rmtree(d, ignore_errors=True)
        self._edit_tmpdirs = []

    def log(self, msg):
        self.logbox.insert("end", msg + "\n"); self.logbox.see("end"); self.update_idletasks()
        _file_logger.info(msg)

    def _collect_cfg(self):
        return {
            "priority": self.priority.get(),
            "security": self.security.get(),
            "report_content": self.report_content.get("1.0", "end-1c"),
            "flow_id": self._flow_by_name.get(self.flow.get()),
            "flow_name": self.flow.get(),
            "flow_nodes_override": self.flow_panel.get_nodes() or None,
            "work_profile_id": self._profile_by_name.get(self.work_profile.get()),
            "work_profile_name": self.work_profile.get(),
            "auto_stamp": self.auto_stamp_var.get(),
            "recv_inside": self.recip.get("inside"),
            "recv_report": self.recip.get("report"),
            "recv_edoc": self.recip.get("edoc"),
            "recv_save": self.recip.get("save"),
            "recv_know": self.recip.get("know"),
            "file_report_main": self.file_report.get(),
            "files_report_extra": self.extra_report.get(),
            "documents": [ds.get() for ds in self.doc_sections],
            "report_id": getattr(self, "_editing_report_id", None),
            "report_existing_attach_ids": getattr(self, "_editing_report_existing_attach_ids", None),
        }

    def _refresh_readiness(self):
        """1 dòng trạng thái cạnh nút CHẠY — liếc là biết đã đủ để bấm chưa, khỏi phải tự rà
        từng ô. Đồng thời gắn dấu ⚠ lên đúng tab (sidebar) còn thiếu, để biết cần quay lại tab
        nào mà không phải tự dò qua cả 4 tab. Tự cập nhật định kỳ (không cần nối callback riêng
        vào từng ô/luồng/nơi nhận — đơn giản hơn, chi phí không đáng kể)."""
        missing_by_tab = {"Phiếu trình": [], "Văn bản": [], "Nơi nhận": [], "Luồng": []}
        if not self.file_report.get():
            missing_by_tab["Phiếu trình"].append("File phiếu trình")
        if not self.report_content.get("1.0", "end-1c").strip():
            missing_by_tab["Phiếu trình"].append("Nội dung phiếu")
        multi = len(self.doc_sections) > 1
        for i, ds in enumerate(self.doc_sections):
            tag = f" (văn bản {i+1})" if multi else ""
            if not ds.abstract.get("1.0", "end-1c").strip():
                missing_by_tab["Văn bản"].append(f"Trích yếu văn bản{tag}")
            if not ds.code.get().strip():
                missing_by_tab["Văn bản"].append(f"Số/ký hiệu{tag}")
        if not any(self.recip.get(c) for c, _ in RecipientBox.CATS):
            missing_by_tab["Nơi nhận"].append("Nơi nhận")
        if not self.flow_panel.is_ready():
            missing_by_tab["Luồng"].append("chọn người ký cho luồng")

        for name, lst in missing_by_tab.items():
            self._set_compose_tab_badge(name, bool(lst))

        missing = [m for lst in missing_by_tab.values() for m in lst]
        if missing:
            self.readiness_label.config(text="⚠ Còn thiếu: " + ", ".join(missing), foreground="#c62828")
        else:
            self.readiness_label.config(text="✓ Sẵn sàng", foreground="#2e7d32")
        self._readiness_after_id = self.after(700, self._refresh_readiness)

    def _maybe_learn_flow_rule(self, flow_id, cfg):
        """Nếu luồng đang chọn KHÔNG do quy tắc nào tự suy ra được (từ khoá lạ, chưa có trong
        sổ), hỏi có muốn nhớ quy tắc mới không — để lần sau ký hiệu tương tự tự chọn đúng luồng
        này, người dùng mới không phải tự sửa tay luong_trinh.json."""
        code = next((d.get("code") for d in (cfg.get("documents") or []) if d.get("code")), None)
        if not code or flow_id_for_code(code, self.flow_store):
            return   # không có ký hiệu để học, hoặc đã có quy tắc khớp sẵn rồi
        kw = flow_keyword_from_code(code)
        if not kw:
            return
        flow_name = cfg.get("flow_name") or flow_id
        if messagebox.askyesno("Nhớ quy tắc chọn luồng?",
                f"Ký hiệu văn bản có '{kw}' — từ nay tự động chọn luồng '{flow_name}' khi gặp "
                f"ký hiệu có '{kw}' không?"):
            self.flow_store.setdefault("rules", []).insert(0, {"keyword": kw, "flowId": flow_id})
            save_flow_store(self.flow_store)
            self.log(f"• Đã nhớ quy tắc: ký hiệu có '{kw}' → luồng '{flow_name}'.")

    def _run(self):
        """Bấm CHẠY: nếu còn file .docx (phiếu trình hoặc bất kỳ văn bản nào) chưa có bản PDF,
        chuyển hết TUẦN TỰ (1 lượt gọi Word duy nhất cho cả phiếu trình, tránh nhiều luồng nền
        tranh nhau 1 tiến trình Word — xem DocumentSection._prepare_worker) rồi mới tiếp tục —
        áp dụng NHƯ NHAU cho cả 'Chỉ kiểm tra' lẫn chạy thật, vì cả 2 đều cần PDF thật (kể cả
        Chỉ kiểm tra cũng đánh số chữ ký lên PDF, xem bước tương ứng trong run_pipeline())."""
        if not self.flow_panel.is_ready():
            messagebox.showwarning("Chưa xong",
                                    "Luồng trình này còn bước chưa xác định được người ký "
                                    "(đang tải hoặc chưa chọn đủ, hoặc danh sách luồng chưa tải xong). "
                                    "Đợi hoặc chọn xong rồi bấm CHẠY lại.")
            return
        if not self.file_report.get():
            messagebox.showwarning("Thiếu", "Chưa chọn file phiếu trình."); return

        pending = []   # [(kind, idx_hoặc_None, path)] — mọi file CHÍNH còn là .docx
        if self.file_report.get().lower().endswith(".docx"):
            pending.append(("report", None, self.file_report.get()))
        for i, ds in enumerate(self.doc_sections):
            p = ds.file_draft.get()
            if p.lower().endswith(".docx"):
                pending.append(("draft", i, p))

        if not pending:
            self._run_after_conversion()
            return

        total = len(pending)
        dlg = _ConvertingDialog(self, f"Đang chuyển sang PDF (1/{total})…")
        def worker():
            for k, (kind, idx, path) in enumerate(pending, start=1):
                self.after(0, lambda p=path, k=k: dlg.set_status(
                    f"Đang chuyển sang PDF ({k}/{total}): {os.path.basename(p)}…"))
                try:
                    pdf_path = convert_office_doc_to_pdf(path)
                except Exception as e:
                    err = str(e)   # tính ngay trong khối except — "e" bị Python xoá khi except kết thúc
                    self.after(0, lambda path=path, err=err: self._conversion_run_failed(dlg, path, err))
                    return
                self.after(0, lambda kind=kind, idx=idx, pdf_path=pdf_path:
                           self._apply_run_conversion(kind, idx, pdf_path))
            self.after(0, lambda: (dlg.close(), self._run_after_conversion()))
        threading.Thread(target=worker, daemon=True).start()

    def _conversion_run_failed(self, dlg, path, err):
        dlg.close()
        messagebox.showerror("Không chuyển được sang PDF",
            f"File: {os.path.basename(path)}\n\nLỗi: {err}\n\n"
            "Chưa thể tiếp tục (kể cả 'Chỉ kiểm tra') khi file này chưa có bản PDF. "
            "Bấm CHẠY thử lại, hoặc tự mở Word 'Save As' sang PDF rồi chọn lại file đó.")

    def _apply_run_conversion(self, kind, idx, pdf_path):
        """Cập nhật đường dẫn sau khi convert xong 1 file, và tự điền NỐT các ô còn trống —
        đọc từ chính bản PDF vừa có (đáng tin hơn/đầy đủ hơn bản đọc nhanh lúc chọn file, vì
        giờ đã chắc chắn có PDF). CHỈ điền ô đang trống, không đè lên bất kỳ ô nào người dùng
        (hoặc bước đọc nhanh lúc chọn file) đã điền trước đó."""
        if kind == "report":
            self.file_report.set(pdf_path)
            try:
                content = extract_phieu_trinh_content(pdf_path)
            except Exception as e:
                self._set_report_warning(f"Không đọc được nội dung phiếu trình từ PDF vừa chuyển: {e}")
                return
            if content and not self.report_content.get("1.0", "end-1c").strip():
                self.report_content.delete("1.0", "end"); self.report_content.insert("1.0", content)
                self._clear_report_warning()
            return

        ds = self.doc_sections[idx]
        ds.file_draft.set(pdf_path)
        try:
            doc_type, code, abstract = extract_draft_fields(pdf_path)
        except Exception as e:
            ds._set_warning(f"Không đọc được văn bản từ PDF vừa chuyển: {e}")
            return
        if doc_type and not ds.doc_type.get():
            ds.doc_type.set(doc_type)
        if code and not ds.code.get().strip():
            ds.code.delete(0, "end"); ds.code.insert(0, code)
            fid = flow_id_for_doc(doc_type, code, self.flow_store)
            fname = self._flow_name_for_id(fid) if fid else None
            if fname:
                self.flow.set(fname)
                self._on_flow_changed()
        if abstract and not ds.abstract.get("1.0", "end-1c").strip():
            ds.abstract.delete("1.0", "end"); ds.abstract.insert("1.0", abstract)
            if not self.report_content.get("1.0", "end-1c").strip():
                self.report_content.delete("1.0", "end"); self.report_content.insert("1.0", abstract)
        ds._clear_warning()

    def _run_after_conversion(self):
        """Phần logic CHẠY thật sự — chạy sau khi mọi file .docx (nếu có) đã chuyển xong sang
        PDF, nên đọc lại cfg ở đây (không dùng cfg cũ) để lấy đúng đường dẫn .pdf mới."""
        flow_id = self._flow_by_name.get(self.flow.get())
        cfg = self._collect_cfg()

        if flow_id:
            bump_flow_freq(self.flow_store, flow_id)   # để lần sau luồng này tự nổi lên đầu combobox
            for n in self.flow_panel.get_nodes():
                if len(n.get("candidates") or []) >= 2:
                    remember_signer_pick(self.flow_store, flow_id, n["nodeId"], n.get("userId"))
            if not self.check_var.get():   # chỉ hỏi ở lần chạy thật, đỡ phiền lúc test "chỉ kiểm tra"
                self._maybe_learn_flow_rule(flow_id, cfg)

        if self.check_var.get():
            # Chế độ an toàn (lần đầu): chạy thẳng như trước, KHÔNG mở khung xem trước, không ghi gì.
            self.logbox.delete("1.0", "end")
            self.log("=== BẮT ĐẦU XỬ LÝ (CHỈ KIỂM TRA) ===")
            s = self.session
            def worker():
                try:
                    run_pipeline(s, cfg, self.log, check_only=True)
                except PipelineError as e:
                    self.log("\n✖ DỪNG: " + str(e))
                except Exception as e:
                    self.log("\n✖ LỖI KHÔNG NGỜ: " + repr(e))
            threading.Thread(target=worker, daemon=True).start()
            return

        if fitz is None:
            messagebox.showerror("Thiếu thư viện",
                                  "Khung xem trước cần thư viện 'pymupdf'. Chạy: pip install pymupdf")
            return
        PreviewWindow(self, cfg, self.session)

    # ---------- TAB "QUẢN LÝ PHIẾU TRÌNH" ----------
    MGMT_TABS = (
        ("processing", "Đang xử lý"),
        ("draft", "Nháp"),
        ("processed", "Hoàn thành"),
    )
    MGMT_COLUMNS = (
        ("date", "Ngày tạo", 120),
        ("content", "Nội dung", 370),
        ("holder", "Người đang giữ", 140),
        ("finish", "Ngày hoàn thành", 120),
    )

    def _build_manage_reports_tab(self, parent):
        top = ttk.Frame(parent); top.pack(fill="x", padx=8, pady=(8, 4))
        now = datetime.now()
        self.mgmt_date_from = tk.StringVar(value=(now - timedelta(days=30)).strftime("%Y-%m-%d"))
        self.mgmt_date_to = tk.StringVar(value=now.strftime("%Y-%m-%d"))
        ttk.Label(top, text="Từ ngày:").pack(side="left")
        ttk.Entry(top, textvariable=self.mgmt_date_from, width=12).pack(side="left", padx=(2, 8))
        ttk.Label(top, text="Đến ngày:").pack(side="left")
        ttk.Entry(top, textvariable=self.mgmt_date_to, width=12).pack(side="left", padx=(2, 8))
        ttk.Label(top, text="(định dạng NĂM-THÁNG-NGÀY, vd 2026-08-01 — mặc định 1 tháng gần nhất)",
                  foreground="gray").pack(side="left")
        ttk.Button(top, text="↻ Làm mới danh sách", command=self._reload_report_lists).pack(side="left", padx=(8, 0))

        # Tìm theo từ khoá trong "Nội dung" — lọc NGAY trên danh sách đã tải (không gọi mạng
        # lại), áp dụng cho cả 3 tab cùng lúc — giúp bớt rối khi danh sách dài mà không phải
        # đổi cách hiển thị từng dòng.
        self.mgmt_search = tk.StringVar()
        self._mgmt_search_job = None
        ttk.Label(top, text="Tìm nội dung:").pack(side="left", padx=(16, 0))
        search_entry = ttk.Entry(top, textvariable=self.mgmt_search, width=24)
        search_entry.pack(side="left", padx=(2, 0))
        search_entry.bind("<KeyRelease>", self._on_mgmt_search_key)

        # Sidebar dọc cho 3 trạng thái (Đang xử lý/Nháp/Hoàn thành) — cùng kiểu nav dọc với
        # "Soạn văn bản" (xem _show_compose_tab), dùng chung màu self._NAV_* đã set ở đó.
        body = ttk.Frame(parent); body.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        nav = tk.Frame(body, bg=self._NAV_BG, width=150); nav.pack(side="left", fill="y")
        nav.pack_propagate(False)
        ttk.Separator(body, orient="vertical").pack(side="left", fill="y")
        content_area = ttk.Frame(body); content_area.pack(side="left", fill="both", expand=True)

        self._mgmt_tab_frames = {}   # which -> Frame nội dung, để _show_mgmt_tab() hiện đúng trang
        self._mgmt_tab_widgets = {}  # which -> {row, accent, label}, để tô màu tab đang chọn
        self.mgmt_trees = {}
        self.mgmt_items = {}
        for which, label in self.MGMT_TABS:
            page = ttk.Frame(content_area)
            self._mgmt_tab_frames[which] = page

            row = tk.Frame(nav, bg=self._NAV_BG); row.pack(fill="x")
            accent = tk.Frame(row, width=4, bg=self._NAV_BG); accent.pack(side="left", fill="y")
            lbl = tk.Label(row, text=label, bg=self._NAV_BG, fg=self._NAV_FG, anchor="w",
                           padx=12, pady=10, font=("", 10))
            lbl.pack(side="left", fill="both", expand=True)
            for w in (row, accent, lbl):
                w.bind("<Button-1>", lambda _e, w=which: self._show_mgmt_tab(w))
            self._mgmt_tab_widgets[which] = {"row": row, "accent": accent, "label": lbl}

            self.mgmt_trees[which] = self._build_report_tree(page, which)
            self.mgmt_items[which] = {}   # iid -> item dict đầy đủ
        self._show_mgmt_tab(self.MGMT_TABS[0][0])
        self._reload_report_lists()   # tự làm mới 1 lần ngay khi mở màn hình chính (mỗi lần đăng nhập)

        ttk.Label(parent, text=AUTHOR_MARK, font=("", 8), foreground="#999999").pack(
            anchor="e", padx=12, pady=(0, 6))

    def _show_mgmt_tab(self, which):
        for w, page in self._mgmt_tab_frames.items():
            wd = self._mgmt_tab_widgets[w]
            active = (w == which)
            bg = self._NAV_BG_ACTIVE if active else self._NAV_BG
            fg = self._NAV_FG_ACTIVE if active else self._NAV_FG
            accent_bg = self._NAV_ACCENT_ACTIVE if active else self._NAV_BG
            wd["row"].config(bg=bg)
            wd["accent"].config(bg=accent_bg)
            wd["label"].config(bg=bg, fg=fg, font=("", 10, "bold" if active else "normal"))
            if active:
                page.pack(fill="both", expand=True)
            else:
                page.pack_forget()
        self._mgmt_active_tab = which

    def _open_manage_reports_tab(self, which):
        """Nhảy sang tab "Quản lý Phiếu trình" + đúng sub-tab `which` ("processing"/"draft"),
        rồi làm mới danh sách — gọi sau khi Lưu/Trình xong (xem PreviewWindow._finish_post_send)."""
        self._notebook.select(self._manage_tab)
        self._show_mgmt_tab(which)
        self._reload_report_lists()

    MGMT_CONTENT_WRAP = 55   # số ký tự/dòng khi bọc cột "Nội dung" — xem _render_mgmt_tree

    def _build_report_tree(self, parent, which):
        wrap = ttk.Frame(parent); wrap.pack(fill="both", expand=True)
        cols = [c[0] for c in self.MGMT_COLUMNS]
        # Style riêng "Reports.Treeview" (không đụng style "Treeview" mặc định — PreviewWindow
        # cũng có 1 Treeview khác, dùng cho cây file, không nên bị tăng rowheight lây) — 2 dòng
        # cho cột "Nội dung" (xem _render_mgmt_tree) cần hàng cao hơn mặc định 1 dòng.
        style = ttk.Style(self)
        line_h = tkfont.nametofont("TkDefaultFont").metrics("linespace")
        style.configure("Reports.Treeview", rowheight=line_h * 2 + 12)
        tree = ttk.Treeview(wrap, columns=cols, show="headings", height=10, style="Reports.Treeview")
        for key, title, width in self.MGMT_COLUMNS:
            tree.heading(key, text=title)
            tree.column(key, width=width, anchor="w")
        # Sọc xen kẽ (xem _compute_nav_colors) — dòng dài na ná nhau (nhiều phiếu cùng loại nội
        # dung) dễ đọc nhầm sang dòng bên cạnh nếu không có gì phân biệt theo hàng.
        tree.tag_configure("odd", background=self._TREE_ZEBRA)
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=vsb.set)
        tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        tree.bind("<Double-1>", lambda e, w=which: self._open_report_detail(w))
        return tree

    def _on_mgmt_search_key(self, _event=None):
        if self._mgmt_search_job is not None:
            self.after_cancel(self._mgmt_search_job)
        self._mgmt_search_job = self.after(200, self._debounced_mgmt_filter)

    def _debounced_mgmt_filter(self):
        self._mgmt_search_job = None
        for which, _label in self.MGMT_TABS:
            self._render_mgmt_tree(which)

    def _render_mgmt_tree(self, which):
        """Vẽ lại bảng của 1 tab trạng thái từ self.mgmt_items[which] (đã tải sẵn, không gọi
        mạng lại) — áp từ khoá tìm (nếu có) + bọc "Nội dung" tối đa 2 dòng thay vì cắt "…" mất
        chữ + tô sọc xen kẽ theo đúng thứ tự dòng đang hiện (không theo thứ tự tải gốc, để sọc
        luôn đều dù đã lọc bớt)."""
        tree = self.mgmt_trees[which]
        tree.delete(*tree.get_children())
        kw = self.mgmt_search.get().strip().lower()
        for i, (iid, it) in enumerate(self.mgmt_items[which].items()):
            content = (it.get("content") or "").strip()
            if kw and kw not in content.lower():
                continue
            date = (it.get("createdDate") or "").replace("T", " ")
            holder = it.get("receiveUser") or ""
            finish = (it.get("finishDate") or "").replace("T", " ")
            lines = textwrap.wrap(content, width=self.MGMT_CONTENT_WRAP) or [""]
            if len(lines) > 2:
                lines = lines[:2]
                lines[-1] = lines[-1].rstrip() + "…"
            wrapped = "\n".join(lines)
            tags = ("odd",) if i % 2 else ()
            tree.insert("", "end", iid=iid, values=(date, wrapped, holder, finish), tags=tags)

    def _apply_report_list(self, which, items):
        self.mgmt_items[which] = {}
        for it in items:
            iid = str(it.get("reportId"))
            self.mgmt_items[which][iid] = it
        self._render_mgmt_tree(which)

    def _reload_report_lists(self):
        date_from, date_to = self.mgmt_date_from.get().strip(), self.mgmt_date_to.get().strip()
        try:
            datetime.strptime(date_from, "%Y-%m-%d")
            datetime.strptime(date_to, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("Sai định dạng ngày",
                                  "Từ ngày/Đến ngày phải theo định dạng NĂM-THÁNG-NGÀY, vd 2026-08-01.")
            return
        # 3 danh sách (Đang xử lý/Nháp/Hoàn thành) tải song song, mỗi cái 1 luồng nền riêng —
        # dùng chung 1 cửa sổ chờ, đóng lại khi cả 3 đã xong (kể cả cái nào lỗi cũng tính là
        # "xong" — không treo cửa sổ chờ mãi chỉ vì 1 danh sách tải hỏng).
        total = len(self.MGMT_TABS)
        dlg = _ConvertingDialog(self, f"Đang tải danh sách phiếu trình (0/{total})…")
        state = {"done": 0}
        for which, _label in self.MGMT_TABS:
            self._load_report_list(which, date_from, date_to, dlg, state, total)

    def _mark_report_list_done(self, dlg, state, total):
        state["done"] += 1
        if state["done"] >= total:
            dlg.close()
        else:
            dlg.set_status(f"Đang tải danh sách phiếu trình ({state['done']}/{total})…")

    def _load_report_list(self, which, date_from, date_to, dlg, state, total):
        s = self.session
        label = dict(self.MGMT_TABS)[which]
        def worker():
            try:
                if which == "processing":
                    items = _search_my_report(s, grid="prepareProcessDocument",
                                               date_from=date_from, date_to=date_to, count=200)
                elif which == "processed":
                    items = _search_processed_report(s, date_from=date_from, date_to=date_to, count=200)
                else:
                    # "Nháp" = mọi thứ không phải "Đang xử lý" (1) / "Hoàn thành" (3) — đã xác
                    # nhận qua 3 file HAR thật: status của "Nháp" không chỉ là 0 (chưa từng
                    # trình), phiếu vừa THU HỒI mang status=4 (và có cả status=2 quan sát được)
                    # — lọc đúng status==0 làm mất các phiếu này. Danh sách grid=None chưa bao
                    # giờ thấy status 1/3 lẫn vào, nên chỉ cần loại đúng 2 giá trị đó.
                    items = [it for it in _search_my_report(s, grid=None, date_from=date_from,
                                                              date_to=date_to, count=200)
                             if it.get("status") not in (1, 3)]
            except Exception as e:
                self.after(0, lambda: self.log(f"• Không tải được danh sách '{label}': {e!r}"))
                self.after(0, lambda: self._mark_report_list_done(dlg, state, total))
                return
            self.after(0, lambda: self._apply_report_list(which, items))
            self.after(0, lambda: self._mark_report_list_done(dlg, state, total))
        threading.Thread(target=worker, daemon=True).start()

    def _open_report_detail(self, which):
        tree = self.mgmt_trees[which]
        sel = tree.selection()
        if not sel:
            return
        item = self.mgmt_items[which].get(sel[0])
        if not item:
            return
        ReportDetailWindow(self, item, which, self.session)

    def _edit_in_compose(self, item):
        """Nhảy sang tab "Soạn văn bản", tự điền toàn bộ thông tin + file của 1 phiếu Nháp đã
        có (gọi từ nút "Sửa" trong ReportDetailWindow) — cho sửa/đổi file rồi bấm CHẠY như
        bình thường để Lưu/Trình lại (xem run_pipeline/save_document/save_report_draft: đã
        nhận report_id/_existing_pid để cập nhật đúng phiếu/văn bản cũ, không tạo mới)."""
        has_data = bool(self.file_report.get().strip() or self.doc_sections[0].file_draft.get().strip())
        if has_data and not messagebox.askyesno(
                "Sửa phiếu trình",
                "Form \"Soạn văn bản\" đang có dữ liệu chưa lưu — chuyển sang sửa phiếu này sẽ "
                "xoá hết dữ liệu đang nhập. Tiếp tục?"):
            return
        self._reset_form()
        self._notebook.select(self._compose_tab)
        report_id = item.get("reportId")
        self._editing_report_id = report_id
        self.log(f"— Đang tải dữ liệu phiếu #{report_id} để sửa… —")

        s = self.session
        tmpdir = tempfile.mkdtemp(prefix="voffice_edit_")
        self._edit_tmpdirs.append(tmpdir)   # dọn ở _cleanup_edit_tmpdirs (Làm mới form / thoát)

        # Toàn bộ việc tải (file phiếu trình + từng văn bản) có thể mất vài giây tới vài chục
        # giây tuỳ số file — hiện cửa sổ chờ nhỏ (giống lúc chuyển .docx->PDF, xem _run) để
        # người dùng biết đang có việc chạy ngầm, không phải chương trình treo/không phản hồi.
        dlg = _ConvertingDialog(self, f"Đang tải dữ liệu phiếu #{report_id} để sửa…")

        def worker():
            result = {"report_local": None, "extra_locals": [], "docs": [], "report_existing_attach_ids": []}
            # 1. File của chính Phiếu trình (+ tài liệu thêm) — link kèm token đã có sẵn
            #    trong attachPathIcons của item (không cần API mới).
            hrefs = re.findall(r"href='([^']+)'[^>]*>\s*<img[^>]*title='([^']*)'",
                                item.get("attachPathIcons") or "")
            n_report_files = len(hrefs)
            for i, (href, title) in enumerate(hrefs):
                self.after(0, lambda i=i, title=title: dlg.set_status(
                    f"Đang tải file phiếu trình ({i+1}/{n_report_files}): {title}…"))
                url = BASE + "/" + html_unescape(href)
                m = re.search(r"attachId=(\d+)", href)
                if m:
                    # Nhớ lại ID file cũ của chính phiếu trình — xoá (removeFile.do) trước khi
                    # lưu lại, cùng lý do/cơ chế như file của từng văn bản (xem
                    # remove_attach_file/run_pipeline) — nếu không sẽ tích file trùng lặp y hệt.
                    result["report_existing_attach_ids"].append(m.group(1))
                # Thư mục con RIÊNG cho từng file (không đổi tên/thêm tiền tố) — tránh trùng tên
                # nếu 2 file tình cờ cùng tên, vẫn giữ nguyên tên gốc thấy được trên form.
                sub = os.path.join(tmpdir, f"phieu_trinh_{i}")
                os.makedirs(sub, exist_ok=True)
                dest = os.path.join(sub, title)
                try:
                    download_attach(s, url, dest, self.log)
                    strip_view_watermark(dest, self._logged_user, self.log)
                    if i == 0:
                        result["report_local"] = dest
                    else:
                        result["extra_locals"].append(dest)
                except Exception as e:
                    self.log(f"   • Không tải được file phiếu trình '{title}': {e!r}")

            # 2. Chi tiết từng văn bản (Loại VB/Số/Trích yếu/Nơi nhận/Khẩn-mật/Người ký).
            self.after(0, lambda: dlg.set_status("Đang tải danh sách văn bản…"))
            attachs = fetch_report_attachs(s, report_id, self.log)
            doc_items = fetch_document_of_report(s, report_id, self.log)
            n_docs = len(doc_items)
            for di, doc_item in enumerate(doc_items, start=1):
                pid = doc_item.get("publishDocumentId")
                self.after(0, lambda di=di, doc_item=doc_item: dlg.set_status(
                    f"Đang tải văn bản ({di}/{n_docs}): "
                    f"{doc_item.get('code') or doc_item.get('documentType') or '(chưa rõ số/loại)'}…"))
                own_files = [a for a in attachs if a.get("documentId") == pid]
                doc = {
                    "doc_type": doc_item.get("documentType") or "",
                    "code": doc_item.get("code") or "",
                    "abstract": doc_item.get("documentAbstract") or "",
                    "priority": doc_item.get("priority") or "",
                    "security": doc_item.get("securityType") or "",
                    "receive_inside": doc_item.get("receiveInside") or "",
                    "receive_report": doc_item.get("receiveReport") or "",
                    "receive_edoc": doc_item.get("receiveEdoc") or "",
                    "receive_know": doc_item.get("receiveToKnow") or "",
                    "receive_save": doc_item.get("receiveSaveDepartment") or "",
                    "existing_pid": pid,
                    "local_file": None,
                    "extra_locals": [],
                    # ID file CŨ của văn bản này — cần xoá (removeFile.do) khi thật sự Lưu/Trình
                    # lại, nếu không file mới upload lại sẽ CỘNG THÊM vào chứ không thay thế (xem
                    # remove_attach_file / run_pipeline).
                    "existing_attach_ids": [f["draftDocumentId"] for f in own_files if f.get("draftDocumentId")],
                }
                main_file = next((a for a in own_files if a.get("documentAbstract")), None) or \
                    (own_files[0] if own_files else None)
                if own_files:
                    # Tải TOÀN BỘ file của văn bản này (không chỉ file chính) — trước đây chỉ
                    # tải main_file nên các "Tài liệu gửi kèm" (vd 1 văn bản có nhiều file phụ,
                    # xác nhận qua HAR có tới 6-7 file/văn bản) bị bỏ sót, không đưa vào lại
                    # form Sửa.
                    tokens = fetch_draft_attach_tokens(s, pid, self.log)
                    for fi, f in enumerate(own_files):
                        info = tokens.get(f.get("draftDocumentId"))
                        name = f.get("draftDocumentName") or "file.pdf"
                        if not info:
                            self.log(f"   • Không tìm thấy token tải file '{name}' (publishDocumentId={pid}) "
                                     "— tự chọn lại file này trước khi bấm CHẠY.")
                            continue
                        # Thư mục con riêng cho từng file — giữ nguyên tên gốc (xem lý do ở
                        # nhánh tải file phiếu trình phía trên).
                        sub = os.path.join(tmpdir, f"van_ban_{pid}_{fi}")
                        os.makedirs(sub, exist_ok=True)
                        dest = os.path.join(sub, name)
                        url = f"{BASE}/uploadiframe!openFile.do?token={info['token']}&attachId={f['draftDocumentId']}"
                        try:
                            download_attach(s, url, dest, self.log)
                            strip_view_watermark(dest, self._logged_user, self.log)
                            if f is main_file:
                                doc["local_file"] = dest
                            else:
                                doc["extra_locals"].append(dest)
                        except Exception as e:
                            self.log(f"   • Không tải được file '{name}' (publishDocumentId={pid}): {e!r} "
                                     "— tự chọn lại file này trước khi bấm CHẠY.")
                result["docs"].append(doc)
            self.after(0, lambda: (dlg.close(), self._apply_edit_data(item, result)))
        threading.Thread(target=worker, daemon=True).start()

    def _apply_edit_data(self, item, result):
        if result["report_local"]:
            self.file_report.set(result["report_local"])
        for p in result["extra_locals"]:
            self.extra_report.add_path(p)
        self._editing_report_existing_attach_ids = result.get("report_existing_attach_ids") or []

        self.report_content.delete("1.0", "end")
        self.report_content.insert("1.0", item.get("content") or "")

        docs = result["docs"] or [{}]
        for i, doc in enumerate(docs):
            if i >= len(self.doc_sections):
                self._add_document_section()
            ds = self.doc_sections[i]
            if doc.get("doc_type"):
                ds.doc_type.set(doc["doc_type"])
            ds.code.delete(0, "end"); ds.code.insert(0, doc.get("code") or "")
            ds.abstract.delete("1.0", "end"); ds.abstract.insert("1.0", doc.get("abstract") or "")
            if doc.get("local_file"):
                ds.file_draft.set(doc["local_file"])
            for p in doc.get("extra_locals") or []:
                ds.extra.add_path(p)
            ds._existing_pid = doc.get("existing_pid")
            ds._existing_attach_ids = doc.get("existing_attach_ids") or []

        first = docs[0]
        if first.get("priority"):
            self.priority.set(first["priority"])
        if first.get("security"):
            self.security.set(first["security"])
        for name, fid in self._profile_by_name.items():
            if fid == item.get("fileId"):
                self.work_profile.set(name); break

        self._fill_recipients_best_effort(first)

        self.log(f"— Đã điền xong dữ liệu phiếu #{item.get('reportId')} — kiểm tra lại rồi bấm CHẠY. "
                  "Luồng trình/Người ký KHÔNG tự chọn lại — tự chọn lại nếu cần. —")
        self._refresh_readiness()
        # Dữ liệu vừa tự điền nằm rải trên nhiều tab (Văn bản/Nơi nhận) — với giao diện 4 tab
        # mới, nếu không tự nhảy sang, người dùng đang đứng ở tab khác (vd "Phiếu trình") sẽ
        # tưởng nhầm là chưa điền được gì. Nhảy sang "Văn bản" trước (file/nội dung cần soát kỹ
        # nhất trước khi bấm CHẠY) — "Nơi nhận" xem badge cảnh báo trên sidebar để biết cần ghé qua.
        self._show_compose_tab("Văn bản")

    def _fill_recipients_best_effort(self, doc):
        """Điền lại Nơi nhận từ TÊN (không có ID) bằng khớp gần nhất trong cây đơn vị — best
        effort, giống triết lý "tự điền chỉ để đỡ gõ tay" đã ghi trong HUONG_DAN.md. Luôn tự
        kiểm tra lại trước khi bấm CHẠY."""
        cat_field_tree = [
            ("inside", "receive_inside", CAY["internal"]["nodes"]),
            ("report", "receive_report", CAY["internal"]["nodes"]),
            ("know", "receive_know", CAY["internal"]["nodes"]),
            ("save", "receive_save", CAY["internal"]["nodes"]),
            ("edoc", "receive_edoc", CAY["lien_thong"]["nodes"]),
        ]
        for cat, field, nodes in cat_field_tree:
            names = [n.strip() for n in (doc.get(field) or "").split(";") if n.strip()]
            for name in names:
                matches = search_nodes(name, nodes, self.store, k=1)
                if matches:
                    nd = matches[0]
                    if not any(_node_key(x) == _node_key(nd) for x in self.recip.buckets[cat]):
                        self.recip.buckets[cat].append(nd)
                else:
                    self.log(f"   • Không khớp được nơi nhận '{name}' trong cây đơn vị — tự thêm tay.")
        self.recip._render_chips()

    def _ask_captcha(self, path):
        try:
            import webbrowser
            webbrowser.open("file://" + os.path.abspath(path))
        except Exception:
            pass
        self._llog(f"   (Ảnh captcha đã lưu: {path})")
        from tkinter import simpledialog
        return simpledialog.askstring("Captcha", "Nhập mã trong ảnh captcha vừa mở:")


# ---------- MÀN XEM TRƯỚC + GỬI ----------
# ---------- Checklist tiến trình + banner kết quả (cho khung Xem trước & Gửi) ----------
# Thay cho việc chỉ có 1 ô log chạy chữ kỹ thuật — người dùng phổ thông không có cách nào yên
# tâm là "đã xong chưa/có lỗi không" nếu phải tự đọc log. Checklist hiện các mốc lớn (xem
# PIPELINE_PHASES) tự tick khi xong; banner hiện kết quả cuối cùng to, rõ, có màu.
class StepChecklist(ttk.Frame):
    ICONS = {"pending": "○", "running": "⏳", "done": "✓", "error": "✕"}
    COLORS = {"pending": "#9e9e9e", "running": "#f57f17", "done": "#2e7d32", "error": "#c62828"}

    def __init__(self, parent, phases):
        """phases: [(key, nhãn tiếng Việt), ...] — xem PIPELINE_PHASES."""
        super().__init__(parent)
        self.order = [k for k, _ in phases]
        self._rows = {}
        for key, label in phases:
            row = ttk.Frame(self); row.pack(fill="x", anchor="w", pady=1)
            icon = ttk.Label(row, text=self.ICONS["pending"], width=2, foreground=self.COLORS["pending"])
            icon.pack(side="left")
            text = ttk.Label(row, text=label, foreground=self.COLORS["pending"])
            text.pack(side="left")
            self._rows[key] = (icon, text)

    def reset(self):
        for key in self.order:
            self.set_status(key, "pending")

    def set_status(self, key, status):
        if key not in self._rows:
            return
        icon, text = self._rows[key]
        icon.config(text=self.ICONS[status], foreground=self.COLORS[status])
        text.config(foreground=self.COLORS[status])

    def mark_running(self, key):
        """Đánh dấu mọi bước TRƯỚC `key` là xong, và `key` đang chạy."""
        idx = self.order.index(key)
        for k in self.order[:idx]:
            self.set_status(k, "done")
        self.set_status(key, "running")

    def complete_all(self):
        for key in self.order:
            self.set_status(key, "done")

    def mark_error(self, key):
        """Bước `key` bị lỗi — các bước trước đó vẫn coi là đã xong (đúng thực tế: chúng đã
        chạy thành công trước khi mắc ở `key`)."""
        if key not in self._rows:
            return
        idx = self.order.index(key)
        for k in self.order[:idx]:
            self.set_status(k, "done")
        self.set_status(key, "error")


class ReportDetailWindow(tk.Toplevel):
    """Xem chi tiết 1 phiếu trình đã gửi (tab "Quản lý Phiếu trình"): thông tin chung + tiến
    trình ký + lịch sử + danh sách file đính kèm. Giai đoạn 1 — chỉ xem + Thu hồi (nếu đang ở
    "Đang xử lý"); chưa xem PDF, chưa Trình lại (xem lý do ở kế hoạch — thiếu dữ liệu để dựng
    lại đúng code/documentType nếu trình lại)."""

    def __init__(self, master, item, which, session):
        super().__init__(master)
        # transient (không grab_set) — chỉ để Windows gom vào cùng nhóm với cửa sổ cha trên
        # taskbar thay vì hiện thành 1 nút riêng biệt, không đổi gì về tương tác.
        self.transient(master)
        self.item = item
        self.which = which   # "processing" | "draft"
        self.session = session
        self.report_id = item.get("reportId")

        self.title(f"Chi tiết phiếu trình #{self.report_id}")
        # Tự tính theo ĐÚNG màn hình đang chạy (giống _show_main) thay vì số cứng — phiếu nào
        # nhiều bước ký/lịch sử dài trước đây dễ bị đẩy tụt phần File đính kèm ra ngoài khung
        # nhìn thấy được, phải tự kéo to cửa sổ mới thấy (xem phản hồi người dùng kèm ảnh chụp).
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        w, h = min(int(sw * 0.55), 760), min(int(sh * 0.8), 860)
        self.geometry(f"{w}x{h}")
        self.minsize(480, 400)

        top = ttk.Frame(self, padding=(10, 8)); top.pack(fill="x")
        if which == "processing":
            self.btn_cancel = ttk.Button(top, text="Thu hồi", command=self._confirm_cancel)
            self.btn_cancel.pack(side="left")
        if which == "draft":
            ttk.Button(top, text="Sửa", command=self._edit_in_compose).pack(side="left")
        ttk.Button(top, text="Đóng", command=self.destroy).pack(side="left", padx=6)
        ttk.Button(top, text="Tải toàn bộ tài liệu",
                   command=self._download_all_attachs).pack(side="left", padx=(6, 0))

        self.status_var = tk.StringVar(value="")
        self.status_label = ttk.Label(self, textvariable=self.status_var, padding=(10, 4),
                                       wraplength=600, justify="left")
        self.status_label.pack(fill="x")

        # Thứ tự chung cho các cấu phần thuộc Quản lý Phiếu trình: Thông tin → File đính kèm →
        # Tiến trình → Lịch sử (ưu tiên cái cần xem/thao tác trước, "Lịch sử" ít cần nhất xếp
        # cuối cùng).
        body = ttk.Frame(self, padding=8); body.pack(fill="both", expand=True)
        self._build_info(body, item)
        self._build_attach_section(body)
        self._build_list_section(body, "Tiến trình ký", "process_list", height=6)
        self._build_list_section(body, "Lịch sử", "history_list", height=6)

        self._attach_items = []   # song song với self.attach_list, cùng chỉ số — xem _apply_detail
        self._attach_tokens = {}  # draftDocumentId -> {"name","token"} — xem _fetch_attach_tokens
        self._load_detail()

    def _row(self, parent, label, value):
        f = ttk.Frame(parent); f.pack(fill="x", pady=1)
        ttk.Label(f, text=label, width=14, font=("", 9, "bold")).pack(side="left")
        ttk.Label(f, text=value or "", wraplength=480, justify="left").pack(side="left", fill="x", expand=True)

    def _build_info(self, parent, item):
        box = ttk.LabelFrame(parent, text="Thông tin", padding=6)
        box.pack(fill="x", pady=(0, 6))
        self._row(box, "Nội dung:", item.get("content"))
        self._row(box, "Người tạo:", item.get("creator"))
        self._row(box, "Ngày tạo:", (item.get("createdDate") or "").replace("T", " "))
        self._row(box, "Đơn vị:", item.get("officeName"))
        self._row(box, "Trạng thái:", item.get("stateName"))

    def _build_list_section(self, parent, title, attr_name, height):
        box = ttk.LabelFrame(parent, text=title, padding=6)
        box.pack(fill="both", expand=True, pady=(0, 6))
        # Thanh cuộn riêng cho khung này — lưới an toàn cho phiếu trình có quá nhiều bước ký/
        # dòng lịch sử để vừa trần % màn hình: dù cửa sổ có to cỡ nào, khung này tự cuộn thay vì
        # đẩy các khung/nút phía dưới ra khỏi tầm nhìn.
        wrap = ttk.Frame(box); wrap.pack(fill="both", expand=True)
        lb = tk.Listbox(wrap, height=height)
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=lb.yview)
        lb.configure(yscrollcommand=vsb.set)
        lb.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        setattr(self, attr_name, lb)

    def _build_attach_section(self, parent):
        """Giống _build_list_section nhưng thêm bấm-đúp để XEM nhanh (giữ nguyên watermark —
        đây là hành động xem thật, giống mở trên web) + nút TẢI file đang chọn (xoá watermark
        trước khi lưu — xem _download_selected_attach). Nút "Tải toàn bộ tài liệu" nằm ở thanh
        trên cùng (xem __init__), không lặp lại ở đây."""
        box = ttk.LabelFrame(parent, text="File đính kèm (Click đúp để xem)", padding=6)
        box.pack(fill="both", expand=True, pady=(0, 6))
        wrap = ttk.Frame(box); wrap.pack(fill="both", expand=True)
        self.attach_list = tk.Listbox(wrap, height=4)
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=self.attach_list.yview)
        self.attach_list.configure(yscrollcommand=vsb.set)
        self.attach_list.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        self.attach_list.bind("<Double-1>", lambda e: self._open_attach())
        btnrow = ttk.Frame(box); btnrow.pack(fill="x", pady=(4, 0))
        ttk.Button(btnrow, text="Tải file này",
                   command=self._download_selected_attach).pack(side="left")

    def _load_detail(self):
        s, rid = self.session, self.report_id
        def worker():
            process = fetch_report_process(s, rid)
            history = fetch_report_history(s, rid)
            attachs = fetch_report_attachs(s, rid)
            self.after(0, lambda: self._apply_detail(process, history, attachs))
        threading.Thread(target=worker, daemon=True).start()

    def _apply_detail(self, process, history, attachs):
        self.process_list.delete(0, "end")
        for it in process:
            mark = " ⏳ đang chờ" if it.get("status") == 1 else ""
            self.process_list.insert(
                "end", f"{it.get('processOrder')}. {it.get('displayPositionName')} — "
                       f"{it.get('receiveUser')}{mark}")
        self.history_list.delete(0, "end")
        for it in history:
            when = (it.get("createAt") or "").replace("T", " ")
            self.history_list.insert("end", f"{when} — {it.get('fullname')}: {it.get('note')}")
        self.attach_list.delete(0, "end")
        self._attach_items = []
        for it in attachs:
            name = it.get("draftDocumentName")
            if name:
                self.attach_list.insert("end", name)
                self._attach_items.append(it)
        self._fetch_attach_tokens()

    def _fetch_attach_tokens(self):
        """Lấy sẵn token tải cho MỌI file đang liệt kê (nhóm theo publishDocumentId để đỡ gọi
        trùng nếu 1 văn bản có nhiều file) — lấy trước ngay khi mở khung, để bấm "Mở"/"Tải về"
        không phải đợi thêm 1 lượt gọi mạng nữa."""
        pids = {it.get("documentId") for it in self._attach_items if it.get("documentId")}
        if not pids:
            return
        s = self.session
        def worker():
            tokens = {}
            for pid in pids:
                tokens.update(fetch_draft_attach_tokens(s, pid))
            self.after(0, lambda: self._attach_tokens.update(tokens))
        threading.Thread(target=worker, daemon=True).start()

    def _selected_attach(self):
        sel = self.attach_list.curselection()
        if not sel:
            messagebox.showinfo("Chưa chọn file", "Chọn 1 file trong danh sách trước.", parent=self)
            return None
        return self._attach_items[sel[0]]

    def _attach_url(self, it):
        aid = it.get("draftDocumentId")
        info = self._attach_tokens.get(aid)
        if not info:
            return None
        return f"{BASE}/uploadiframe!openFile.do?token={info['token']}&attachId={aid}"

    def _open_attach(self):
        """Bấm đúp — xem nhanh bằng ứng dụng mặc định của máy, GIỮ NGUYÊN watermark (đây là
        hành động xem thật, hệ thống tự đóng dấu y như khi bạn xem trên web — không xoá gì)."""
        it = self._selected_attach()
        if it is None:
            return
        url = self._attach_url(it)
        if not url:
            messagebox.showwarning("Chưa sẵn sàng",
                                    "Đang lấy đường dẫn tải file, thử lại sau vài giây.", parent=self)
            return
        name = it.get("draftDocumentName") or "file"
        dest = os.path.join(_new_gen_tmpdir("voffice_view_"), name)
        dlg = _ConvertingDialog(self, f"Đang tải để xem: {name}…")
        s = self.session
        def worker():
            try:
                download_attach(s, url, dest)
            except Exception as e:
                self.after(0, lambda: (dlg.close(),
                                        messagebox.showerror("Không tải được file", str(e), parent=self)))
                return
            def do_open():
                dlg.close()
                try:
                    open_file_with_default_app(dest)
                except Exception as e:
                    messagebox.showerror("Không mở được file", str(e), parent=self)
            self.after(0, do_open)
        threading.Thread(target=worker, daemon=True).start()

    def _download_selected_attach(self):
        """Nút "Tải file này" — khác _open_attach ở chỗ LƯU VĨNH VIỄN theo đường dẫn tự chọn,
        và XOÁ WATERMARK trước khi lưu (đây là tải về để dùng thật, không phải xem thoáng qua —
        xem strip_view_watermark)."""
        it = self._selected_attach()
        if it is None:
            return
        url = self._attach_url(it)
        if not url:
            messagebox.showwarning("Chưa sẵn sàng",
                                    "Đang lấy đường dẫn tải file, thử lại sau vài giây.", parent=self)
            return
        name = it.get("draftDocumentName") or "file"
        dest = filedialog.asksaveasfilename(parent=self, initialfile=name,
                                             defaultextension=os.path.splitext(name)[1])
        if not dest:
            return
        dlg = _ConvertingDialog(self, f"Đang tải: {name}…")
        s, username = self.session, getattr(self.master, "_logged_user", None)
        def worker():
            try:
                download_attach(s, url, dest)
                strip_view_watermark(dest, username)
            except Exception as e:
                self.after(0, lambda: (dlg.close(),
                                        messagebox.showerror("Không tải được file", str(e), parent=self)))
                return
            self.after(0, lambda: (dlg.close(), messagebox.showinfo(
                "Đã tải về", f"Đã lưu bản sạch (không watermark) tại:\n{dest}", parent=self)))
        threading.Thread(target=worker, daemon=True).start()

    def _download_all_attachs(self):
        """Nút "Tải toàn bộ các file" — tải hết danh sách vào 1 thư mục tự chọn, mỗi file cũng
        được xoá watermark trước khi lưu như _download_selected_attach. Lỗi từng file (nếu có)
        không dừng cả loạt — báo gộp lại cuối cùng (best-effort, giống triết lý các chỗ tải
        hàng loạt khác trong chương trình)."""
        if not self._attach_items:
            messagebox.showinfo("Không có file", "Phiếu trình này chưa có file đính kèm nào.", parent=self)
            return
        folder = filedialog.askdirectory(parent=self, title="Chọn thư mục lưu toàn bộ file")
        if not folder:
            return
        items = list(self._attach_items)
        n = len(items)
        dlg = _ConvertingDialog(self, f"Đang tải file (0/{n})…")
        s, username = self.session, getattr(self.master, "_logged_user", None)
        def worker():
            ok, failed = [], []
            for i, it in enumerate(items, start=1):
                name = it.get("draftDocumentName") or f"file_{i}"
                self.after(0, lambda i=i, name=name: dlg.set_status(f"Đang tải file ({i}/{n}): {name}…"))
                url = self._attach_url(it)
                if not url:
                    failed.append(f"{name} (chưa lấy được đường dẫn tải)")
                    continue
                dest = os.path.join(folder, name)
                try:
                    download_attach(s, url, dest)
                    strip_view_watermark(dest, username)
                    ok.append(name)
                except Exception as e:
                    failed.append(f"{name} ({e})")
            def done():
                dlg.close()
                msg = f"Đã tải {len(ok)}/{n} file (bản sạch, không watermark) vào:\n{folder}"
                if failed:
                    msg += "\n\nKhông tải được:\n" + "\n".join(failed)
                messagebox.showinfo("Tải toàn bộ file", msg, parent=self)
            self.after(0, done)
        threading.Thread(target=worker, daemon=True).start()

    def _confirm_cancel(self):
        if not messagebox.askyesno(
                "Thu hồi phiếu trình",
                f"Thu hồi (hủy trình ký) phiếu trình #{self.report_id}?\n"
                "Đây là hành động ghi lên hệ thống thật — không thể tự hoàn tác từ chương trình.",
                parent=self):
            return
        self.btn_cancel.config(state="disabled")
        self.status_var.set("• Đang thu hồi…")
        s, rid = self.session, self.report_id
        def worker():
            try:
                ok, note = cancel_report(s, rid)
            except Exception as e:
                self.after(0, lambda: self._cancel_done(False, repr(e)))
                return
            msg = (note or {}).get("note") if ok else "chưa xác minh được — tự kiểm tra lại trên web"
            self.after(0, lambda: self._cancel_done(ok, msg))
        threading.Thread(target=worker, daemon=True).start()

    def _cancel_done(self, ok, msg):
        color = "#2e7d32" if ok else "#c62828"
        prefix = "✔" if ok else "⚠"
        self.status_var.set(f"{prefix} {msg}")
        self.status_label.config(foreground=color)
        self.btn_cancel.config(state="normal")
        if ok:
            self.master._reload_report_lists()
            self.destroy()

    def _edit_in_compose(self):
        self.master._edit_in_compose(self.item)
        self.destroy()


class PreviewWindow(tk.Toplevel):
    """Xem lại thông tin + file đã chuẩn bị trước khi gửi thật.
    - Trái: thông tin đã lưu + danh sách file (2 nhóm: Phiếu trình / Văn bản).
    - Phải: xem PDF; với file CHÍNH của mỗi nhóm, các vị trí đánh số chữ ký (tính theo
      Luồng trình) hiện thành các dấu tròn có thể kéo/sửa/xoá/thêm ngay trên trang.
    - File không phải PDF: bấm vào chỉ mở thư mục chứa file (Explorer)."""

    MARK_R = 11
    BANNER_STYLE = {   # kind -> (nền, chữ)
        "success": ("#e8f5e9", "#1b5e20"),
        "ambiguous": ("#fff8e1", "#7a5b00"),
        "error": ("#fdecea", "#b71c1c"),
    }

    def __init__(self, master, cfg, session):
        super().__init__(master)
        # transient (không grab_set) — chỉ để Windows gom vào cùng nhóm với cửa sổ cha trên
        # taskbar thay vì hiện thành 1 nút riêng biệt, không đổi gì về tương tác.
        self.transient(master)
        self.cfg = cfg
        self.session = session
        self._stamps_cache = {}     # path -> [stamp dict]
        self._current_doc = None
        self._current_path = None
        self._current_page = 0
        self._current_kind = None
        self._editable = False
        self._add_mode = False
        self._zoom = 1.0
        self._mark_items = {}       # idx -> (oval_id, text_id)
        self._drag_idx = None
        self._drag_last = None
        self._tree_paths = {}       # iid -> (path, kind)
        self._sending = False       # True trong lúc luồng nền đang Lưu/Trình — chặn đóng cửa
                                     # sổ "lặng lẽ" giữa chừng (xem _on_close)
        self._submitted_ok = False  # True nếu đã Lưu/Trình THÀNH CÔNG ít nhất 1 lần — báo
                                     # cho _on_close biết cần làm mới form chính hay không
        self._current_phase = None  # phase (xem PIPELINE_PHASES) đang chạy — để biết tô đỏ
                                     # đúng bước nào nếu lỗi
        self._log_window = None     # SendLogWindow độc lập, mở ra mỗi lần bấm Lưu/Trình (_send)
        self._post_send_target_tab = None   # "processing"/"draft" — xem _start_post_send_close

        self.title("Xem trước & Gửi")
        self.geometry("1180x760")
        self.minsize(820, 520)
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        # ---- Thanh trên: nút Lưu dự thảo / Trình văn bản ----
        top = ttk.Frame(self, padding=(10, 8)); top.pack(fill="x")
        self.btn_send = ttk.Button(top, text="Lưu dự thảo", command=lambda: self._send("0"))
        self.btn_send.pack(side="left")
        self.btn_submit = ttk.Button(top, text="Trình văn bản", command=lambda: self._send("1"))
        self.btn_submit.pack(side="left", padx=(6, 0))
        ttk.Button(top, text="Đóng", command=self._on_close).pack(side="left", padx=6)
        if not cfg.get("auto_stamp", True):
            ttk.Label(top, foreground="#c62828",
                      text="⚠ Đang TẮT 'Tự đánh số chữ ký' — các dấu bên dưới sẽ KHÔNG được ghi khi Gửi.")\
                .pack(side="left", padx=12)
        self.status_var = tk.StringVar(value="")
        ttk.Label(top, textvariable=self.status_var, foreground="gray").pack(side="right")

        # ---- Khung tiến trình: checklist các mốc lớn + banner kết quả cuối ----
        # Mục đích: người dùng phổ thông bấm Lưu/Trình xong phải THẤY NGAY (không cần đọc log
        # kỹ thuật) đang chạy tới đâu, và cuối cùng có chắc chắn thành công hay không.
        progress = ttk.Frame(self, padding=(10, 4)); progress.pack(fill="x")
        self.checklist = StepChecklist(progress, PIPELINE_PHASES)
        self.checklist.pack(fill="x", anchor="w")
        self.banner_var = tk.StringVar(value="")
        self.banner_label = tk.Label(progress, textvariable=self.banner_var, justify="left",
                                      anchor="w", padx=10, pady=8, wraplength=1000, font=("", 10))
        # banner ẩn cho tới khi có kết quả (xem _set_banner) — không choán chỗ lúc mới mở khung
        btnrow = ttk.Frame(progress); btnrow.pack(fill="x", pady=(4, 0))
        self.btn_toggle_log = ttk.Button(btnrow, text="▾ Xem chi tiết kỹ thuật",
                                          command=self._toggle_log)
        self.btn_toggle_log.pack(side="left")

        body = ttk.Frame(self); body.pack(fill="both", expand=True)
        left = ttk.Frame(body, width=340); left.pack(side="left", fill="y")
        left.pack_propagate(False)
        right = ttk.Frame(body, padding=(0, 8, 8, 8)); right.pack(side="left", fill="both", expand=True)

        # Chia cột trái thành 2 khung có thể kéo giãn: "Thông tin đã lưu" (tự cuộn riêng nếu
        # dài) ở trên, "File" ở dưới — để danh sách file LUÔN có chỗ hiện ra, không bị đẩy
        # khuất khi phần thông tin (trích yếu dài, nhiều nơi nhận...) chiếm hết chỗ.
        vpaned = ttk.PanedWindow(left, orient="vertical")
        vpaned.pack(fill="both", expand=True, padx=8, pady=8)
        info_pane = ttk.Frame(vpaned)
        tree_pane = ttk.Frame(vpaned)
        vpaned.add(info_pane, weight=1)
        vpaned.add(tree_pane, weight=2)

        def _set_initial_sash():
            try:
                total = vpaned.winfo_height()
                if total > 50:
                    vpaned.sashpos(0, min(280, int(total * 0.4)))
            except tk.TclError:
                pass
        self.after(50, _set_initial_sash)

        self._build_info(info_pane)
        self._build_tree(tree_pane, cfg)
        self._build_viewer(right)

        # Log gửi kỹ thuật — ẩn mặc định (xem progress/btn_toggle_log ở trên), người phổ thông
        # không cần thấy; bấm "Xem chi tiết kỹ thuật" mới hiện.
        self.logbox = tk.Text(self, height=6, bg="#101418", fg="#d0d0d0")

        # Quét sẵn file chính của phiếu trình + MỌI văn bản (không chờ người dùng bấm) + chọn
        # file mặc định để hiện (ưu tiên văn bản đầu tiên có file, không thì phiếu trình).
        report_main = cfg.get("file_report_main")
        if report_main and report_main.lower().endswith(".pdf"):
            self._ensure_scanned(report_main)
        draft_mains = [d.get("file_draft_main") for d in (cfg.get("documents") or [])]
        for dm in draft_mains:
            if dm and dm.lower().endswith(".pdf"):
                self._ensure_scanned(dm)

        init_path = next((dm for dm in draft_mains if dm), None) or report_main
        for iid, (p, k) in self._tree_paths.items():
            if p == init_path:
                self.tree.selection_set(iid); self.tree.see(iid)
                break

    # ---------- Trái: thông tin + cây file ----------
    def _build_info(self, parent):
        parent = make_scrollable_frame(parent)   # tự cuộn riêng khi trích yếu/nơi nhận dài
        info = ttk.LabelFrame(parent, text="Thông tin đã lưu", padding=8)
        info.pack(fill="x", pady=(0, 8))
        rows = [
            ("Độ khẩn", self.cfg.get("priority")),
            ("Độ mật", self.cfg.get("security")),
            ("Nội dung phiếu", self.cfg.get("report_content")),
            ("Luồng trình", self.cfg.get("flow_name")),
            ("Hồ sơ công việc", self.cfg.get("work_profile_name")),
        ]
        for label, val in rows:
            r = ttk.Frame(info); r.pack(fill="x", anchor="w", pady=1)
            ttk.Label(r, text=label + ":", width=13, foreground="gray").pack(side="left", anchor="n")
            ttk.Label(r, text=val or "(trống)", wraplength=200, justify="left").pack(
                side="left", fill="x", expand=True)

        documents = self.cfg.get("documents") or []
        if documents:
            docs_frame = ttk.LabelFrame(parent, text=f"Văn bản ({len(documents)})", padding=8)
            docs_frame.pack(fill="x", pady=(0, 8))
            for i, doc in enumerate(documents):
                r = ttk.Frame(docs_frame); r.pack(fill="x", anchor="w", pady=(0 if i == 0 else 6, 0))
                summary = f"{doc.get('doc_type') or '(chưa chọn loại)'} — {doc.get('code') or '(chưa có số)'}"
                ttk.Label(r, text=f"VB {i+1}:", width=13, foreground="gray").pack(side="left", anchor="n")
                ttk.Label(r, text=summary, wraplength=200, justify="left").pack(
                    side="left", fill="x", expand=True)
                if doc.get("abstract"):
                    r2 = ttk.Frame(docs_frame); r2.pack(fill="x", anchor="w")
                    ttk.Label(r2, text="", width=13).pack(side="left")
                    ttk.Label(r2, text=doc["abstract"], wraplength=200, justify="left",
                              foreground="#444").pack(side="left", fill="x", expand=True)

        cats = [("recv_inside", "Nhận nội bộ"), ("recv_report", "Báo cáo"), ("recv_edoc", "Liên thông"),
                ("recv_save", "Nơi lưu"), ("recv_know", "Để biết")]
        any_recv = any(self.cfg.get(k) for k, _ in cats)
        if any_recv:
            rec = ttk.LabelFrame(parent, text="Nơi nhận", padding=8)
            rec.pack(fill="x", pady=(0, 8))
            for key, label in cats:
                nodes = self.cfg.get(key) or []
                if not nodes:
                    continue
                names = ", ".join(nd["name"].strip() for nd in nodes)
                r = ttk.Frame(rec); r.pack(fill="x", anchor="w", pady=1)
                ttk.Label(r, text=label + ":", width=13, foreground="gray").pack(side="left", anchor="n")
                ttk.Label(r, text=names, wraplength=200, justify="left").pack(
                    side="left", fill="x", expand=True)

    def _build_tree(self, parent, cfg):
        box = ttk.LabelFrame(parent, text="File", padding=4)
        box.pack(fill="both", expand=True)
        tree_wrap = ttk.Frame(box); tree_wrap.pack(fill="both", expand=True)
        tsb = ttk.Scrollbar(tree_wrap, orient="vertical")
        self.tree = ttk.Treeview(tree_wrap, show="tree", height=14, yscrollcommand=tsb.set)
        tsb.config(command=self.tree.yview)
        tsb.pack(side="right", fill="y")
        self.tree.pack(side="left", fill="both", expand=True)
        bind_mousewheel_scroll(self.tree)

        def add_group(title, entries):
            gid = self.tree.insert("", "end", text=title, open=True)
            for path, kind in entries:
                if not path:
                    continue
                iid = self.tree.insert(gid, "end", text=os.path.basename(path))
                self._tree_paths[iid] = (path, kind)

        add_group("📄 PHIẾU TRÌNH (không gửi đi)",
                   [(cfg.get("file_report_main"), "report_main")] +
                   [(p, "report_extra") for p in (cfg.get("files_report_extra") or [])])
        documents = cfg.get("documents") or []
        n_docs = len(documents)
        for i, doc in enumerate(documents):
            title = f"📤 VĂN BẢN {i+1}/{n_docs} (gửi đi)" if n_docs > 1 else "📤 VĂN BẢN (gửi đi)"
            add_group(title,
                      [(doc.get("file_draft_main"), f"draft_main:{i}")] +
                      [(p, f"draft_extra:{i}") for p in (doc.get("files_draft_extra") or [])])

        self.tree.bind("<<TreeviewSelect>>", self._on_tree_select)
        ttk.Label(parent, foreground="gray", wraplength=320, justify="left",
                  text="PDF của Phiếu trình/Dự thảo: kéo dấu để di chuyển, nhấp đúp để sửa số/xoá, "
                       "bấm '➕ Thêm dấu' rồi nhấp vào trang để thêm. File khác/PDF phụ: bấm để mở "
                       "thư mục chứa file.").pack(anchor="w", pady=(6, 0))

    def _on_tree_select(self, _e):
        sel = self.tree.selection()
        if not sel or sel[0] not in self._tree_paths:
            return
        path, kind = self._tree_paths[sel[0]]
        self._select_file(path, kind)

    # ---------- Phải: khung xem PDF ----------
    def _build_viewer(self, parent):
        bar = ttk.Frame(parent); bar.pack(fill="x")
        ttk.Button(bar, text="◀", width=3, command=self._prev_page).pack(side="left")
        self.page_var = tk.StringVar(value="—")
        ttk.Label(bar, textvariable=self.page_var).pack(side="left", padx=4)
        ttk.Button(bar, text="▶", width=3, command=self._next_page).pack(side="left")
        self.btn_addmark = ttk.Button(bar, text="➕ Thêm dấu", command=self._toggle_add_mode, state="disabled")
        self.btn_addmark.pack(side="left", padx=12)

        wrap = ttk.Frame(parent); wrap.pack(fill="both", expand=True, pady=(6, 0))
        xsb = ttk.Scrollbar(wrap, orient="horizontal")
        ysb = ttk.Scrollbar(wrap, orient="vertical")
        self.canvas = tk.Canvas(wrap, bg="#5c5c5c",
                                 xscrollcommand=xsb.set, yscrollcommand=ysb.set)
        xsb.config(command=self.canvas.xview); ysb.config(command=self.canvas.yview)
        ysb.pack(side="right", fill="y"); xsb.pack(side="bottom", fill="x")
        self.canvas.pack(side="left", fill="both", expand=True)
        self.canvas.bind("<Button-1>", self._canvas_click)
        bind_mousewheel_scroll(self.canvas, horizontal=True)

        self._page_img = None   # giữ tham chiếu tránh bị garbage-collect

    # ---------- Chọn file trong cây ----------
    def _select_file(self, path, kind):
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            self._open_folder(path)
            return
        try:
            doc = fitz.open(path)
        except Exception as e:
            messagebox.showerror("Lỗi mở PDF", f"Không mở được file:\n{e}", parent=self)
            return
        if self._current_doc is not None:
            try: self._current_doc.close()
            except Exception: pass
        self._current_doc = doc
        self._current_path = path
        self._current_kind = kind
        self._current_page = 0
        is_main = kind == "report_main" or kind.startswith("draft_main")
        self._editable = is_main and self.cfg.get("auto_stamp", True)
        self._add_mode = False
        self.btn_addmark.config(state=("normal" if self._editable else "disabled"),
                                 text="➕ Thêm dấu")
        self._render_current()
        if is_main:
            self._ensure_scanned(path)
        else:
            self._draw_marks()

    def _open_folder(self, path):
        try:
            if sys.platform == "win32":
                subprocess.Popen(["explorer", f"/select,{os.path.normpath(path)}"])
            elif sys.platform == "darwin":
                subprocess.Popen(["open", "-R", path])   # hiện trong Finder, tương đương /select trên Windows
            else:
                subprocess.Popen(["xdg-open", os.path.dirname(path)])
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không mở được thư mục:\n{e}", parent=self)

    # ---------- Vẽ trang ----------
    def _render_current(self):
        if self._current_doc is None:
            return
        page = self._current_doc[self._current_page]
        target_w = max(self.canvas.winfo_width(), 640)
        self._zoom = max(0.5, min(2.2, target_w / max(page.rect.width, 1)))
        pix = page.get_pixmap(matrix=fitz.Matrix(self._zoom, self._zoom))
        self._page_img = tk.PhotoImage(data=pix.tobytes("ppm"))
        self.canvas.delete("page")
        self.canvas.create_image(0, 0, anchor="nw", image=self._page_img, tags="page")
        self.canvas.config(scrollregion=(0, 0, pix.width, pix.height))
        self.page_var.set(f"Trang {self._current_page + 1}/{len(self._current_doc)}")
        self._draw_marks()

    def _prev_page(self):
        if self._current_doc and self._current_page > 0:
            self._current_page -= 1
            self._render_current()

    def _next_page(self):
        if self._current_doc and self._current_page < len(self._current_doc) - 1:
            self._current_page += 1
            self._render_current()

    # ---------- Quét vị trí ký (nền) ----------
    def _ensure_scanned(self, path):
        if path in self._stamps_cache:
            if path == self._current_path:
                self._draw_marks()
            return
        self.status_var.set(f"Đang quét vị trí ký: {os.path.basename(path)}…")

        def worker():
            try:
                d = fitz.open(path)
                try:
                    stamps = find_signature_stamps(d, self.cfg.get("flow_nodes_override"), self._plog)
                finally:
                    d.close()
            except Exception as e:
                err_msg = str(e)   # tính ngay trong khối except — "e" bị Python xoá khi except kết thúc
                self.after(0, lambda: self._scan_failed(path, err_msg))
                return
            self.after(0, lambda: self._scan_done(path, stamps))
        threading.Thread(target=worker, daemon=True).start()

    def _scan_done(self, path, stamps):
        self._stamps_cache[path] = stamps
        self.status_var.set("")
        if path == self._current_path:
            self._draw_marks()

    def _scan_failed(self, path, e):
        self.status_var.set("")
        self._plog(f"✖ Không quét được vị trí ký trong {os.path.basename(path)}: {e}")

    # ---------- Vẽ + tương tác các dấu ----------
    def _draw_marks(self):
        self.canvas.delete("markgrp")
        self._mark_items = {}
        if not self._editable or self._current_path not in self._stamps_cache:
            return
        stamps = self._stamps_cache[self._current_path]
        r = self.MARK_R
        for i, st in enumerate(stamps):
            if st["page"] != self._current_page:
                continue
            cx, cy = st["x"] * self._zoom, st["y"] * self._zoom
            oval = self.canvas.create_oval(cx - r, cy - r, cx + r, cy + r,
                                            fill="#ffca28", outline="#e65100", width=2)
            txt = self.canvas.create_text(cx, cy, text=str(st["number"]),
                                           fill="#000", font=("", 10, "bold"))
            tag = f"mark{i}"
            self.canvas.itemconfig(oval, tags=("markgrp", tag))
            self.canvas.itemconfig(txt, tags=("markgrp", tag))
            self._mark_items[i] = (oval, txt)
            self.canvas.tag_bind(tag, "<ButtonPress-1>", lambda e, idx=i: self._mark_press(e, idx))
            self.canvas.tag_bind(tag, "<B1-Motion>", lambda e, idx=i: self._mark_drag(e, idx))
            self.canvas.tag_bind(tag, "<ButtonRelease-1>", lambda e, idx=i: self._mark_release(e, idx))
            self.canvas.tag_bind(tag, "<Double-Button-1>", lambda e, idx=i: self._mark_edit(idx))

    def _mark_press(self, e, idx):
        self._drag_idx = idx
        self._drag_last = (e.x, e.y)

    def _mark_drag(self, e, idx):
        if self._drag_idx != idx:
            return
        dx, dy = e.x - self._drag_last[0], e.y - self._drag_last[1]
        oval, txt = self._mark_items[idx]
        self.canvas.move(oval, dx, dy)
        self.canvas.move(txt, dx, dy)
        self._drag_last = (e.x, e.y)

    def _mark_release(self, e, idx):
        if self._drag_idx != idx:
            return
        self._drag_idx = None
        oval, _txt = self._mark_items[idx]
        x0, y0, x1, y1 = self.canvas.coords(oval)
        cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
        st = self._stamps_cache[self._current_path][idx]
        st["x"], st["y"] = cx / self._zoom, cy / self._zoom

    def _mark_edit(self, idx):
        self._open_role_dialog(idx=idx)

    # ---------- Thêm dấu mới ----------
    def _toggle_add_mode(self):
        self._add_mode = not self._add_mode
        self.btn_addmark.config(text="✋ Nhấp vào trang…" if self._add_mode else "➕ Thêm dấu")

    def _canvas_click(self, e):
        if not self._add_mode:
            return
        if "markgrp" in self.canvas.gettags("current"):
            return   # nhấp trúng 1 dấu có sẵn — để binding của dấu đó xử lý
        self._add_mode = False
        self.btn_addmark.config(text="➕ Thêm dấu")
        px, py = self.canvas.canvasx(e.x) / self._zoom, self.canvas.canvasy(e.y) / self._zoom
        self._open_role_dialog(new=True, page=self._current_page, x=px, y=py)

    # ---------- Hộp thoại chọn chức danh/số ----------
    def _open_role_dialog(self, idx=None, new=False, page=None, x=None, y=None):
        stamps = self._stamps_cache.setdefault(self._current_path, [])
        st = None if new else stamps[idx]

        role_map = build_role_number_map(self.cfg.get("flow_nodes_override"))
        role_keys = list(role_map.keys()) or [SIG_CHUYEN_VIEN_ROLE]

        dlg = tk.Toplevel(self)
        dlg.title("Thêm dấu ký" if new else "Sửa dấu ký")
        dlg.transient(self); dlg.grab_set()
        ttk.Label(dlg, text="Chức danh:").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        role_var = tk.StringVar(value=(st["role"] if st and st.get("role") in role_keys else role_keys[0]))
        combo = ttk.Combobox(dlg, values=role_keys, textvariable=role_var, state="readonly", width=22)
        combo.grid(row=0, column=1, padx=8, pady=6)

        ttk.Label(dlg, text="Số:").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        num_var = tk.StringVar(value=(str(st["number"]) if st else ""))
        ttk.Entry(dlg, textvariable=num_var, width=8).grid(row=1, column=1, sticky="w", padx=8, pady=6)

        def _fill_number(*_):
            n = role_map.get(role_var.get())
            num_var.set(str(n) if n is not None else "")
        combo.bind("<<ComboboxSelected>>", _fill_number)
        if new:
            _fill_number()

        btns = ttk.Frame(dlg); btns.grid(row=2, column=0, columnspan=2, pady=8)

        def _close_dlg():
            # Giải phóng "grab" TƯỜNG MINH trước khi đóng — trên macOS, destroy() không phải
            # lúc nào cũng tự giải phóng grab sạch sẽ, để lại cả ứng dụng bị khoá input (mọi ô
            # nhập khác, kể cả ở màn hình chính, ngừng nhận phím) dù cửa sổ đã đóng từ lâu.
            try:
                dlg.grab_release()
            except Exception:
                pass
            dlg.destroy()

        def _ok():
            try:
                number = int(num_var.get())
            except ValueError:
                messagebox.showwarning("Số không hợp lệ", "Số phải là số nguyên.", parent=dlg)
                return
            role = role_var.get()
            if new:
                stamps.append({"page": page, "x": x, "y": y, "number": number, "role": role,
                                "title": f"{role} (thêm tay)"})
            else:
                st["number"] = number
                st["role"] = role
            _close_dlg()
            self._draw_marks()
        dlg.protocol("WM_DELETE_WINDOW", _close_dlg)   # bấm nút [X] của cửa sổ cũng phải giải phóng grab
        ttk.Button(btns, text="OK", command=_ok).pack(side="left", padx=4)
        if not new:
            def _del():
                del stamps[idx]
                _close_dlg()
                self._draw_marks()
            ttk.Button(btns, text="Xóa", command=_del).pack(side="left", padx=4)
        ttk.Button(btns, text="Hủy", command=_close_dlg).pack(side="left", padx=4)

    # ---------- Tiến trình + kết quả (checklist/banner thân thiện — xem StepChecklist) ----------
    def _toggle_log(self):
        if self.logbox.winfo_ismapped():
            self.logbox.pack_forget()
            self.btn_toggle_log.config(text="▾ Xem chi tiết kỹ thuật")
        else:
            self.logbox.pack(fill="x", side="bottom", padx=10, pady=(0, 10))
            self.btn_toggle_log.config(text="▴ Ẩn chi tiết kỹ thuật")

    def _set_banner(self, kind, lines):
        bg, fg = self.BANNER_STYLE[kind]
        self.banner_var.set("\n".join(lines))
        self.banner_label.config(bg=bg, fg=fg)
        self.banner_label.pack(fill="x", pady=(6, 0))

    def _set_phase(self, key):
        self._current_phase = key
        self.checklist.mark_running(key)
        if self._log_window is not None:
            label = dict(PIPELINE_PHASES).get(key, key)
            self._log_window.append(f"• {label}...")

    def _render_result_banner(self, result):
        """`result`: dict trả về từ run_pipeline (xem verify_report_saved) — KHÔNG suy đoán từ
        mã HTTP, chỉ hiện những gì đã tự đọc lại và xác nhận được với server."""
        sign = result.get("submit_sign")
        if not result.get("verified"):
            self._set_banner("ambiguous", [
                "⚠ ĐÃ GỬI YÊU CẦU, NHƯNG CHƯA XÁC MINH ĐƯỢC",
                "Hệ thống chưa phản hồi rõ trong lúc kiểm tra lại — nhiều khả năng vẫn đã thành công.",
                "→ Mở web, vào thùng nháp/đang trình để tự kiểm tra trước khi thử gửi lại (tránh tạo trùng).",
            ])
            return
        if sign != "1":
            self._set_banner("success", [
                "✅ ĐÃ LƯU NHÁP THÀNH CÔNG",
                "Hệ thống đã xác nhận phiếu trình có tồn tại (chưa gửi cho ai).",
                "→ Mở web, vào thùng nháp để trình khi sẵn sàng.",
            ])
            return
        if not result.get("in_process"):
            # Thấy phiếu trình tồn tại (khớp content+thời gian) nhưng KHÔNG thấy trong đúng hộp
            # "đang trình" — với sign=1 lẽ ra phải thấy ở cả 2 nơi. Đáng ngờ riêng, không nên
            # gộp chung với "✅ thành công" — có thể lệnh Trình chưa thật sự vào luồng.
            self._set_banner("ambiguous", [
                "⚠ ĐÃ LƯU, NHƯNG CHƯA THẤY TRONG HỘP \"ĐANG TRÌNH\"",
                "Hệ thống xác nhận phiếu trình tồn tại, nhưng chưa thấy nó ở danh sách đang xử lý "
                "— có thể chỉ là chậm cập nhật, cũng có thể lệnh Trình chưa thật sự vào luồng.",
                "→ Mở web, vào hộp \"đang trình\" kiểm tra lại cho chắc trước khi coi là xong.",
            ])
            return
        line2 = "Hệ thống đã xác nhận phiếu trình đã vào luồng ký duyệt"
        note = result.get("history_note")
        if note and note.get("createAt"):
            line2 += f", lúc {note['createAt'][11:16]}"
        line2 += "."
        lines = ["✅ ĐÃ TRÌNH THÀNH CÔNG", line2]
        pending = result.get("pending")
        if pending:
            who = pending.get("receiveUser") or "?"
            role = pending.get("displayPositionName") or ""
            lines.append(f"Đang chờ: {who}" + (f" ({role})" if role else "") + " xử lý.")
        self._set_banner("success", lines)

    def _render_error_banner(self, err_text):
        phase_label = dict(PIPELINE_PHASES).get(self._current_phase, self._current_phase or "?")
        if self._current_phase:
            self.checklist.mark_error(self._current_phase)
        self._set_banner("error", [
            "❌ CHƯA HOÀN TẤT",
            f"Mắc ở bước: {phase_label}",
            f"Lý do: {err_text[:300]}",
            "→ Xem \"chi tiết kỹ thuật\" bên dưới để rõ hơn. Nếu bước \"Lưu văn bản\" đã xong "
            "trước khi lỗi, kiểm tra kỹ thùng nháp trên web trước khi thử lại, tránh tạo trùng.",
        ])

    # ---------- Gửi thật ----------
    def _plog(self, msg):
        # Log kỹ thuật (HTTP, phản hồi server...) chỉ vào logbox ẩn ("Xem chi tiết kỹ thuật")
        # như cũ — SendLogWindow chỉ hiện đúng các bước lớn (xem _set_phase), không lặp lại log
        # chi tiết ở đây để giữ cửa sổ đó gọn, dễ đọc.
        self.logbox.insert("end", msg + "\n"); self.logbox.see("end"); self.update_idletasks()

    def _send(self, sign):
        cfg = dict(self.cfg)
        cfg["submit_sign"] = sign   # "0" = Lưu dự thảo, "1" = Trình văn bản (xem run_pipeline)
        cfg["stamps_report_override"] = self._stamps_cache.get(cfg.get("file_report_main"))
        docs = []
        for doc in (self.cfg.get("documents") or []):
            d = dict(doc)
            main = d.get("file_draft_main")
            if main:
                d["stamps_override"] = self._stamps_cache.get(main)
            docs.append(d)
        cfg["documents"] = docs
        self.btn_send.config(state="disabled")
        self.btn_submit.config(state="disabled")
        self._sending = True
        self._current_phase = None
        self.checklist.reset()
        self.banner_label.pack_forget()
        self.logbox.delete("1.0", "end")
        # Cửa sổ nhật ký riêng (xem SendLogWindow) — độc lập với PreviewWindow, để thấy tiến
        # trình dù cửa sổ này có bị che khuất; _plog/_set_phase tự mirror sang đây bên dưới.
        title = "Nhật ký — Trình văn bản" if sign == "1" else "Nhật ký — Lưu dự thảo"
        self._log_window = SendLogWindow(self, title, self._finish_post_send)
        self._plog("=== BẮT ĐẦU TRÌNH ===" if sign == "1" else "=== BẮT ĐẦU LƯU DỰ THẢO ===")

        def phase_cb(key):
            self.after(0, lambda k=key: self._set_phase(k))

        def worker():
            try:
                result = run_pipeline(self.session, cfg, self._plog, check_only=False, phase_cb=phase_cb)
            except PipelineError as e:
                err_text = str(e)   # tính ngay trong khối except — "e" bị Python xoá khi except kết thúc
                self._plog("\n✖ DỪNG: " + err_text)
                self.after(0, lambda: self._render_error_banner(err_text))
                self.after(0, lambda: self._start_post_send_close(None))
            except Exception as e:
                err_text = repr(e)
                self._plog("\n✖ LỖI KHÔNG NGỜ: " + err_text)
                self.after(0, lambda: self._render_error_banner(err_text))
                self.after(0, lambda: self._start_post_send_close(None))
            else:
                self._submitted_ok = True   # đã ghi lên hệ thống thật (kể cả banner "ambiguous")
                self.after(0, lambda: (self.checklist.complete_all(), self._render_result_banner(result)))
                self.after(0, lambda: self._start_post_send_close(result))
            finally:
                self._sending = False
                self.after(0, lambda: self.btn_send.config(state="normal"))
                self.after(0, lambda: self.btn_submit.config(state="normal"))
        threading.Thread(target=worker, daemon=True).start()

    def _start_post_send_close(self, result):
        """Gọi ngay sau khi có banner kết quả (thành công/ambiguous, `result` là dict trả về từ
        run_pipeline) hoặc lỗi (`result=None`) — quyết định tab đích bên "Quản lý Phiếu trình"
        + có tự mở trình duyệt hay không, rồi cho SendLogWindow đếm ngược tự đóng (kéo theo
        đóng luôn PreviewWindow, xem _finish_post_send)."""
        self._post_send_target_tab = (
            "processing" if (result and result.get("submit_sign") == "1"
                              and result.get("verified") and result.get("in_process"))
            else "draft")
        if result is not None:
            import webbrowser
            # KHÔNG có link tĩnh dẫn thẳng vào đúng phiếu trình (đã kiểm qua HAR thật: reportId
            # chỉ xuất hiện trong AJAX ngầm kèm token dùng 1 lần, không hề có trên URL nhìn thấy
            # được) — mở đúng MÀN HÌNH "Quản lý phiếu trình" trên web, tự tìm phiếu trong đó.
            url = BASE + "/Index.do?request_locale=en_US&mainMenu=3&trId=2.2"
            try:
                webbrowser.open(url)
                self._log_window.append("\n• Đã mở trình duyệt tới màn hình \"Quản lý phiếu trình\" trên web.")
            except Exception as e:
                self._log_window.append(f"\n• Không mở được trình duyệt: {e!r}")
        else:
            self._log_window.append("\n• Chưa hoàn tất — không tự mở trình duyệt.")
        self._log_window.start_close_countdown(5)

    def _finish_post_send(self):
        """Gọi khi SendLogWindow đóng (tự đếm ngược hoặc bấm tay Đóng) — đóng luôn PreviewWindow
        rồi nhảy sang đúng tab/sub-tab bên "Quản lý Phiếu trình" (xem _start_post_send_close)."""
        target_tab = self._post_send_target_tab
        self._on_close()
        self.master._open_manage_reports_tab(target_tab)

    def _on_close(self):
        if self._sending:
            if not messagebox.askyesno(
                    "Đang gửi",
                    "Đang gửi phiếu trình — đóng cửa sổ này thì việc gửi vẫn tiếp tục chạy ngầm, "
                    "nhưng bạn sẽ không thấy kết quả nữa. Vẫn đóng?", parent=self):
                return
        if self._log_window is not None:
            # Đóng PreviewWindow bằng đường khác (nút "Đóng" chính, hoặc bấm X) trong lúc
            # SendLogWindow vẫn còn mở (giữa chừng gửi, hoặc đang đếm ngược) — đóng nốt luôn,
            # tránh để cửa sổ nhật ký mồ côi không còn PreviewWindow phía sau.
            try: self._log_window.destroy()
            except Exception: pass
        if self._current_doc is not None:
            try: self._current_doc.close()
            except Exception: pass
        master = self.master
        submitted_ok = self._submitted_ok
        self.destroy()
        # Trên macOS, đóng 1 Toplevel không tự trả bàn phím về đúng cửa sổ cha — màn hình chính
        # nhìn vẫn "active" nhưng không nhận phím nữa (vd ô "Nơi nhận" gõ không ăn) cho tới khi
        # tự bấm vào. Ép lại tường minh để khỏi phải bấm.
        try:
            master.lift()
            master.focus_force()
        except Exception:
            pass
        if submitted_ok:
            # Đã gửi thành công (Lưu dự thảo/Trình văn bản) — làm mới form chính để tránh dữ
            # liệu của phiếu này còn dính lại khi bắt đầu điền phiếu tiếp theo.
            master._reset_form()


if __name__ == "__main__":
    App().mainloop()
